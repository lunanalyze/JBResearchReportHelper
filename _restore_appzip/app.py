from __future__ import annotations

import datetime as dt
import base64
import ctypes
import ctypes.wintypes
import json
import os
import re
import shutil
import sys
import threading
import time
import traceback
import urllib.parse
import urllib.request
import webbrowser
import csv
import queue
import uuid
from copy import deepcopy
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

sys.dont_write_bytecode = True

import collector
import paths
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from openpyxl import Workbook


HOST = "127.0.0.1"
PORT = 8765
APP_DIR = paths.APP_DIR
APP_DATA_DIR = paths.APP_DATA_DIR
CONFIG_DIR = paths.CONFIG_DIR
PROMPT_PATH = paths.RESOURCES_DIR / "report_prompt.md"
TEMPLATE_PATH = paths.RESOURCES_DIR / "report_template.docx"
KEY_PATH = CONFIG_DIR / "openai_key.bin"
NEWS_KEYWORDS_PATH = CONFIG_DIR / "news_keywords.json"
OPENAI_MODEL = "gpt-5.4-mini"
DEFAULT_REPORT_PROMPT = """Return only one JSON object for filling the weekly report template.
Use the provided metadata and attached PDFs.

Required top-level keys:
- REPORT_WEEK
- DEPARTMENT_NAME
- BANK_SECTION_PAGE
- AGENCY_SECTION_PAGE
- RESEARCH_SECTION_PAGE
- BANK_DETAIL_ITEMS
- AGENCY_DETAIL_ITEMS
- RESEARCH_DETAIL_ITEMS

Each DETAIL_ITEMS value should be an array of objects with ITEM_NO, TITLE, URL, SOURCE_NAME, PUBLISHED_MM_DD, SUMMARY_BULLET_1, SUMMARY_BULLET_2, SUMMARY_BULLET_3.
"""

paths.ensure_app_dirs()
paths.migrate_legacy_data()
paths.copy_default_resource("report_prompt.md", DEFAULT_REPORT_PROMPT)
paths.copy_default_resource("report_template.docx")

STATE_LOCK = threading.Lock()
STATE: dict = {"items": [], "run_dir": None, "report_path": ""}
JOBS: dict[str, dict] = {}


INDEX_HTML = r"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>조사연구 도우미</title>
  <style>
    :root {
      color-scheme: light;
      --bg: #f6f7f9;
      --panel: #ffffff;
      --text: #17202a;
      --muted: #6b7280;
      --line: #d8dee8;
      --brand: #1f6feb;
      --brand-dark: #1557bd;
      --ok: #0f7b3d;
      --warn: #b45309;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: "Malgun Gothic", "Segoe UI", sans-serif;
      background: var(--bg);
      color: var(--text);
    }
    header {
      padding: 18px 24px 14px;
      background: var(--panel);
      border-bottom: 1px solid var(--line);
    }
    .brand {
      display: flex;
      align-items: center;
      gap: 12px;
    }
    .brand-mark {
      width: 38px;
      height: 38px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      border-radius: 8px;
      background: #eef4ff;
      color: #1557bd;
      border: 1px solid #c7ddff;
      font-size: 16px;
      font-weight: 900;
    }
    .brand-copy {
      min-width: 0;
    }
    h1 {
      margin: 0;
      font-size: 21px;
      font-weight: 700;
      letter-spacing: 0;
    }
    main { padding: 18px 24px 28px; }
    .controls {
      display: grid;
      grid-template-columns: repeat(3, minmax(260px, 1fr));
      gap: 12px;
      align-items: stretch;
      margin-bottom: 14px;
    }
    .group {
      min-height: 238px;
      display: flex;
      flex-direction: column;
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 12px;
    }
    .group-title {
      font-size: 13px;
      font-weight: 700;
      margin-bottom: 10px;
    }
    .group-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      margin-bottom: 10px;
      min-height: 38px;
    }
    .group-head .group-title {
      margin-bottom: 0;
    }
    .header-action {
      height: 38px;
      min-width: 96px;
      padding: 0 10px;
      font-size: 13px;
    }
    .header-placeholder {
      width: 96px;
      height: 38px;
      flex: 0 0 auto;
    }
    .date-row {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 8px;
    }
    .limit-row {
      margin-top: 10px;
    }
    .card-spacer {
      display: none;
    }
    .button-row {
      display: flex;
      align-items: center;
      justify-content: flex-start;
      gap: 8px;
      margin-top: auto;
      padding-top: 10px;
    }
    .button-row button {
      width: auto;
      min-width: 96px;
      padding: 0 10px;
    }
    .collect-row {
      margin-top: auto;
      padding-top: 10px;
    }
    .collect-row .button-row {
      margin-top: 0;
      padding-top: 0;
    }
    label {
      display: block;
      font-size: 12px;
      color: var(--muted);
    }
    input[type="date"], input[type="number"] {
      width: 100%;
      margin-top: 4px;
      padding: 8px 9px;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
      background: #fff;
    }
    .side {
      display: flex;
      justify-content: flex-end;
      margin: -2px 0 14px;
    }
    .side button {
      min-width: 104px;
      padding: 0 12px;
    }
    button {
      height: 38px;
      border: 1px solid var(--brand);
      border-radius: 6px;
      background: var(--brand);
      color: #fff;
      font-weight: 700;
      cursor: pointer;
    }
    button:hover { background: var(--brand-dark); }
    button.secondary {
      background: var(--brand);
      color: #fff;
      border-color: var(--brand);
    }
    button.secondary:hover { background: var(--brand-dark); }
    button.jb {
      background: #eef4ff;
      color: #1557bd;
      border-color: #c7ddff;
    }
    button.jb:hover {
      background: #dceaff;
      border-color: #9cc5ff;
    }
    button.ghost {
      background: #fff;
      color: var(--text);
      border-color: var(--line);
    }
    button.ghost:hover {
      background: #eef2f7;
    }
    button.danger {
      background: var(--brand);
      color: #fff;
      border-color: var(--brand);
    }
    button.danger:hover { background: var(--brand-dark); }
    button:disabled {
      opacity: 0.55;
      cursor: not-allowed;
    }
    .status {
      min-height: 28px;
      margin: 8px 0 12px;
      color: var(--muted);
      font-size: 13px;
    }
    .toolbar {
      display: flex;
      gap: 8px;
      align-items: center;
      justify-content: space-between;
      margin: 10px 0;
    }
    .summary {
      color: var(--muted);
      font-size: 13px;
    }
    .table-wrap {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: auto;
      max-height: calc(100vh - 270px);
      cursor: default;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
      cursor: default;
    }
    th, td {
      border-bottom: 1px solid var(--line);
      padding: 8px 10px;
      vertical-align: middle;
      white-space: nowrap;
      cursor: default;
    }
    th {
      position: sticky;
      top: 0;
      background: #f1f4f8;
      text-align: left;
      z-index: 1;
    }
    td.title {
      min-width: 420px;
      max-width: 760px;
      white-space: normal;
      line-height: 1.35;
    }
    .action-cell {
      width: 64px;
      text-align: center;
    }
    .open-item-btn {
      width: 34px;
      height: 30px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      padding: 0;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .open-item-btn:hover {
      background: #eef4ff;
      border-color: #9cc5ff;
      color: #1557bd;
    }
    .open-item-btn svg {
      width: 17px;
      height: 17px;
      stroke: currentColor;
      stroke-width: 2.2;
      fill: none;
    }
    tr:hover { background: #f7fbff; }
    .tabs {
      display: flex;
      flex-wrap: wrap;
      gap: 6px;
      margin: 8px 0 10px;
    }
    .tab {
      height: 32px;
      padding: 0 11px;
      border: 1px solid var(--line);
      border-radius: 999px;
      background: #fff;
      color: var(--text);
      font-weight: 600;
      cursor: pointer;
    }
    .tab.active {
      background: #e8f1ff;
      border-color: #9cc5ff;
      color: #1557bd;
    }
    .badge {
      display: inline-flex;
      min-width: 24px;
      height: 22px;
      align-items: center;
      justify-content: center;
      border-radius: 999px;
      background: #eef4ff;
      color: #1d4ed8;
      font-weight: 700;
      font-size: 12px;
    }
    .kind {
      color: var(--muted);
      font-size: 12px;
    }
    .modal-backdrop {
      position: fixed;
      inset: 0;
      display: none;
      align-items: center;
      justify-content: center;
      padding: 24px;
      background: rgba(15, 23, 42, 0.42);
      z-index: 20;
    }
    .modal-backdrop.open { display: flex; }
    .modal {
      width: min(760px, 100%);
      max-height: min(720px, calc(100vh - 48px));
      display: grid;
      grid-template-rows: auto auto 1fr;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: 0 24px 70px rgba(15, 23, 42, 0.28);
      overflow: hidden;
    }
    .modal-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 16px 18px;
      border-bottom: 1px solid var(--line);
    }
    .modal-title {
      font-size: 17px;
      font-weight: 800;
    }
    .modal-close {
      width: 34px;
      height: 34px;
      padding: 0;
      background: #fff;
      color: var(--text);
      border-color: var(--line);
      font-size: 18px;
      line-height: 1;
    }
    .modal-close:hover { background: #eef2f7; }
    .modal-help {
      padding: 10px 18px;
      color: var(--muted);
      font-size: 13px;
    }
    .history-list {
      overflow: auto;
      padding: 8px;
    }
    .history-item {
      width: 100%;
      height: auto;
      display: grid;
      grid-template-columns: 1fr auto;
      gap: 12px;
      align-items: center;
      padding: 12px 14px;
      margin: 0 0 6px;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 8px;
      text-align: left;
      font-weight: 600;
    }
    .history-actions {
      display: inline-flex;
      gap: 8px;
      align-items: center;
      justify-content: flex-end;
      flex-wrap: wrap;
    }
    .history-item:hover {
      background: #f7fbff;
      border-color: #9cc5ff;
    }
    .history-name {
      font-size: 14px;
      font-weight: 800;
    }
    .history-meta {
      margin-top: 4px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 500;
    }
    .history-tags {
      display: inline-flex;
      gap: 6px;
      align-items: center;
      justify-content: flex-end;
      flex-wrap: wrap;
    }
    .history-tag {
      display: inline-flex;
      min-height: 24px;
      align-items: center;
      padding: 3px 8px;
      border-radius: 999px;
      background: #eef4ff;
      color: #1d4ed8;
      font-size: 12px;
      font-weight: 800;
    }
    .history-tag.report {
      background: #ecfdf3;
      color: #0f7b3d;
    }
    .history-delete {
      height: 30px;
      padding: 0 10px;
      background: #fff;
      color: #b42318;
      border: 1px solid #f1b9b4;
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .history-delete:hover {
      background: #fff1f0;
      border-color: #e0776e;
    }
    .danger-outline {
      background: #fff;
      color: #b42318;
      border: 1px solid #f1b9b4;
    }
    .danger-outline:hover {
      background: #fff1f0;
      border-color: #e0776e;
      color: #b42318;
    }
    .keyword-modal {
      width: min(720px, calc(100vw - 40px));
    }
    .keyword-body {
      padding: 20px 22px 22px;
    }
    .keyword-editor {
      display: grid;
      grid-template-columns: minmax(0, 1fr) 86px;
      gap: 10px;
      margin-bottom: 14px;
    }
    .keyword-editor input {
      width: 100%;
      height: 38px;
      padding: 0 12px;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
    }
    .keyword-tabs {
      display: flex;
      gap: 8px;
      margin-bottom: 14px;
    }
    .keyword-tab {
      height: 32px;
      padding: 0 10px;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .keyword-tab.active {
      color: #fff;
      border-color: var(--brand);
      background: var(--brand);
    }
    .keyword-list {
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      max-height: 300px;
      overflow: auto;
      padding: 12px;
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #f8fafc;
    }
    .keyword-item {
      width: auto;
      min-height: 32px;
      padding: 0 11px;
      text-align: center;
      color: var(--text);
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 999px;
      font-size: 13px;
      font-weight: 700;
    }
    .keyword-item.selected {
      background: #eef4ff;
      border-color: #1f6feb;
      color: #174ea6;
    }
    .source-list {
      display: grid;
      gap: 8px;
      max-height: 340px;
      overflow: auto;
    }
    .source-item {
      min-height: 42px;
      display: flex;
      align-items: center;
      padding: 10px 12px;
      color: var(--text);
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
      font-weight: 700;
    }
    .modal-actions {
      display: flex;
      justify-content: space-between;
      gap: 8px;
      margin-top: 16px;
    }
    .empty-state {
      padding: 36px 18px;
      color: var(--muted);
      text-align: center;
      font-size: 13px;
    }
    @media (max-width: 1100px) {
      .controls { grid-template-columns: 1fr; }
    }
  </style>
</head>
<body>
  <header>
    <div class="brand">
      <div class="brand-mark">JB</div>
      <div class="brand-copy">
        <h1>조사연구 도우미</h1>
      </div>
    </div>
  </header>
  <main>
    <section class="controls">
      <div class="group">
        <div class="group-head">
          <div class="group-title">뉴스</div>
          <button class="jb header-action" id="keywordBtn" type="button">검색 키워드</button>
        </div>
        <div class="date-row">
          <label>시작일<input id="newsStart" type="date"></label>
          <label>종료일<input id="newsEnd" type="date"></label>
        </div>
        <div class="date-row limit-row">
          <label>은행/지주사 최대 건수<input id="newsBankMax" type="number" min="1" max="100" value="10"></label>
          <label>그외 최대 건수<input id="newsOtherMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="button-row">
          <button class="jb" id="collectNewsBtn" type="button">자료 수집</button>
        </div>
      </div>
      <div class="group">
        <div class="group-head">
          <div class="group-title">국가기관</div>
          <button class="jb header-action" id="agencyListBtn" type="button">기관 목록</button>
        </div>
        <div class="date-row">
          <label>시작일<input id="agencyStart" type="date"></label>
          <label>종료일<input id="agencyEnd" type="date"></label>
        </div>
        <div class="limit-row">
          <label>기관별 최대 건수<input id="agencyMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="card-spacer"></div>
        <div class="limit-row collect-row">
          <div class="button-row">
            <button class="jb" id="collectAgencyBtn" type="button">자료 수집</button>
          </div>
        </div>
      </div>
      <div class="group">
        <div class="group-head">
          <div class="group-title">금융연구소</div>
          <button class="jb header-action" id="researchListBtn" type="button">연구소 목록</button>
        </div>
        <div class="date-row">
          <label>시작일<input id="researchStart" type="date"></label>
          <label>종료일<input id="researchEnd" type="date"></label>
        </div>
        <div class="limit-row">
          <label>연구소별 최대 건수<input id="researchMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="card-spacer"></div>
        <div class="limit-row collect-row">
          <div class="button-row">
            <button class="jb" id="collectResearchBtn" type="button">자료 수집</button>
          </div>
        </div>
      </div>
    </section>

    <div class="side">
      <button class="secondary" id="collectBtn">일괄 수집</button>
    </div>

    <div class="status" id="status">기간을 선택한 뒤 수집을 시작하세요.</div>

    <div class="toolbar">
      <div class="summary" id="summary">선택 0건 / 전체 0건</div>
      <div>
        <button class="ghost" id="historyBtn">실행 기록</button>
        <button class="secondary" id="reportBtn">보고서 생성</button>
        <button class="ghost" id="openReportBtn" disabled>보고서 열기</button>
      </div>
    </div>

    <div class="tabs" id="tabs"></div>

    <div class="table-wrap">
      <table>
        <thead>
          <tr>
            <th><input id="headCheck" type="checkbox"></th>
            <th>구분</th>
            <th>언론/기관/연구소</th>
            <th>일자</th>
            <th>제목</th>
            <th>형식</th>
            <th></th>
          </tr>
        </thead>
        <tbody id="tbody"></tbody>
      </table>
    </div>
  </main>

  <div class="modal-backdrop" id="historyModal" aria-hidden="true">
    <section class="modal" role="dialog" aria-modal="true" aria-labelledby="historyTitle">
      <div class="modal-head">
        <div class="modal-title" id="historyTitle">실행 기록</div>
        <button class="modal-close" id="historyCloseBtn" type="button" title="닫기">×</button>
      </div>
      <div class="modal-help">불러올 실행 기록을 선택하세요. 보고서가 생성된 기록은 별도로 표시됩니다.</div>
      <div class="history-list" id="historyList"></div>
    </section>
  </div>

  <div class="modal-backdrop" id="keywordModal" aria-hidden="true">
    <section class="modal keyword-modal" role="dialog" aria-modal="true" aria-labelledby="keywordTitle">
      <div class="modal-head">
        <div class="modal-title" id="keywordTitle">뉴스 검색 키워드</div>
        <button class="modal-close" id="keywordCloseBtn" type="button" title="닫기">×</button>
      </div>
      <div class="modal-help">검색할 문구를 추가하거나 삭제하세요. 저장된 키워드는 다음 수집과 앱 재실행 때 그대로 사용됩니다.</div>
      <div class="keyword-body">
        <div class="keyword-tabs">
          <button class="keyword-tab active" id="keywordBankTab" type="button">은행/지주사</button>
          <button class="keyword-tab" id="keywordOtherTab" type="button">그외</button>
        </div>
        <div class="keyword-editor">
          <input id="keywordInput" type="text" placeholder="예: 전북은행">
          <button id="keywordAddBtn" type="button">추가</button>
        </div>
        <div class="keyword-list" id="keywordList"></div>
        <div class="modal-actions">
          <button class="danger-outline" id="keywordDeleteBtn" type="button">삭제</button>
          <button class="jb" id="keywordSaveBtn" type="button">저장</button>
        </div>
      </div>
    </section>
  </div>

  <div class="modal-backdrop" id="sourceModal" aria-hidden="true">
    <section class="modal keyword-modal" role="dialog" aria-modal="true" aria-labelledby="sourceTitle">
      <div class="modal-head">
        <div class="modal-title" id="sourceTitle">수집 대상</div>
        <button class="modal-close" id="sourceCloseBtn" type="button" title="닫기">×</button>
      </div>
      <div class="modal-help" id="sourceHelp">현재 설정된 수집 대상 목록입니다.</div>
      <div class="keyword-body">
        <div class="source-list" id="sourceList"></div>
      </div>
    </section>
  </div>

  <script>
    let rows = [];
    let activeTab = "all";
    let currentReportPath = "";
    let keywordState = {bank: [], other: []};
    let activeKeywordKind = "bank";
    let selectedKeywordIndex = -1;
    const sourceLists = {
      agency: ["금융감독원", "금융위원회", "한국은행"],
      research: [
        "한국금융연구원",
        "하나금융연구소 연구보고서",
        "하나금융연구소 정기보고서",
        "KB경영연구소",
        "KDB미래전략연구소",
        "우리금융경영연구소 연구보고서",
        "우리금융경영연구소 동남아 Review"
      ]
    };

    function isoDate(d) {
      return d.toISOString().slice(0, 10);
    }
    function setDefaults() {
      const end = new Date();
      const start = new Date();
      start.setDate(end.getDate() - 7);
      for (const id of ["newsStart", "agencyStart", "researchStart"]) {
        document.getElementById(id).value = isoDate(start);
      }
      for (const id of ["newsEnd", "agencyEnd", "researchEnd"]) {
        document.getElementById(id).value = isoDate(end);
      }
    }
    function payload(targets = {news: true, agency: true, research: true}) {
      return {
        news_start: document.getElementById("newsStart").value,
        news_end: document.getElementById("newsEnd").value,
        agency_start: document.getElementById("agencyStart").value,
        agency_end: document.getElementById("agencyEnd").value,
        research_start: document.getElementById("researchStart").value,
        research_end: document.getElementById("researchEnd").value,
        news_bank_max: Number(document.getElementById("newsBankMax").value || 10),
        news_other_max: Number(document.getElementById("newsOtherMax").value || 10),
        agency_max: Number(document.getElementById("agencyMax").value || 10),
        research_max: Number(document.getElementById("researchMax").value || 10),
        include_news: Boolean(targets.news),
        include_agency: Boolean(targets.agency),
        include_research: Boolean(targets.research)
      };
    }
    function updateSummary() {
      const selected = rows.filter(r => r.checked).length;
      const visible = visibleRows();
      const visibleSelected = visible.filter(r => r.checked).length;
      document.getElementById("summary").textContent = `선택 ${selected}건 / 전체 ${rows.length}건 · 현재 탭 ${visibleSelected}/${visible.length}건`;
      document.getElementById("headCheck").checked = visible.length > 0 && visibleSelected === visible.length;
      updateReportControls();
    }
    function setReportPath(path) {
      currentReportPath = path || "";
      updateReportControls();
    }
    function updateReportControls() {
      const selected = rows.filter(r => r.checked).length;
      const reportBtn = document.getElementById("reportBtn");
      reportBtn.disabled = selected === 0;
      reportBtn.textContent = currentReportPath ? "재생성" : "보고서 생성";
      document.getElementById("openReportBtn").disabled = !currentReportPath;
    }
    function tabKey(row) {
      if (row.category === "그외") return "뉴스 그외";
      if (row.category === "은행/지주사") return "뉴스 은행/지주사";
      if (row.category === "국가기관") return "국가기관";
      if (row.category === "금융연구소") return "금융연구소";
      return row.category || "기타";
    }
    function visibleRows() {
      if (activeTab === "all") return rows;
      return rows.filter(row => tabKey(row) === activeTab);
    }
    function renderTabs() {
      const tabs = document.getElementById("tabs");
      const counts = new Map();
      for (const row of rows) {
        const key = tabKey(row);
        counts.set(key, (counts.get(key) || 0) + 1);
      }
      const tabOrder = ["뉴스 그외", "뉴스 은행/지주사", "국가기관", "금융연구소"];
      const ordered = ["all", ...tabOrder.filter(key => counts.has(key))];
      if (activeTab !== "all" && !ordered.includes(activeTab)) activeTab = "all";
      tabs.innerHTML = "";
      for (const key of ordered) {
        const btn = document.createElement("button");
        btn.className = "tab" + (key === activeTab ? " active" : "");
        btn.textContent = key === "all" ? `전체 (${rows.length})` : `${key} (${counts.get(key)})`;
        btn.addEventListener("click", () => {
          activeTab = key;
          render();
        });
        tabs.appendChild(btn);
      }
    }
    function render() {
      renderTabs();
      const tbody = document.getElementById("tbody");
      tbody.innerHTML = "";
      visibleRows().forEach((row) => {
        const index = rows.findIndex(r => r.id === row.id);
        const tr = document.createElement("tr");
        tr.dataset.id = row.id;
        tr.innerHTML = `
          <td><input type="checkbox" ${row.checked ? "checked" : ""}></td>
          <td>${escapeHtml(row.category || "")}</td>
          <td>${escapeHtml(row.source_name || "")}</td>
          <td>${escapeHtml(row.published_date || "")}</td>
          <td class="title">${escapeHtml(row.title || "")}</td>
          <td><span class="kind">${escapeHtml(row.file_type || "")}</span></td>
          <td class="action-cell"><button class="open-item-btn" type="button" title="${row.file_type === "article" ? "기사 링크 열기" : "PDF 열기"}">${openItemIcon(row)}</button></td>
        `;
        tr.querySelector("input").addEventListener("change", e => {
          rows[index].checked = e.target.checked;
          updateReportControls();
          updateSummary();
        });
        tr.querySelector(".open-item-btn").addEventListener("click", () => openItem(row));
        tbody.appendChild(tr);
      });
      updateSummary();
    }
    function openItemIcon(row) {
      if (row.file_type === "article") {
        return `<svg viewBox="0 0 24 24" aria-hidden="true"><path d="M7 17 17 7"></path><path d="M9 7h8v8"></path><path d="M5 5v14h14"></path></svg>`;
      }
      return `<svg viewBox="0 0 24 24" aria-hidden="true"><path d="M7 3h7l5 5v13H7z"></path><path d="M14 3v6h5"></path><path d="M9 14h6"></path><path d="M9 17h6"></path></svg>`;
    }
    function escapeHtml(value) {
      return String(value)
        .replaceAll("&", "&amp;")
        .replaceAll("<", "&lt;")
        .replaceAll(">", "&gt;")
        .replaceAll('"', "&quot;");
    }
    function setCollectButtonsDisabled(disabled) {
      for (const id of ["collectBtn", "collectNewsBtn", "collectAgencyBtn", "collectResearchBtn"]) {
        document.getElementById(id).disabled = disabled;
      }
    }
    async function collect(targets = {news: true, agency: true, research: true}) {
      setCollectButtonsDisabled(true);
      setReportPath("");
      updateReportControls();
      let streaming = false;
      document.getElementById("status").textContent = "수집 중입니다. 사이트 응답에 따라 1~3분 정도 걸릴 수 있습니다.";
      try {
        const res = await fetch("/collect-start", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify(payload(targets))
        });
        const data = await res.json();
        if (!res.ok || !data.ok) throw new Error(data.error || "failed to start collection");
        const events = new EventSource(`/collect-events?id=${encodeURIComponent(data.job_id)}`);
        events.onmessage = (event) => {
          const msg = JSON.parse(event.data);
          if (msg.type === "source_start") {
            document.getElementById("status").textContent = `Collecting: ${msg.source}`;
          } else if (msg.type === "source_done") {
            document.getElementById("status").textContent = `${msg.source}: ${msg.count} item(s) found`;
          } else if (msg.type === "source_error") {
            document.getElementById("status").textContent = `${msg.source}: ${msg.error}`;
          } else if (msg.type === "decode_start") {
            document.getElementById("status").textContent = `Decoding article URL (${msg.index}/${msg.total}): ${msg.title || ""}`;
          } else if (msg.type === "decode_done") {
            document.getElementById("status").textContent = `Decoded article URL (${msg.index}/${msg.total})`;
          } else if (msg.type === "download_start") {
            document.getElementById("status").textContent = `Saving PDF (${msg.index}/${msg.total}): ${msg.source}`;
          } else if (msg.type === "complete") {
            rows = msg.items.map(x => ({...x, checked: true}));
            setReportPath("");
            activeTab = "all";
            render();
            document.getElementById("status").textContent = `Collection complete: ${rows.length} item(s) · ${msg.run_dir}`;
            events.close();
            setCollectButtonsDisabled(false);
          } else if (msg.type === "fatal") {
            document.getElementById("status").textContent = "Error: " + msg.error;
            events.close();
            setCollectButtonsDisabled(false);
          }
        };
        events.onerror = () => {
          document.getElementById("status").textContent = "Progress connection closed.";
          events.close();
          setCollectButtonsDisabled(false);
        };
        streaming = true;
        return;
        if (!res.ok || !data.ok) throw new Error(data.error || "수집 실패");
        rows = data.items.map(x => ({...x, checked: true}));
        activeTab = "all";
        render();
        document.getElementById("status").textContent = `수집 완료: ${rows.length}건 · 저장 폴더: ${data.run_dir}`;
      } catch (err) {
        document.getElementById("status").textContent = "오류: " + err.message;
      } finally {
        if (!streaming) setCollectButtonsDisabled(false);
      }
    }
    async function openItem(row) {
      const res = await fetch("/open", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({id: row.id})
      });
      const data = await res.json();
      if (data.open_in_client && data.target) {
        window.open(data.target, "_blank", "noopener");
        return;
      }
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "자료 열기 오류: " + (data.error || "열 수 없습니다.");
      }
    }
    async function openReport() {
      if (!currentReportPath) return;
      const res = await fetch("/open-report", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "보고서 열기 오류: " + (data.error || "열 수 없습니다.");
      }
    }
    async function saveSelection() {
      const ids = rows.filter(r => r.checked).map(r => r.id);
      const res = await fetch("/save-selection", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({ids})
      });
      const data = await res.json();
      if (data.ok) {
        document.getElementById("status").textContent = `선택 항목 저장 완료: ${ids.length}건 · ${data.csv}`;
      } else {
        document.getElementById("status").textContent = "저장 오류: " + data.error;
      }
    }
    async function ensureApiKey() {
      const statusRes = await fetch("/api-key-status");
      const status = await statusRes.json();
      if (status.has_key) {
        const useExisting = confirm("저장된 OpenAI API key를 사용할까요? 취소를 누르면 새 key를 입력합니다.");
        if (useExisting) return true;
      }
      const key = prompt("OpenAI API key를 입력하세요. 로컬 파일에 암호화 저장됩니다.");
      if (!key) return false;
      const res = await fetch("/api-key", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({api_key: key})
      });
      const data = await res.json();
      if (!data.ok) {
        alert("API key 저장 실패: " + data.error);
        return false;
      }
      return true;
    }
    async function generateReport() {
      const ids = rows.filter(r => r.checked).map(r => r.id);
      if (!ids.length) {
        alert("보고서에 포함할 자료를 하나 이상 선택하세요.");
        return;
      }
      const selectedRows = rows.filter(r => r.checked);
      const missingSections = [];
      if (!selectedRows.some(r => r.category === "은행/지주사" || r.category === "그외")) {
        missingSections.push("뉴스");
      }
      if (!selectedRows.some(r => r.category === "국가기관")) {
        missingSections.push("국가기관");
      }
      if (!selectedRows.some(r => r.category === "금융연구소")) {
        missingSections.push("금융연구소");
      }
      if (missingSections.length) {
        const okToContinue = confirm(`${missingSections.join(", ")} 수집 자료가 없습니다.\n자료가 없는 섹션은 비우고 보고서를 생성하시겠습니까?`);
        if (!okToContinue) return;
      }
      const ok = await ensureApiKey();
      if (!ok) return;
      const btn = document.getElementById("reportBtn");
      btn.disabled = true;
      const isRegenerate = Boolean(currentReportPath);
      document.getElementById("status").textContent = isRegenerate
        ? "보고서 재생성 중입니다. 체크된 자료 기준으로 기존 결과물을 대체합니다."
        : "보고서 생성 중입니다. 선택하지 않은 PDF를 정리하고 OpenAI API를 호출합니다.";
      try {
        const res = await fetch("/generate-report-start", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({ids})
        });
        const data = await res.json();
        if (!res.ok || !data.ok) throw new Error(data.error || "보고서 생성 시작 실패");
        const events = new EventSource(`/collect-events?id=${encodeURIComponent(data.job_id)}`);
        events.onmessage = (event) => {
          const msg = JSON.parse(event.data);
          if (msg.type === "generate_status") {
            document.getElementById("status").textContent = msg.message || "보고서 생성 중입니다.";
          } else if (msg.type === "generate_upload") {
            document.getElementById("status").textContent = `자료 업로드 중 (${msg.index}/${msg.total}): ${msg.title || ""}`;
          } else if (msg.type === "complete") {
            rows = msg.items.map(x => ({...x, checked: true}));
            setReportPath(msg.docx || "");
            activeTab = "all";
            render();
            document.getElementById("status").textContent = `${isRegenerate ? "보고서 재생성" : "보고서 생성"} 완료. 보고서 열기 버튼을 눌러 확인하세요.`;
            events.close();
            updateReportControls();
          } else if (msg.type === "fatal") {
            document.getElementById("status").textContent = "보고서 생성 오류: " + msg.error;
            events.close();
            updateReportControls();
          }
        };
        events.onerror = () => {
          document.getElementById("status").textContent = "보고서 생성 상태 연결이 종료되었습니다.";
          events.close();
          updateReportControls();
        };
      } catch (err) {
        document.getElementById("status").textContent = "보고서 생성 오류: " + err.message;
        updateReportControls();
      }
    }
    async function loadHistory() {
      const res = await fetch("/runs");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "실행 기록 오류: " + data.error;
        return;
      }
      const label = data.runs.map((r, i) => `${i + 1}. ${r.name} (${r.count}건)`).join("\n");
      const picked = prompt("불러올 실행 기록 번호를 입력하세요.\n\n" + label);
      const index = Number(picked) - 1;
      if (!Number.isInteger(index) || index < 0 || index >= data.runs.length) return;
      const run = data.runs[index];
      const detail = await fetch(`/run?id=${encodeURIComponent(run.name)}`);
      const loaded = await detail.json();
      if (!loaded.ok) {
        document.getElementById("status").textContent = "실행 기록 불러오기 오류: " + loaded.error;
        return;
      }
      rows = loaded.items.map(x => ({...x, checked: true}));
      setReportPath(loaded.docx || "");
      activeTab = "all";
      render();
      document.getElementById("status").textContent = `실행 기록 불러오기 완료: ${run.name}`;
    }
    async function loadHistoryModal() {
      const res = await fetch("/runs");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "실행 기록 오류: " + data.error;
        return;
      }
      showHistoryModal(data.runs);
    }
    function showHistoryModal(runs) {
      const modal = document.getElementById("historyModal");
      const list = document.getElementById("historyList");
      list.innerHTML = "";
      if (!runs.length) {
        list.innerHTML = `<div class="empty-state">저장된 실행 기록이 없습니다.</div>`;
      }
      runs.forEach((run) => {
        const btn = document.createElement("button");
        btn.type = "button";
        btn.className = "history-item";
        btn.innerHTML = `
          <div>
            <div class="history-name">${escapeHtml(formatRunName(run.name))}</div>
            <div class="history-meta">${escapeHtml(run.path || "")}</div>
          </div>
          <div class="history-actions">
            <div class="history-tags">
              <span class="history-tag">${Number(run.count || 0)}건</span>
              ${run.has_report ? `<span class="history-tag report">보고서</span>` : ""}
            </div>
            <button class="history-delete" type="button" title="실행 기록 삭제">삭제</button>
          </div>
        `;
        btn.addEventListener("click", () => openRunFromModal(run.name));
        btn.querySelector(".history-delete").addEventListener("click", (event) => {
          event.stopPropagation();
          deleteRunFromModal(run.name);
        });
        list.appendChild(btn);
      });
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
    }
    function closeHistoryModal() {
      const modal = document.getElementById("historyModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }
    function formatRunName(name) {
      const m = String(name).match(/^(\d{2})(\d{2})(\d{2})_(\d{2})(\d{2})(?:_(\d+))?$/);
      if (!m) return name;
      const suffix = m[6] ? ` #${m[6]}` : "";
      return `20${m[1]}.${m[2]}.${m[3]} ${m[4]}:${m[5]}${suffix}`;
    }
    async function openRunFromModal(runName) {
      closeHistoryModal();
      const detail = await fetch(`/run?id=${encodeURIComponent(runName)}`);
      const loaded = await detail.json();
      if (!loaded.ok) {
        document.getElementById("status").textContent = "실행 기록 불러오기 오류: " + loaded.error;
        return;
      }
      rows = loaded.items.map(x => ({...x, checked: true}));
      setReportPath(loaded.docx || "");
      activeTab = "all";
      render();
      document.getElementById("status").textContent = `실행 기록 불러오기 완료: ${runName}`;
    }
    async function deleteRunFromModal(runName) {
      if (!confirm(`${formatRunName(runName)} 실행 기록을 삭제할까요?\n\n해당 폴더의 PDF, JSON, 생성 보고서가 모두 삭제됩니다.`)) {
        return;
      }
      const res = await fetch("/delete-run", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({id: runName})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "실행 기록 삭제 오류: " + (data.error || "삭제할 수 없습니다.");
        return;
      }
      if (data.cleared_current) {
        rows = [];
        setReportPath("");
        activeTab = "all";
        render();
      }
      document.getElementById("status").textContent = `실행 기록 삭제 완료: ${runName}`;
      await loadHistoryModal();
    }

    async function openKeywordModal() {
      const res = await fetch("/news-keywords");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "키워드 불러오기 오류: " + data.error;
        return;
      }
      keywordState = {
        bank: [...(data.keywords.bank || [])],
        other: [...(data.keywords.other || [])]
      };
      activeKeywordKind = "bank";
      selectedKeywordIndex = -1;
      renderKeywordModal();
      const modal = document.getElementById("keywordModal");
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
      document.getElementById("keywordInput").focus();
    }

    function closeKeywordModal() {
      const modal = document.getElementById("keywordModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }

    function setKeywordKind(kind) {
      activeKeywordKind = kind;
      selectedKeywordIndex = -1;
      renderKeywordModal();
    }

    function renderKeywordModal() {
      document.getElementById("keywordBankTab").className = "keyword-tab" + (activeKeywordKind === "bank" ? " active" : "");
      document.getElementById("keywordOtherTab").className = "keyword-tab" + (activeKeywordKind === "other" ? " active" : "");
      const list = document.getElementById("keywordList");
      const items = keywordState[activeKeywordKind] || [];
      list.innerHTML = "";
      if (!items.length) {
        list.innerHTML = `<div class="empty-state">저장된 키워드가 없습니다.</div>`;
        return;
      }
      items.forEach((keyword, index) => {
        const btn = document.createElement("button");
        btn.type = "button";
        btn.className = "keyword-item" + (index === selectedKeywordIndex ? " selected" : "");
        btn.textContent = keyword;
        btn.addEventListener("click", () => {
          selectedKeywordIndex = index;
          renderKeywordModal();
        });
        list.appendChild(btn);
      });
    }

    function addKeyword() {
      const input = document.getElementById("keywordInput");
      const value = input.value.trim();
      if (!value) return;
      const items = keywordState[activeKeywordKind] || [];
      if (!items.includes(value)) {
        items.push(value);
        keywordState[activeKeywordKind] = items;
        selectedKeywordIndex = items.length - 1;
      }
      input.value = "";
      renderKeywordModal();
      input.focus();
    }

    function deleteKeyword() {
      const items = keywordState[activeKeywordKind] || [];
      if (selectedKeywordIndex < 0 || selectedKeywordIndex >= items.length) return;
      items.splice(selectedKeywordIndex, 1);
      selectedKeywordIndex = -1;
      renderKeywordModal();
    }

    async function saveKeywords() {
      const res = await fetch("/news-keywords", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({keywords: keywordState})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "키워드 저장 오류: " + (data.error || "저장할 수 없습니다.");
        return;
      }
      keywordState = data.keywords;
      closeKeywordModal();
      document.getElementById("status").textContent = "뉴스 검색 키워드 저장 완료";
    }

    function openSourceModal(kind) {
      const title = kind === "agency" ? "기관 목록" : "연구소 목록";
      const help = kind === "agency"
        ? "국가기관 자료 수집 시 아래 기관에서 자료를 가져옵니다."
        : "금융연구소 자료 수집 시 아래 연구소에서 자료를 가져옵니다.";
      document.getElementById("sourceTitle").textContent = title;
      document.getElementById("sourceHelp").textContent = help;
      const list = document.getElementById("sourceList");
      list.innerHTML = "";
      for (const source of sourceLists[kind] || []) {
        const item = document.createElement("div");
        item.className = "source-item";
        item.textContent = source;
        list.appendChild(item);
      }
      const modal = document.getElementById("sourceModal");
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
    }

    function closeSourceModal() {
      const modal = document.getElementById("sourceModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }

    document.getElementById("collectBtn").addEventListener("click", () => collect({news: true, agency: true, research: true}));
    document.getElementById("collectNewsBtn").addEventListener("click", () => collect({news: true, agency: false, research: false}));
    document.getElementById("collectAgencyBtn").addEventListener("click", () => collect({news: false, agency: true, research: false}));
    document.getElementById("collectResearchBtn").addEventListener("click", () => collect({news: false, agency: false, research: true}));
    document.getElementById("reportBtn").addEventListener("click", generateReport);
    document.getElementById("openReportBtn").addEventListener("click", openReport);
    document.getElementById("historyBtn").addEventListener("click", loadHistoryModal);
    document.getElementById("keywordBtn").addEventListener("click", openKeywordModal);
    document.getElementById("agencyListBtn").addEventListener("click", () => openSourceModal("agency"));
    document.getElementById("researchListBtn").addEventListener("click", () => openSourceModal("research"));
    document.getElementById("historyCloseBtn").addEventListener("click", closeHistoryModal);
    document.getElementById("keywordCloseBtn").addEventListener("click", closeKeywordModal);
    document.getElementById("sourceCloseBtn").addEventListener("click", closeSourceModal);
    document.getElementById("keywordBankTab").addEventListener("click", () => setKeywordKind("bank"));
    document.getElementById("keywordOtherTab").addEventListener("click", () => setKeywordKind("other"));
    document.getElementById("keywordAddBtn").addEventListener("click", addKeyword);
    document.getElementById("keywordDeleteBtn").addEventListener("click", deleteKeyword);
    document.getElementById("keywordSaveBtn").addEventListener("click", saveKeywords);
    document.getElementById("keywordInput").addEventListener("keydown", e => {
      if (e.key === "Enter") addKeyword();
    });
    document.getElementById("historyModal").addEventListener("click", e => {
      if (e.target.id === "historyModal") closeHistoryModal();
    });
    document.getElementById("keywordModal").addEventListener("click", e => {
      if (e.target.id === "keywordModal") closeKeywordModal();
    });
    document.getElementById("sourceModal").addEventListener("click", e => {
      if (e.target.id === "sourceModal") closeSourceModal();
    });
    document.addEventListener("keydown", e => {
      if (e.key === "Escape") {
        closeHistoryModal();
        closeKeywordModal();
        closeSourceModal();
      }
    });
    document.getElementById("headCheck").addEventListener("change", e => {
      visibleRows().forEach(r => r.checked = e.target.checked);
      render();
    });
    setDefaults();
    updateReportControls();
  </script>
</body>
</html>
"""


class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt: str, *args) -> None:
        print(f"[app] {self.address_string()} {fmt % args}")

    def do_GET(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/":
            self.send_text(INDEX_HTML, "text/html; charset=utf-8")
            return
        if parsed.path == "/collect-events":
            self.handle_collect_events(parsed)
            return
        if parsed.path == "/api-key-status":
            self.send_json({"ok": True, "has_key": KEY_PATH.exists()})
            return
        if parsed.path == "/news-keywords":
            self.send_json({"ok": True, "keywords": load_news_keywords()})
            return
        if parsed.path == "/runs":
            self.handle_runs()
            return
        if parsed.path == "/run":
            self.handle_run(parsed)
            return
        self.send_json({"ok": False, "error": "not found"}, status=404)

    def do_POST(self) -> None:
        try:
            if self.path == "/collect-start":
                self.handle_collect_start()
            elif self.path == "/collect":
                self.handle_collect()
            elif self.path == "/open":
                self.handle_open()
            elif self.path == "/open-report":
                self.handle_open_report()
            elif self.path == "/save-selection":
                self.handle_save_selection()
            elif self.path == "/api-key":
                self.handle_api_key()
            elif self.path == "/news-keywords":
                self.handle_news_keywords()
            elif self.path == "/generate-report-start":
                self.handle_generate_report_start()
            elif self.path == "/generate-report":
                self.handle_generate_report()
            elif self.path == "/delete-run":
                self.handle_delete_run()
            else:
                self.send_json({"ok": False, "error": "not found"}, status=404)
        except Exception as exc:
            traceback.print_exc()
            self.send_json({"ok": False, "error": str(exc)}, status=500)

    def handle_collect(self) -> None:
        data = self.read_json()
        news_start = parse_iso_date(data["news_start"])
        news_end = parse_iso_date(data["news_end"])
        agency_start = parse_iso_date(data["agency_start"])
        agency_end = parse_iso_date(data["agency_end"])
        research_start = parse_iso_date(data["research_start"])
        research_end = parse_iso_date(data["research_end"])
        news_bank_max = int(data.get("news_bank_max") or 10)
        news_other_max = int(data.get("news_other_max") or 10)
        agency_max = int(data.get("agency_max") or 10)
        research_max = int(data.get("research_max") or 10)
        news_keywords = load_news_keywords()
        include_news = bool(data.get("include_news", True))
        include_agency = bool(data.get("include_agency", True))
        include_research = bool(data.get("include_research", True))
        run_dir = collector.make_run_dir()

        items = collector.collect_by_ranges(
            news_start,
            news_end,
            agency_start,
            agency_end,
            research_start,
            research_end,
            max_per_source=None,
            dry_run=False,
            news_bank_max=news_bank_max,
            news_other_max=news_other_max,
            agency_max=agency_max,
            research_max=research_max,
            output_dir=run_dir,
            news_keywords=news_keywords,
            include_news=include_news,
            include_agency=include_agency,
            include_research=include_research,
        )
        normalize_article_urls(items)
        collector.write_outputs(items, output_dir=run_dir)
        payload = []
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = ""
            for idx, item in enumerate(items):
                row = row_for_client(item)
                row["id"] = idx
                payload.append(row)
        self.send_json({"ok": True, "items": payload, "run_dir": str(run_dir)})

    def handle_collect_start(self) -> None:
        data = self.read_json()
        job_id = uuid.uuid4().hex
        job_queue: queue.Queue = queue.Queue()
        job = {"queue": job_queue, "done": False}
        with STATE_LOCK:
            JOBS[job_id] = job
            STATE["report_path"] = ""
        thread = threading.Thread(target=run_collect_job, args=(job_id, data), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_collect_events(self, parsed) -> None:
        params = urllib.parse.parse_qs(parsed.query)
        job_id = params.get("id", [""])[0]
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if not job:
            self.send_json({"ok": False, "error": "invalid job id"}, status=404)
            return

        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Connection", "keep-alive")
        self.end_headers()

        job_queue = job["queue"]
        while True:
            try:
                event = job_queue.get(timeout=15)
            except queue.Empty:
                event = {"type": "ping"}
            data = f"data: {json.dumps(event, ensure_ascii=False)}\n\n".encode("utf-8")
            try:
                self.wfile.write(data)
                self.wfile.flush()
            except (BrokenPipeError, ConnectionResetError):
                break
            if event.get("type") in {"complete", "fatal"}:
                with STATE_LOCK:
                    JOBS.pop(job_id, None)
                break

    def handle_open(self) -> None:
        try:
            data = self.read_json()
            item_id = int(data["id"])
            with STATE_LOCK:
                items = STATE.get("items", [])
                run_dir_value = str(STATE.get("run_dir") or "")
                if item_id < 0 or item_id >= len(items):
                    self.send_json({"ok": False, "error": "invalid id"}, status=400)
                    return
                item = items[item_id]
            if item.file_type == "article":
                target_url = resolve_article_open_url(item)
                if target_url:
                    item.original_url = target_url
                    self.send_json({"ok": True, "target": target_url, "open_in_client": True})
                    return
                self.send_json({"ok": False, "error": "기사 URL이 없습니다."}, status=404)
                return
            run_dir = Path(run_dir_value) if run_dir_value else None
            path = resolve_item_local_path(item, run_dir)
            if path:
                os.startfile(str(path))
                self.send_json({"ok": True, "target": str(path)})
                return
            self.send_json({"ok": False, "error": "실행 기록 폴더에서 PDF 파일을 찾을 수 없습니다."}, status=404)
        except Exception as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=500)

    def handle_open_report(self) -> None:
        with STATE_LOCK:
            report_path = str(STATE.get("report_path") or "")
            run_dir_value = str(STATE.get("run_dir") or "")
        candidates = []
        if report_path:
            candidates.append(Path(report_path))
        if run_dir_value:
            candidates.append(Path(run_dir_value) / "generated_report.docx")
        for path in candidates:
            if path.exists() and path.is_file():
                os.startfile(path)
                self.send_json({"ok": True, "docx": str(path)})
                return
        self.send_json({"ok": False, "error": "생성된 보고서 파일을 찾을 수 없습니다."}, status=404)

    def handle_save_selection(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        with STATE_LOCK:
            selected = [item for i, item in enumerate(STATE.get("items", [])) if i in ids]
            run_dir_value = STATE.get("run_dir")
        run_dir = Path(run_dir_value) if run_dir_value else collector.OUT_DIR
        selected_csv = run_dir / "selected_metadata.csv"
        selected_xlsx = run_dir / "selected_metadata.xlsx"
        write_items(selected, selected_csv, selected_xlsx)
        self.send_json(
            {
                "ok": True,
                "count": len(selected),
                "csv": str(selected_csv),
                "xlsx": str(selected_xlsx),
            }
        )

    def handle_api_key(self) -> None:
        data = self.read_json()
        api_key = str(data.get("api_key") or "").strip()
        if not api_key:
            self.send_json({"ok": False, "error": "api_key is required"}, status=400)
            return
        save_api_key(api_key)
        self.send_json({"ok": True})

    def handle_news_keywords(self) -> None:
        data = self.read_json()
        try:
            keywords = normalize_news_keywords(data.get("keywords", data))
        except ValueError as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        save_news_keywords(keywords)
        self.send_json({"ok": True, "keywords": keywords})

    def handle_generate_report(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        result = perform_generate_report(ids)
        self.send_json(result)

    def handle_generate_report_start(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        job_id = uuid.uuid4().hex
        job_queue: queue.Queue = queue.Queue()
        job = {"queue": job_queue, "done": False}
        with STATE_LOCK:
            JOBS[job_id] = job
        thread = threading.Thread(target=run_generate_report_job, args=(job_id, ids), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_runs(self) -> None:
        runs = []
        if collector.RUNS_DIR.exists():
            for path in sorted((p for p in collector.RUNS_DIR.iterdir() if p.is_dir()), reverse=True):
                metadata = path / "metadata.csv"
                count = 0
                if metadata.exists():
                    with metadata.open("r", encoding="utf-8-sig", newline="") as f:
                        count = sum(1 for _ in csv.DictReader(f))
                runs.append(
                    {
                        "name": path.name,
                        "path": str(path),
                        "count": count,
                        "has_report": (path / "generated_report.docx").exists(),
                    }
                )
        self.send_json({"ok": True, "runs": runs})

    def handle_run(self, parsed) -> None:
        run_id = urllib.parse.parse_qs(parsed.query).get("id", [""])[0]
        run_dir = resolve_run_dir(run_id)
        if run_dir is None:
            self.send_json({"ok": False, "error": "invalid run id"}, status=400)
            return
        metadata = run_dir / "metadata.csv"
        if not metadata.exists():
            self.send_json({"ok": False, "error": "metadata.csv not found"}, status=404)
            return
        items: list[collector.Item] = []
        with metadata.open("r", encoding="utf-8-sig", newline="") as f:
            for idx, row in enumerate(csv.DictReader(f)):
                item = item_from_row(row)
                items.append(item)
        payload = []
        for idx, item in enumerate(items):
            out = row_for_client(item)
            out["id"] = idx
            payload.append(out)
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = str(run_dir / "generated_report.docx") if (run_dir / "generated_report.docx").exists() else ""
        self.send_json(
            {
                "ok": True,
                "items": payload,
                "run_dir": str(run_dir),
                "docx": str(run_dir / "generated_report.docx") if (run_dir / "generated_report.docx").exists() else "",
            }
        )

    def handle_delete_run(self) -> None:
        data = self.read_json()
        run_id = str(data.get("id") or "")
        run_dir = resolve_run_dir(run_id)
        if run_dir is None:
            self.send_json({"ok": False, "error": "invalid run id"}, status=400)
            return
        if not run_dir.exists() or not run_dir.is_dir():
            self.send_json({"ok": False, "error": "run not found"}, status=404)
            return
        cleared_current = False
        with STATE_LOCK:
            current = str(STATE.get("run_dir") or "")
            if current and same_path(current, run_dir):
                STATE["items"] = []
                STATE["run_dir"] = None
                STATE["report_path"] = ""
                cleared_current = True
        shutil.rmtree(run_dir)
        self.send_json({"ok": True, "deleted": str(run_dir), "cleared_current": cleared_current})

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0"))
        raw = self.rfile.read(length).decode("utf-8")
        return json.loads(raw or "{}")

    def send_text(self, body: str, content_type: str) -> None:
        data = body.encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_json(self, data: dict, status: int = 200) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def perform_generate_report(ids: set[int], progress=None) -> dict:
    def emit(message: str) -> None:
        if progress:
            progress("generate_status", {"message": message})

    emit("선택 항목을 확인하는 중입니다.")
    with STATE_LOCK:
        current_items = list(STATE.get("items", []))
        source_run_dir = str(STATE.get("run_dir") or "")
    if not current_items:
        raise RuntimeError("수집된 목록이 없습니다.")
    selected = [item for i, item in enumerate(current_items) if i in ids]
    unselected = [item for i, item in enumerate(current_items) if i not in ids]
    existing_run_dir = resolve_existing_run_dir(source_run_dir)
    if not selected:
        raise RuntimeError("선택된 항목이 없습니다.")

    emit("기사 원문 링크를 확인하는 중입니다.")
    normalize_article_urls(selected)

    emit("OpenAI API key를 확인하는 중입니다.")
    api_key = load_api_key()
    if not api_key:
        raise RuntimeError("OpenAI API key가 저장되어 있지 않습니다.")

    emit("선택하지 않은 자료를 정리하는 중입니다.")
    if existing_run_dir:
        for item in unselected:
            delete_local_file_under(item.local_path, existing_run_dir)
    else:
        for item in unselected:
            delete_local_file(item.local_path)

    run_dir = existing_run_dir or collector.make_run_dir()
    emit("실행 기록 폴더를 준비하는 중입니다.")
    if not existing_run_dir:
        move_selected_files(selected, run_dir, move=False)
    enrich_article_texts(selected, progress=progress)
    collector.write_outputs(selected, output_dir=run_dir)

    emit("프롬프트와 입력 자료를 준비하는 중입니다.")
    prompt_path = ensure_prompt_template()
    prompt = prompt_path.read_text(encoding="utf-8")
    report_json = call_openai_report(api_key, prompt, selected, progress=progress)
    llm_output_text = str(report_json.pop("_openai_output_text", "") or "")
    llm_output_path = run_dir / "llm_output.txt"
    if llm_output_text:
        emit("LLM output text를 저장하는 중입니다.")
        llm_output_path.write_text(llm_output_text, encoding="utf-8")
    report_json_path = run_dir / "report_data.json"
    emit("LLM 응답 JSON을 저장하는 중입니다.")
    report_json_path.write_text(
        json.dumps(report_json, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    docx_path = run_dir / "generated_report.docx"
    emit("워드 보고서를 생성하는 중입니다.")
    generate_docx(report_json, selected, docx_path)

    payload = []
    with STATE_LOCK:
        STATE["items"] = selected
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = str(docx_path)
        for idx, item in enumerate(selected):
            row = row_for_client(item)
            row["id"] = idx
            payload.append(row)
    return {
        "ok": True,
        "items": payload,
        "run_dir": str(run_dir),
        "json": str(report_json_path),
        "llm_output": str(llm_output_path) if llm_output_text else "",
        "docx": str(docx_path),
    }


def enrich_article_texts(items: list[collector.Item], progress=None) -> None:
    article_items = [
        item for item in items
        if item.file_type == "article" and not item.extra.get("article_text") and (item.original_url or item.url)
    ]
    if not article_items:
        return
    normalize_article_urls(article_items)
    http = collector.Http()
    total = len(article_items)
    for index, item in enumerate(article_items, start=1):
        if progress:
            progress("generate_status", {"message": f"기사 본문 추출 중 ({index}/{total}): {item.title}"})
        text, note = collector.fetch_article_text(
            http,
            item.original_url or item.url,
            referer=item.extra.get("google_news_url", "") if isinstance(item.extra, dict) else "",
        )
        if text:
            item.extra["article_text"] = text
        if note:
            item.notes = "; ".join(x for x in [item.notes, note] if x)


def run_generate_report_job(job_id: str, ids: set[int]) -> None:
    def emit(event_type: str, payload: dict | None = None) -> None:
        payload = dict(payload or {})
        payload["type"] = event_type
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if job:
            job["queue"].put(payload)

    try:
        result = perform_generate_report(ids, progress=emit)
        emit("complete", result)
    except Exception as exc:
        traceback.print_exc()
        emit("fatal", {"error": str(exc)})


def parse_iso_date(value: str) -> dt.date:
    return dt.date.fromisoformat(value)


def ensure_prompt_template() -> Path:
    if not PROMPT_PATH.exists():
        PROMPT_PATH.write_text(DEFAULT_REPORT_PROMPT, encoding="utf-8")
    return PROMPT_PATH


def default_news_keywords() -> dict:
    return {
        "bank": list(collector.DEFAULT_NEWS_KEYWORDS["bank"]),
        "other": list(collector.DEFAULT_NEWS_KEYWORDS["other"]),
    }


def normalize_news_keywords(value) -> dict:
    if not isinstance(value, dict):
        raise ValueError("keywords must be an object")
    defaults = default_news_keywords()
    out = {}
    for key in ["bank", "other"]:
        raw_items = value.get(key, defaults[key])
        if not isinstance(raw_items, list):
            raise ValueError(f"{key} keywords must be a list")
        items = []
        seen = set()
        for raw in raw_items:
            text = str(raw or "").strip().strip('"').strip("'").strip()
            if not text or text in seen:
                continue
            seen.add(text)
            items.append(text)
        out[key] = items or defaults[key]
    return out


def load_news_keywords() -> dict:
    if NEWS_KEYWORDS_PATH.exists():
        try:
            return normalize_news_keywords(json.loads(NEWS_KEYWORDS_PATH.read_text(encoding="utf-8")))
        except Exception:
            pass
    keywords = default_news_keywords()
    save_news_keywords(keywords)
    return keywords


def save_news_keywords(keywords: dict) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    NEWS_KEYWORDS_PATH.write_text(
        json.dumps(normalize_news_keywords(keywords), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def delete_local_file(path_value: str) -> None:
    if not path_value:
        return
    path = Path(path_value)
    try:
        if path.exists() and path.is_file():
            path.unlink()
    except Exception:
        pass


def resolve_existing_run_dir(path_value: str) -> Path | None:
    if not path_value:
        return None
    try:
        path = Path(path_value).resolve()
        runs_root = collector.RUNS_DIR.resolve()
        if path.exists() and path.is_dir() and str(path).startswith(str(runs_root)):
            return path
    except Exception:
        return None
    return None


def resolve_run_dir(run_id: str) -> Path | None:
    if not re.fullmatch(r"\d{6}_\d{4}(?:_\d+)?", run_id or ""):
        return None
    try:
        runs_root = collector.RUNS_DIR.resolve()
        path = (runs_root / run_id).resolve()
        path.relative_to(runs_root)
        return path
    except Exception:
        return None


def same_path(left: str | Path, right: str | Path) -> bool:
    try:
        return Path(left).resolve() == Path(right).resolve()
    except Exception:
        return False


def delete_local_file_under(path_value: str, root: Path) -> None:
    if not path_value:
        return
    try:
        path = Path(path_value).resolve()
        root = root.resolve()
        if path.exists() and path.is_file() and str(path).startswith(str(root)):
            path.unlink()
    except Exception:
        pass


def resolve_item_local_path(item: collector.Item, run_dir: Path | None = None) -> Path | None:
    candidates: list[Path] = []
    if item.local_path:
        raw = Path(item.local_path)
        candidates.append(raw)
        if run_dir and not raw.is_absolute():
            candidates.append(run_dir / raw)
        if run_dir and raw.name:
            candidates.append(run_dir / raw.name)
            candidates.append(run_dir / "raw_pdfs" / raw.name)
    for candidate in candidates:
        try:
            path = candidate.resolve()
            if path.exists() and path.is_file():
                return path
        except Exception:
            continue
    if run_dir and item.local_path:
        name = Path(item.local_path).name
        if name:
            try:
                run_root = run_dir.resolve()
                for path in run_root.rglob(name):
                    if path.is_file():
                        path.resolve().relative_to(run_root)
                        return path
            except Exception:
                return None
    return None


def resolve_article_open_url(item: collector.Item) -> str:
    if item.original_url:
        return item.original_url
    google_url = ""
    if isinstance(item.extra, dict):
        google_url = str(item.extra.get("google_news_url") or "").strip()
    if google_url:
        decoded = collector.original_url_from_google_news(collector.Http(), google_url)
        if decoded:
            return decoded
    return item.url or ""


def normalize_article_urls(items: list[collector.Item]) -> None:
    collector.populate_original_urls(items)


def move_selected_files(items: list[collector.Item], run_dir: Path, move: bool = True) -> None:
    raw_root = run_dir / "raw_pdfs"
    for item in items:
        if not item.local_path:
            continue
        src = Path(item.local_path)
        if not src.exists() or not src.is_file():
            continue
        source_dir = raw_root / collector.clean_filename(item.source_name, item.source_name).replace(".pdf", "")
        source_dir.mkdir(parents=True, exist_ok=True)
        dst = source_dir / src.name
        if dst.exists():
            dst = source_dir / f"{dst.stem}_{uuid.uuid4().hex[:6]}{dst.suffix}"
        if move:
            shutil.move(str(src), str(dst))
        else:
            shutil.copy2(str(src), str(dst))
        item.local_path = str(dst)


def item_from_row(row: dict) -> collector.Item:
    extra = {}
    try:
        extra = json.loads(row.get("extra_json") or "{}")
    except json.JSONDecodeError:
        extra = {}
    item = collector.Item(
        category=row.get("category", ""),
        source_name=row.get("source_name", ""),
        title=row.get("title", ""),
        url=row.get("url", ""),
        published_date=row.get("published_date", ""),
        file_type=row.get("file_type", "pdf"),
        download_url=row.get("download_url", ""),
        local_path=row.get("local_path", ""),
        priority=row.get("priority", ""),
        priority_reason=row.get("priority_reason", ""),
        matched_keywords=row.get("matched_keywords", ""),
        notes=row.get("notes", ""),
        original_url=row.get("original_url", ""),
        extra=extra,
    )
    if item.file_type == "article" and (item.source_name or "").lower() in {"v.daum.net", "daum", "daum 뉴스"}:
        try:
            collector.normalize_article_source_name(collector.Http(), item)
        except Exception:
            pass
    return item


def protect_with_dpapi(secret: str) -> bytes:
    data = secret.encode("utf-8")
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptProtectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData)
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def unprotect_with_dpapi(data: bytes) -> str:
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptUnprotectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData).decode("utf-8")
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def save_api_key(api_key: str) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    try:
        KEY_PATH.write_bytes(b"dpapi:" + protect_with_dpapi(api_key))
    except Exception:
        KEY_PATH.write_bytes(b"b64:" + base64.b64encode(api_key.encode("utf-8")))


def load_api_key() -> str:
    if not KEY_PATH.exists():
        return ""
    data = KEY_PATH.read_bytes()
    if data.startswith(b"dpapi:"):
        return unprotect_with_dpapi(data[6:])
    if data.startswith(b"b64:"):
        return base64.b64decode(data[4:]).decode("utf-8")
    return ""


class DATA_BLOB(ctypes.Structure):
    _fields_ = [("cbData", ctypes.wintypes.DWORD), ("pbData", ctypes.POINTER(ctypes.c_char))]


def multipart_request(url: str, fields: dict, files: list[tuple[str, Path]], api_key: str) -> dict:
    boundary = "----codexboundary" + uuid.uuid4().hex
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"),
                str(value).encode("utf-8"),
                b"\r\n",
            ]
        )
    for name, path in files:
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"; filename="{path.name}"\r\n'.encode("utf-8"),
                b"Content-Type: application/pdf\r\n\r\n",
                path.read_bytes(),
                b"\r\n",
            ]
        )
    chunks.append(f"--{boundary}--\r\n".encode("utf-8"))
    body = b"".join(chunks)
    req = urllib.request.Request(
        url,
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
    )
    opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
    with opener.open(req, timeout=180) as resp:
        return json.loads(resp.read().decode("utf-8"))


def openai_json_request(path: str, payload: dict, api_key: str, timeout: int = 300) -> dict:
    req = urllib.request.Request(
        f"https://api.openai.com/v1{path}",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
    )
    try:
        opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
        with opener.open(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", "replace")
        raise RuntimeError(f"OpenAI API error {exc.code}: {detail}") from exc


def call_openai_report(api_key: str, prompt: str, items: list[collector.Item], progress=None) -> dict:
    def emit(event_type: str, payload: dict | None = None) -> None:
        if progress:
            progress(event_type, payload or {})

    emit("generate_status", {"message": "LLM 입력 JSON 골격을 만드는 중입니다."})
    report_base = build_report_base(items)
    report_items = list(iter_report_items(report_base))
    merged_bullets = {"ITEMS": []}
    response_logs = []
    response_ids = []
    batch_size = 6
    total_batches = max((len(report_items) + batch_size - 1) // batch_size, 1)

    for batch_index, start in enumerate(range(0, len(report_items), batch_size), start=1):
        batch = report_items[start:start + batch_size]
        expected_ids = [item.get("ITEM_ID", "") for item in batch]
        pending = batch
        collected: dict[str, dict] = {}
        for attempt in range(1, 3):
            if not pending:
                break
            emit(
                "generate_status",
                {
                    "message": (
                        f"OpenAI에 bullet 생성을 요청하는 중입니다. "
                        f"({batch_index}/{total_batches}, {len(pending)}건)"
                    )
                },
            )
            response, text, parsed = request_openai_bullets(api_key, prompt, pending, emit)
            response_ids.append(response.get("id", ""))
            response_logs.append({
                "batch": batch_index,
                "attempt": attempt,
                "expected_item_ids": [item.get("ITEM_ID", "") for item in pending],
                "response_id": response.get("id", ""),
                "output_text": text,
            })
            for raw in extract_bullet_items(parsed):
                if not isinstance(raw, dict):
                    continue
                item_id = clean_cell_value(raw.get("ITEM_ID"))
                if item_id in expected_ids:
                    collected[item_id] = raw
            missing = [item_id for item_id in expected_ids if item_id not in collected]
            if not missing:
                break
            pending = [item for item in batch if item.get("ITEM_ID", "") in missing]
            emit("generate_status", {"message": f"누락 bullet {len(missing)}건을 재요청하는 중입니다."})
        merged_bullets["ITEMS"].extend(collected[item_id] for item_id in expected_ids if item_id in collected)

    emit("generate_status", {"message": "LLM 응답을 JSON으로 정리하는 중입니다."})
    merged = merge_bullets_into_report(report_base, merged_bullets)
    merged["_openai_response_id"] = ", ".join(x for x in response_ids if x)
    merged["_openai_output_text"] = json.dumps(
        {
            "batch_size": batch_size,
            "expected_count": len(report_items),
            "received_count": len(merged_bullets["ITEMS"]),
            "responses": response_logs,
        },
        ensure_ascii=False,
        indent=2,
    )
    return merged


def request_openai_bullets(api_key: str, prompt: str, report_items: list[dict], emit) -> tuple[dict, str, dict]:
    materials = build_bullet_request_context(report_items)
    if "{{INPUT_MATERIALS}}" in prompt:
        prompt_text = prompt.replace("{{INPUT_MATERIALS}}", materials)
    else:
        prompt_text = prompt + "\n\n[입력 자료]\n" + materials
    content = [{"type": "input_text", "text": prompt_text}]
    pdf_items = [item for item in report_items if item.get("_LOCAL_PATH")]
    for index, item in enumerate(pdf_items, start=1):
        path = Path(item.get("_LOCAL_PATH", ""))
        if not path or not path.exists() or path.suffix.lower() != ".pdf":
            continue
        emit("generate_upload", {"index": index, "total": len(pdf_items), "title": item.get("TITLE", "")})
        uploaded = multipart_request(
            "https://api.openai.com/v1/files",
            {"purpose": "user_data"},
            [("file", path)],
            api_key,
        )
        content.append({"type": "input_file", "file_id": uploaded["id"]})
    response = openai_json_request(
        "/responses",
        {
            "model": OPENAI_MODEL,
            "input": [{"role": "user", "content": content}],
            "text": {"format": {"type": "json_object"}},
        },
        api_key,
    )
    text = extract_response_text(response)
    parsed = extract_json_object(text)
    return response, text, parsed


def build_bullet_request_context(report_base_or_items) -> str:
    rows = []
    source_items = report_base_or_items if isinstance(report_base_or_items, list) else iter_report_items(report_base_or_items)
    for item in source_items:
        rows.append({
            "ITEM_ID": item.get("ITEM_ID", ""),
            "SECTION": item.get("_SECTION", ""),
            "NO": item.get("NO", ""),
            "SOURCE_NAME": item.get("SOURCE_NAME", ""),
            "TITLE": item.get("TITLE", ""),
            "URL": item.get("URL", ""),
            "PUBLISHED_MM_DD": item.get("PUBLISHED_MM_DD", ""),
            "LOCAL_PATH": item.get("_LOCAL_PATH", ""),
            "CONTENT_TEXT": item.get("_CONTENT_TEXT", ""),
        })
    return json.dumps(rows, ensure_ascii=False, indent=2)


def build_report_base(items: list[collector.Item]) -> dict:
    bank_items = [item for item in items if item.category == "은행/지주사"]
    other_items = [item for item in items if item.category == "그외"]
    agency_items = [item for item in items if item.category == "국가기관"]
    research_items = [item for item in items if item.category == "금융연구소"]
    return {
        "BANK_SECTION": {
            "ITEMS": metadata_items(bank_items, "bank", "1", "BANK_SECTION.ITEMS"),
            "OTHER_ITEMS": metadata_items(other_items, "other", "", "BANK_SECTION.OTHER_ITEMS"),
        },
        "GOVERNMENT_SECTION": {
            "ITEMS": metadata_items(agency_items, "agency", "2", "GOVERNMENT_SECTION.ITEMS"),
        },
        "FINANCIAL_RESEARCH_SECTION": {
            "ITEMS": metadata_items(research_items, "research", "3", "FINANCIAL_RESEARCH_SECTION.ITEMS"),
        },
    }


def metadata_items(items: list[collector.Item], id_prefix: str, section_no: str, section_name: str) -> list[dict]:
    out = []
    for idx, item in enumerate(items, start=1):
        url = item.original_url or item.url
        out.append({
            "ITEM_ID": f"{id_prefix}_{idx}",
            "_SECTION": section_name,
            "_LOCAL_PATH": item.local_path,
            "NO": f"{section_no}.{idx}" if section_no else "",
            "SOURCE_NAME": clean_cell_value(item.source_name),
            "TITLE": clean_title(item.title, item.source_name),
            "URL": clean_cell_value(url),
            "PUBLISHED_MM_DD": clean_mm_dd(item.published_mm_dd),
            "_CONTENT_TEXT": content_text_for_item(item),
            "SUMMARY_BULLET_1": "",
            "SUMMARY_BULLET_2": "",
            "SUMMARY_BULLET_3": "",
        })
    return out


def content_text_for_item(item: collector.Item, limit: int = 5000) -> str:
    text = ""
    if isinstance(item.extra, dict):
        text = item.extra.get("article_text") or item.extra.get("rss_description") or ""
    text = clean_cell_value(text)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + " ..."


def merge_bullets_into_report(report_base: dict, bullet_json: dict) -> dict:
    by_id = {item.get("ITEM_ID", ""): item for item in iter_report_items(report_base)}
    by_url = {item.get("URL", ""): item for item in iter_report_items(report_base) if item.get("URL")}
    by_title = {item.get("TITLE", ""): item for item in iter_report_items(report_base) if item.get("TITLE")}
    for raw in extract_bullet_items(bullet_json):
        if not isinstance(raw, dict):
            continue
        item_id = clean_cell_value(raw.get("ITEM_ID"))
        target = by_id.get(item_id)
        if not target:
            target = by_url.get(clean_cell_value(raw.get("URL")))
        if not target:
            target = by_title.get(clean_title(raw.get("TITLE"), raw.get("SOURCE_NAME")))
        if not target:
            continue
        for key in ["SUMMARY_BULLET_1", "SUMMARY_BULLET_2", "SUMMARY_BULLET_3"]:
            target[key] = clean_cell_value(raw.get(key))
    return strip_internal_report_fields(report_base)


def extract_bullet_items(data: dict) -> list:
    if isinstance(data.get("ITEMS"), list):
        return data["ITEMS"]
    if isinstance(data.get("BULLETS"), list):
        return data["BULLETS"]
    out = []
    for section in ["BANK_SECTION", "GOVERNMENT_SECTION", "FINANCIAL_RESEARCH_SECTION"]:
        section_data = data.get(section, {})
        if not isinstance(section_data, dict):
            continue
        if isinstance(section_data.get("ITEMS"), list):
            out.extend(section_data["ITEMS"])
        if isinstance(section_data.get("OTHER_ITEMS"), list):
            out.extend(section_data["OTHER_ITEMS"])
    return out


def iter_report_items(report_data: dict):
    yield from report_data.get("BANK_SECTION", {}).get("ITEMS", [])
    yield from report_data.get("BANK_SECTION", {}).get("OTHER_ITEMS", [])
    yield from report_data.get("GOVERNMENT_SECTION", {}).get("ITEMS", [])
    yield from report_data.get("FINANCIAL_RESEARCH_SECTION", {}).get("ITEMS", [])


def strip_internal_report_fields(report_data: dict) -> dict:
    cleaned = deepcopy(report_data)
    for item in iter_report_items(cleaned):
        for key in ["ITEM_ID", "_SECTION", "_LOCAL_PATH", "_CONTENT_TEXT"]:
            item.pop(key, None)
    return cleaned


def extract_response_text(response: dict) -> str:
    if response.get("output_text"):
        return str(response["output_text"])
    parts = []
    for output in response.get("output", []):
        for content in output.get("content", []):
            if "text" in content:
                parts.append(str(content["text"]))
    return "\n".join(parts)


def extract_json_object(text: str) -> dict:
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            return json.loads(match.group(0))
        raise


def generate_docx(report_json: dict, items: list[collector.Item], output_path: Path) -> None:
    template = TEMPLATE_PATH
    doc = Document(template) if template.exists() else Document()
    add_section_bookmarks(doc)
    request_word_field_update(doc)
    report_data = normalize_report_json(report_json, items)
    populate_report_tables(doc, report_data)
    replace_docx_placeholders(
        doc,
        {
            "BANK_SECTION_PAGE": "__PAGEREF_BANK_SECTION__",
            "AGENCY_SECTION_PAGE": "__PAGEREF_AGENCY_SECTION__",
            "RESEARCH_SECTION_PAGE": "__PAGEREF_RESEARCH_SECTION__",
        },
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(f".{output_path.stem}.{uuid.uuid4().hex[:8]}.tmp{output_path.suffix}")
    try:
        doc.save(tmp_path)
        replace_file_with_retry(tmp_path, output_path)
    finally:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


def replace_file_with_retry(src: Path, dst: Path, attempts: int = 5) -> None:
    last_exc: Exception | None = None
    for attempt in range(attempts):
        try:
            os.replace(src, dst)
            return
        except PermissionError as exc:
            last_exc = exc
            time.sleep(0.5 * (attempt + 1))
    raise PermissionError(
        f"{dst} 파일을 덮어쓸 수 없습니다. 보고서가 열려 있다면 닫은 뒤 다시 재생성하세요."
    ) from last_exc


def build_template_values(report_json: dict, items: list[collector.Item]) -> dict:
    values = {k: stringify_template_value(v) for k, v in report_json.items()}
    values.setdefault("REPORT_WEEK", dt.date.today().strftime("%Y-%m-%d"))
    values.setdefault("DEPARTMENT_NAME", "조사연구")
    groups = {
        "BANK_DETAIL_ITEMS": [x for x in items if "뉴스" in x.category or "은행" in x.category or "지주" in x.category],
        "AGENCY_DETAIL_ITEMS": [x for x in items if "기관" in x.category],
        "RESEARCH_DETAIL_ITEMS": [x for x in items if "연구" in x.category],
    }
    for key, group_items in groups.items():
        values.setdefault(key, format_detail_items(report_json.get(key), group_items))
    values["BANK_SECTION_PAGE"] = "__PAGEREF_BANK_SECTION__"
    values["AGENCY_SECTION_PAGE"] = "__PAGEREF_AGENCY_SECTION__"
    values["RESEARCH_SECTION_PAGE"] = "__PAGEREF_RESEARCH_SECTION__"
    values.setdefault("ITEM_NO", "")
    values.setdefault("TITLE", "")
    values.setdefault("URL", "")
    values.setdefault("SOURCE_NAME", "")
    values.setdefault("PUBLISHED_MM_DD", "")
    values.setdefault("SUMMARY_BULLET_1", "")
    values.setdefault("SUMMARY_BULLET_2", "")
    values.setdefault("SUMMARY_BULLET_3", "")
    values.setdefault("NO", "")
    return values


def normalize_report_json(report_json: dict, items: list[collector.Item]) -> dict:
    data = {
        "BANK_SECTION": {
            "ITEMS": normalize_items(report_json.get("BANK_SECTION", {}).get("ITEMS", []), "1"),
            "OTHER_ITEMS": normalize_items(report_json.get("BANK_SECTION", {}).get("OTHER_ITEMS", []), ""),
        },
        "GOVERNMENT_SECTION": {
            "ITEMS": normalize_items(report_json.get("GOVERNMENT_SECTION", {}).get("ITEMS", []), "2"),
        },
        "FINANCIAL_RESEARCH_SECTION": {
            "ITEMS": normalize_items(report_json.get("FINANCIAL_RESEARCH_SECTION", {}).get("ITEMS", []), "3"),
        },
    }
    if not any(data[section]["ITEMS"] for section in ["BANK_SECTION", "GOVERNMENT_SECTION", "FINANCIAL_RESEARCH_SECTION"]):
        data = fallback_report_data(items)
    return data


def normalize_items(raw_items, section_no: str) -> list[dict]:
    if not isinstance(raw_items, list):
        return []
    out = []
    for idx, raw in enumerate(raw_items, start=1):
        if not isinstance(raw, dict):
            continue
        item = {
            "NO": clean_cell_value(raw.get("NO") or (f"{section_no}.{idx}" if section_no else "")),
            "SOURCE_NAME": clean_cell_value(raw.get("SOURCE_NAME")),
            "TITLE": clean_title(raw.get("TITLE"), raw.get("SOURCE_NAME")),
            "URL": clean_cell_value(raw.get("URL")),
            "PUBLISHED_MM_DD": clean_mm_dd(raw.get("PUBLISHED_MM_DD")),
            "SUMMARY_BULLET_1": clean_cell_value(raw.get("SUMMARY_BULLET_1")),
            "SUMMARY_BULLET_2": clean_cell_value(raw.get("SUMMARY_BULLET_2")),
            "SUMMARY_BULLET_3": clean_cell_value(raw.get("SUMMARY_BULLET_3")),
        }
        out.append(item)
    return out


def fallback_report_data(items: list[collector.Item]) -> dict:
    groups = {
        "BANK_SECTION": [x for x in items if "뉴스" in x.category or "은행" in x.category or "지주" in x.category],
        "GOVERNMENT_SECTION": [x for x in items if "기관" in x.category],
        "FINANCIAL_RESEARCH_SECTION": [x for x in items if "연구" in x.category],
    }
    return {
        "BANK_SECTION": {
            "ITEMS": fallback_items(groups["BANK_SECTION"], "1"),
            "OTHER_ITEMS": [],
        },
        "GOVERNMENT_SECTION": {"ITEMS": fallback_items(groups["GOVERNMENT_SECTION"], "2")},
        "FINANCIAL_RESEARCH_SECTION": {"ITEMS": fallback_items(groups["FINANCIAL_RESEARCH_SECTION"], "3")},
    }


def fallback_items(items: list[collector.Item], section_no: str) -> list[dict]:
    out = []
    for idx, item in enumerate(items, start=1):
        url = item.original_url or item.url
        out.append(
            {
                "NO": f"{section_no}.{idx}",
                "SOURCE_NAME": item.source_name,
                "TITLE": clean_title(item.title, item.source_name),
                "URL": url,
                "PUBLISHED_MM_DD": item.published_mm_dd,
                "SUMMARY_BULLET_1": "",
                "SUMMARY_BULLET_2": "",
                "SUMMARY_BULLET_3": "",
            }
        )
    return out


def clean_cell_value(value) -> str:
    return "" if value is None else re.sub(r"\s+", " ", str(value)).strip()


def clean_mm_dd(value) -> str:
    text = clean_cell_value(value)
    match = re.fullmatch(r"0?(\d{1,2})[.\-/]0?(\d{1,2})", text)
    if not match:
        return ""
    month = int(match.group(1))
    day = int(match.group(2))
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return ""
    return f"{month}.{day}"


def clean_title(value, source_name: str = "") -> str:
    text = clean_cell_value(value)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s+[가-힣]{2,4}(?:\s*,\s*[가-힣]{2,4}){0,5}\s+20\d{2}[.\-/]\d{1,2}[.\-/]\d{1,2}\s+조회수\s*\d+\s*$", "", text)
    text = re.sub(r"\s+20\d{2}[.\-/]\d{1,2}[.\-/]\d{1,2}\s+조회수\s*\d+\s*$", "", text)
    text = re.sub(r"\s+(작성자|등록일|작성일|조회수)\s*[:：]?\s*.*$", "", text)
    source_name = clean_cell_value(source_name)
    if source_name:
        suffix = f" - {source_name}"
        while text.endswith(suffix):
            text = text[: -len(suffix)].strip()
    return text.lstrip("▣").strip()


def row_for_client(item: collector.Item) -> dict:
    row = item.row()
    row["title"] = clean_title(row.get("title", ""), row.get("source_name", ""))
    row["published_date"] = collector.normalize_date(row.get("published_date", "")) or row.get("published_date", "")
    return row


def populate_report_tables(doc: Document, data: dict) -> None:
    if len(doc.tables) < 10:
        return
    fill_summary_table(doc.tables[2], data["BANK_SECTION"]["ITEMS"], keep_min=1)
    fill_other_summary_table(doc.tables[3], data["BANK_SECTION"]["OTHER_ITEMS"])
    fill_summary_table(doc.tables[4], data["GOVERNMENT_SECTION"]["ITEMS"], keep_min=1)
    fill_summary_table(doc.tables[5], data["FINANCIAL_RESEARCH_SECTION"]["ITEMS"], keep_min=1)
    fill_detail_table(doc.tables[6], data["BANK_SECTION"]["ITEMS"], keep_min=1)
    fill_other_detail_table(doc.tables[7], data["BANK_SECTION"]["OTHER_ITEMS"])
    fill_detail_table(doc.tables[8], data["GOVERNMENT_SECTION"]["ITEMS"], keep_min=1)
    fill_detail_table(doc.tables[9], data["FINANCIAL_RESEARCH_SECTION"]["ITEMS"], keep_min=1)


def fill_summary_table(table, items: list[dict], keep_min: int = 1) -> None:
    desired = max(len(items), keep_min)
    adjust_table_rows(table, header_rows=2, desired_data_rows=desired)
    set_fixed_table_widths(table, [650, 1600, 7550])
    rows = items + [blank_item()] * (desired - len(items))
    for row, item in zip(table.rows[2:], rows):
        set_cell_text(row.cells[0], item.get("NO", ""))
        set_cell_text(row.cells[1], item.get("SOURCE_NAME", ""))
        set_cell_text(row.cells[2], summary_title(item))


def fill_other_summary_table(table, items: list[dict]) -> None:
    adjust_table_rows(table, header_rows=0, desired_data_rows=1)
    set_fixed_table_widths(table, [650, 1600, 7550])
    row = table.rows[0]
    set_cell_text(row.cells[0], "그\n외")
    set_cell_text(row.cells[1], "\n".join(item.get("SOURCE_NAME", "") for item in items))
    set_cell_text(row.cells[2], "\n".join(summary_title(item) for item in items))


def fill_detail_table(table, items: list[dict], keep_min: int = 1) -> None:
    desired = max(len(items), keep_min)
    adjust_detail_rows(table, desired)
    set_fixed_table_widths(table, [650, 9150])
    rows = items + [blank_item()] * (desired - len(items))
    for idx, item in enumerate(rows):
        title_row = table.rows[1 + idx * 2]
        bullet_row = table.rows[2 + idx * 2]
        no = item.get("NO", "")
        set_cell_text(title_row.cells[0], no)
        set_cell_text(title_row.cells[1], detail_title(item))
        set_cell_text(bullet_row.cells[0], "")
        set_cell_text(bullet_row.cells[1], detail_summary(item))


def fill_other_detail_table(table, items: list[dict]) -> None:
    desired = max(len(items), 1)
    adjust_other_detail_rows(table, desired)
    set_fixed_table_widths(table, [650, 9150])
    rows = items + [blank_item()] * (desired - len(items))
    for idx, item in enumerate(rows):
        title_row = table.rows[idx * 2]
        bullet_row = table.rows[idx * 2 + 1]
        set_cell_text(title_row.cells[0], "그\n외" if idx == 0 else "")
        set_cell_text(title_row.cells[1], detail_title(item))
        set_cell_text(bullet_row.cells[0], "")
        set_cell_text(bullet_row.cells[1], detail_summary(item))


def adjust_table_rows(table, header_rows: int, desired_data_rows: int) -> None:
    while len(table.rows) < header_rows + desired_data_rows:
        clone_row(table, table.rows[-1])
    while len(table.rows) > header_rows + desired_data_rows:
        delete_row(table, table.rows[-1])


def adjust_detail_rows(table, desired_items: int) -> None:
    if len(table.rows) < 3:
        return
    title_row_xml = deepcopy(table.rows[1]._tr)
    bullet_row_xml = deepcopy(table.rows[2]._tr)
    remove_vertical_merges(title_row_xml)
    remove_vertical_merges(bullet_row_xml)
    while len(table.rows) > 1:
        delete_row(table, table.rows[-1])
    for _ in range(desired_items):
        table._tbl.append(deepcopy(title_row_xml))
        table._tbl.append(deepcopy(bullet_row_xml))


def adjust_other_detail_rows(table, desired_items: int) -> None:
    if len(table.rows) < 2:
        return
    title_row_xml = deepcopy(table.rows[0]._tr)
    bullet_row_xml = deepcopy(table.rows[1]._tr)
    remove_vertical_merges(title_row_xml)
    remove_vertical_merges(bullet_row_xml)
    while len(table.rows) > 0:
        delete_row(table, table.rows[-1])
    for _ in range(desired_items):
        table._tbl.append(deepcopy(title_row_xml))
        table._tbl.append(deepcopy(bullet_row_xml))


def clone_row(table, row) -> None:
    table._tbl.append(deepcopy(row._tr))


def delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def remove_vertical_merges(row_xml) -> None:
    for tc_pr in row_xml.iter(qn("w:tcPr")):
        for v_merge in list(tc_pr.findall(qn("w:vMerge"))):
            tc_pr.remove(v_merge)


def set_fixed_table_widths(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        seen = set()
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths) or id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


URL_RE = re.compile(r"https?://[^\s)]+")


def add_hyperlink(paragraph, url: str, text: str | None = None) -> None:
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_hyperlinks(paragraph, line: str) -> None:
    pos = 0
    for match in URL_RE.finditer(line):
        if match.start() > pos:
            paragraph.add_run(line[pos:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        pos = match.end()
    if pos < len(line):
        paragraph.add_run(line[pos:])


def set_cell_text(cell, text: str) -> None:
    text = text or ""
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    remove_paragraph_numbering(first)
    for child in list(first._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            first._p.remove(child)
    lines = str(text).split("\n")
    for index, line in enumerate(lines):
        if index:
            first.add_run().add_break()
        add_text_with_hyperlinks(first, line)


def remove_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def blank_item() -> dict:
    return {
        "NO": "",
        "SOURCE_NAME": "",
        "TITLE": "",
        "URL": "",
        "PUBLISHED_MM_DD": "",
        "SUMMARY_BULLET_1": "",
        "SUMMARY_BULLET_2": "",
        "SUMMARY_BULLET_3": "",
    }


def summary_title(item: dict) -> str:
    title = item.get("TITLE", "")
    if not title:
        return ""
    mmdd = item.get("PUBLISHED_MM_DD", "")
    return f"▣ {title}" + (f" ({mmdd})" if mmdd else "")


def detail_title(item: dict) -> str:
    title = item.get("TITLE", "")
    if not title:
        return ""
    url = item.get("URL", "")
    meta = " / ".join(x for x in [item.get("SOURCE_NAME", ""), item.get("PUBLISHED_MM_DD", "")] if x)
    second = f"{url} ({meta})" if url and meta else url or (f"({meta})" if meta else "")
    return f"▣ {title}" + (f"\n{second}" if second else "")


def detail_summary(item: dict) -> str:
    bullets = [
        item.get("SUMMARY_BULLET_1", ""),
        item.get("SUMMARY_BULLET_2", ""),
        item.get("SUMMARY_BULLET_3", ""),
    ]
    return "\n".join(format_detail_bullet(b) for b in bullets if clean_cell_value(b))


def format_detail_bullet(value: str) -> str:
    text = clean_cell_value(value)
    text = re.sub(r"^(?:[\-•·–—]\s*)+", "", text).strip()
    return f"- {text}" if text else ""


def detail_block(item: dict) -> str:
    parts = [detail_title(item), detail_summary(item)]
    return "\n".join(part for part in parts if part)


def compact_url(url: str, limit: int = 76) -> str:
    url = clean_cell_value(url)
    if not url:
        return ""
    parsed = urllib.parse.urlparse(url)
    if parsed.netloc:
        value = parsed.netloc + parsed.path
        if parsed.query:
            value += "?..."
    else:
        value = url
    if len(value) <= limit:
        return value
    return value[: max(0, limit - 3)].rstrip("/") + "..."


def stringify_template_value(value) -> str:
    if isinstance(value, (list, dict)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return "" if value is None else str(value)


def format_detail_items(generated, fallback_items: list[collector.Item]) -> str:
    source = generated if isinstance(generated, list) else []
    lines = []
    if source:
        for idx, item in enumerate(source, start=1):
            if not isinstance(item, dict):
                lines.append(str(item))
                continue
            lines.extend(
                [
                    f"{idx}. {item.get('TITLE') or item.get('title') or ''}",
                    f"{item.get('URL') or item.get('url') or ''} ({item.get('SOURCE_NAME') or item.get('source_name') or ''} / {item.get('PUBLISHED_MM_DD') or ''})",
                    f"- {item.get('SUMMARY_BULLET_1') or ''}",
                    f"- {item.get('SUMMARY_BULLET_2') or ''}",
                    f"- {item.get('SUMMARY_BULLET_3') or ''}",
                    "",
                ]
            )
    else:
        for idx, item in enumerate(fallback_items, start=1):
            lines.extend([
                f"{idx}. {item.title}",
                f"{item.original_url or item.url} ({item.source_name} / {item.published_mm_dd})",
                "",
            ])
    return "\n".join(lines).strip()


def replace_docx_placeholders(doc: Document, values: dict) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, values)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, values)


def replace_in_paragraph(paragraph, values: dict) -> None:
    text = paragraph.text
    if "{{" not in text:
        return
    for key, value in values.items():
        text = text.replace("{{" + key + "}}", str(value))
    paragraph.clear()
    add_text_with_pageref_fields(paragraph, text)


def add_text_with_pageref_fields(paragraph, text: str) -> None:
    markers = {
        "__PAGEREF_BANK_SECTION__": "bank_section",
        "__PAGEREF_AGENCY_SECTION__": "agency_section",
        "__PAGEREF_RESEARCH_SECTION__": "research_section",
    }
    pattern = "(" + "|".join(re.escape(marker) for marker in markers) + ")"
    for part in re.split(pattern, text):
        if not part:
            continue
        bookmark = markers.get(part)
        if bookmark:
            add_pageref_field(paragraph, bookmark)
        else:
            paragraph.add_run(part)


def add_pageref_field(paragraph, bookmark_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    placeholder = OxmlElement("w:t")
    placeholder.text = "0"
    run._r.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_section_bookmarks(doc: Document) -> None:
    targets = {
        "1. 은행/지주사": "bank_section",
        "2. 국가기관": "agency_section",
        "3. 금융연구소": "research_section",
    }
    candidates: dict[str, list] = {bookmark: [] for bookmark in targets.values()}
    for paragraph in iter_doc_paragraphs(doc):
        text = paragraph.text.strip()
        for label, bookmark in targets.items():
            if text == label:
                candidates[bookmark].append(paragraph)
    bookmark_id = 100
    for bookmark, paragraphs in candidates.items():
        if not paragraphs:
            continue
        add_bookmark_to_paragraph(paragraphs[-1], bookmark, bookmark_id)
        bookmark_id += 1


def iter_doc_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def add_bookmark_to_paragraph(paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._p
    if not paragraph.runs:
        paragraph.add_run("")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def request_word_field_update(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def run_collect_job(job_id: str, data: dict) -> None:
    def emit(event_type: str, payload: dict | None = None) -> None:
        payload = dict(payload or {})
        payload["type"] = event_type
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if job:
            job["queue"].put(payload)

    try:
        news_start = parse_iso_date(data["news_start"])
        news_end = parse_iso_date(data["news_end"])
        agency_start = parse_iso_date(data["agency_start"])
        agency_end = parse_iso_date(data["agency_end"])
        research_start = parse_iso_date(data["research_start"])
        research_end = parse_iso_date(data["research_end"])
        news_bank_max = int(data.get("news_bank_max") or 10)
        news_other_max = int(data.get("news_other_max") or 10)
        agency_max = int(data.get("agency_max") or 10)
        research_max = int(data.get("research_max") or 10)
        news_keywords = load_news_keywords()
        include_news = bool(data.get("include_news", True))
        include_agency = bool(data.get("include_agency", True))
        include_research = bool(data.get("include_research", True))
        run_dir = collector.make_run_dir()
        emit("started", {"run_dir": str(run_dir)})

        items = collector.collect_by_ranges(
            news_start,
            news_end,
            agency_start,
            agency_end,
            research_start,
            research_end,
            max_per_source=None,
            dry_run=False,
            news_bank_max=news_bank_max,
            news_other_max=news_other_max,
            agency_max=agency_max,
            research_max=research_max,
            output_dir=run_dir,
            progress=emit,
            news_keywords=news_keywords,
            include_news=include_news,
            include_agency=include_agency,
            include_research=include_research,
        )
        normalize_article_urls(items)
        collector.write_outputs(items, output_dir=run_dir)
        payload = []
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = ""
            for idx, item in enumerate(items):
                row = row_for_client(item)
                row["id"] = idx
                payload.append(row)
        emit("complete", {"items": payload, "run_dir": str(run_dir)})
    except Exception as exc:
        traceback.print_exc()
        emit("fatal", {"error": str(exc)})


def write_items(items: list[collector.Item], csv_path: Path, xlsx_path: Path) -> None:
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = list(collector.Item("", "", "", "").row().keys())
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for item in items:
            writer.writerow(row_for_client(item))

    wb = Workbook()
    ws = wb.active
    ws.title = "selected_metadata"
    ws.append(fieldnames)
    for item in items:
        row = row_for_client(item)
        ws.append([row[k] for k in fieldnames])
    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col[:100])
        ws.column_dimensions[col[0].column_letter].width = min(max(max_len + 2, 12), 45)
    wb.save(xlsx_path)


def main() -> int:
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    print(f"조사연구 도우미: http://{HOST}:{PORT}")
    if os.environ.get("RRA_NO_BROWSER") != "1":
        threading.Timer(0.8, lambda: webbrowser.open(f"http://{HOST}:{PORT}/")).start()
    server.serve_forever()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
