from __future__ import annotations

import datetime as dt
import base64
import ctypes
import ctypes.wintypes
import json
import os
import re
import shutil
import sys
import threading
import time
import traceback
import urllib.parse
import urllib.request
import webbrowser
import csv
import queue
import uuid
from copy import deepcopy
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

sys.dont_write_bytecode = True

import collector
import paths
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from openpyxl import Workbook


HOST = "127.0.0.1"
PORT = 8765
APP_DIR = paths.APP_DIR
APP_DATA_DIR = paths.APP_DATA_DIR
CONFIG_DIR = paths.CONFIG_DIR
PROMPT_PATH = paths.RESOURCES_DIR / "report_prompt.md"
TEMPLATE_PATH = paths.RESOURCES_DIR / "report_template.docx"
KEY_PATH = CONFIG_DIR / "openai_key.bin"
NEWS_KEYWORDS_PATH = CONFIG_DIR / "news_keywords.json"
OPENAI_MODEL = "gpt-5.4-mini"
DEFAULT_REPORT_PROMPT = """Return only one JSON object for filling the weekly report template.
Use the provided metadata and attached PDFs.

Required top-level keys:
- REPORT_WEEK
- DEPARTMENT_NAME
- BANK_SECTION_PAGE
- AGENCY_SECTION_PAGE
- RESEARCH_SECTION_PAGE
- BANK_DETAIL_ITEMS
- AGENCY_DETAIL_ITEMS
- RESEARCH_DETAIL_ITEMS

Each DETAIL_ITEMS value should be an array of objects with ITEM_NO, TITLE, URL, SOURCE_NAME, PUBLISHED_MM_DD, SUMMARY_BULLET_1, SUMMARY_BULLET_2, SUMMARY_BULLET_3.
"""

paths.ensure_app_dirs()
paths.migrate_legacy_data()
paths.copy_default_resource("report_prompt.md", DEFAULT_REPORT_PROMPT)
paths.copy_default_resource("report_template.docx")

STATE_LOCK = threading.Lock()
STATE: dict = {"items": [], "run_dir": None, "report_path": ""}
JOBS: dict[str, dict] = {}


INDEX_HTML = r"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>議곗궗?곌뎄 ?꾩슦誘?/title>
  <style>
    :root {
      color-scheme: light;
      --bg: #f6f7f9;
      --panel: #ffffff;
      --text: #17202a;
      --muted: #6b7280;
      --line: #d8dee8;
      --brand: #1f6feb;
      --brand-dark: #1557bd;
      --ok: #0f7b3d;
      --warn: #b45309;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: "Malgun Gothic", "Segoe UI", sans-serif;
      background: var(--bg);
      color: var(--text);
    }
    header {
      padding: 18px 24px 14px;
      background: var(--panel);
      border-bottom: 1px solid var(--line);
    }
    .brand {
      display: flex;
      align-items: center;
      gap: 12px;
    }
    .brand-mark {
      width: 38px;
      height: 38px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      border-radius: 8px;
      background: #eef4ff;
      color: #1557bd;
      border: 1px solid #c7ddff;
      font-size: 16px;
      font-weight: 900;
    }
    .brand-copy {
      min-width: 0;
    }
    h1 {
      margin: 0;
      font-size: 21px;
      font-weight: 700;
      letter-spacing: 0;
    }
    main { padding: 18px 24px 28px; }
    .controls {
      display: grid;
      grid-template-columns: repeat(3, minmax(260px, 1fr));
      gap: 12px;
      align-items: stretch;
      margin-bottom: 14px;
    }
    .group {
      min-height: 238px;
      display: flex;
      flex-direction: column;
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 12px;
    }
    .group-title {
      font-size: 13px;
      font-weight: 700;
      margin-bottom: 10px;
    }
    .group-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      margin-bottom: 10px;
      min-height: 38px;
    }
    .group-head .group-title {
      margin-bottom: 0;
    }
    .header-action {
      height: 38px;
      min-width: 96px;
      padding: 0 10px;
      font-size: 13px;
    }
    .header-placeholder {
      width: 96px;
      height: 38px;
      flex: 0 0 auto;
    }
    .date-row {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 8px;
    }
    .limit-row {
      margin-top: 10px;
    }
    .card-spacer {
      display: none;
    }
    .button-row {
      display: flex;
      align-items: center;
      justify-content: flex-start;
      gap: 8px;
      margin-top: auto;
      padding-top: 10px;
    }
    .button-row button {
      width: auto;
      min-width: 96px;
      padding: 0 10px;
    }
    .collect-row {
      margin-top: auto;
      padding-top: 10px;
    }
    .collect-row .button-row {
      margin-top: 0;
      padding-top: 0;
    }
    label {
      display: block;
      font-size: 12px;
      color: var(--muted);
    }
    input[type="date"], input[type="number"] {
      width: 100%;
      margin-top: 4px;
      padding: 8px 9px;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
      background: #fff;
    }
    .side {
      display: flex;
      justify-content: flex-end;
      margin: -2px 0 14px;
    }
    .side button {
      min-width: 104px;
      padding: 0 12px;
    }
    button {
      height: 38px;
      border: 1px solid var(--brand);
      border-radius: 6px;
      background: var(--brand);
      color: #fff;
      font-weight: 700;
      cursor: pointer;
    }
    button:hover { background: var(--brand-dark); }
    button.secondary {
      background: var(--brand);
      color: #fff;
      border-color: var(--brand);
    }
    button.secondary:hover { background: var(--brand-dark); }
    button.jb {
      background: #eef4ff;
      color: #1557bd;
      border-color: #c7ddff;
    }
    button.jb:hover {
      background: #dceaff;
      border-color: #9cc5ff;
    }
    button.ghost {
      background: #fff;
      color: var(--text);
      border-color: var(--line);
    }
    button.ghost:hover {
      background: #eef2f7;
    }
    button.danger {
      background: var(--brand);
      color: #fff;
      border-color: var(--brand);
    }
    button.danger:hover { background: var(--brand-dark); }
    button:disabled {
      opacity: 0.55;
      cursor: not-allowed;
    }
    .status {
      min-height: 28px;
      margin: 8px 0 12px;
      color: var(--muted);
      font-size: 13px;
    }
    .toolbar {
      display: flex;
      gap: 8px;
      align-items: center;
      justify-content: space-between;
      margin: 10px 0;
    }
    .summary {
      color: var(--muted);
      font-size: 13px;
    }
    .table-wrap {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: auto;
      max-height: calc(100vh - 270px);
      cursor: default;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
      cursor: default;
    }
    th, td {
      border-bottom: 1px solid var(--line);
      padding: 8px 10px;
      vertical-align: middle;
      white-space: nowrap;
      cursor: default;
    }
    th {
      position: sticky;
      top: 0;
      background: #f1f4f8;
      text-align: left;
      z-index: 1;
    }
    td.title {
      min-width: 420px;
      max-width: 760px;
      white-space: normal;
      line-height: 1.35;
    }
    .action-cell {
      width: 64px;
      text-align: center;
    }
    .open-item-btn {
      width: 34px;
      height: 30px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      padding: 0;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .open-item-btn:hover {
      background: #eef4ff;
      border-color: #9cc5ff;
      color: #1557bd;
    }
    .open-item-btn svg {
      width: 17px;
      height: 17px;
      stroke: currentColor;
      stroke-width: 2.2;
      fill: none;
    }
    tr:hover { background: #f7fbff; }
    .tabs {
      display: flex;
      flex-wrap: wrap;
      gap: 6px;
      margin: 8px 0 10px;
    }
    .tab {
      height: 32px;
      padding: 0 11px;
      border: 1px solid var(--line);
      border-radius: 999px;
      background: #fff;
      color: var(--text);
      font-weight: 600;
      cursor: pointer;
    }
    .tab.active {
      background: #e8f1ff;
      border-color: #9cc5ff;
      color: #1557bd;
    }
    .badge {
      display: inline-flex;
      min-width: 24px;
      height: 22px;
      align-items: center;
      justify-content: center;
      border-radius: 999px;
      background: #eef4ff;
      color: #1d4ed8;
      font-weight: 700;
      font-size: 12px;
    }
    .kind {
      color: var(--muted);
      font-size: 12px;
    }
    .modal-backdrop {
      position: fixed;
      inset: 0;
      display: none;
      align-items: center;
      justify-content: center;
      padding: 24px;
      background: rgba(15, 23, 42, 0.42);
      z-index: 20;
    }
    .modal-backdrop.open { display: flex; }
    .modal {
      width: min(760px, 100%);
      max-height: min(720px, calc(100vh - 48px));
      display: grid;
      grid-template-rows: auto auto 1fr;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: 0 24px 70px rgba(15, 23, 42, 0.28);
      overflow: hidden;
    }
    .modal-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 16px 18px;
      border-bottom: 1px solid var(--line);
    }
    .modal-title {
      font-size: 17px;
      font-weight: 800;
    }
    .modal-close {
      width: 34px;
      height: 34px;
      padding: 0;
      background: #fff;
      color: var(--text);
      border-color: var(--line);
      font-size: 18px;
      line-height: 1;
    }
    .modal-close:hover { background: #eef2f7; }
    .modal-help {
      padding: 10px 18px;
      color: var(--muted);
      font-size: 13px;
    }
    .history-list {
      overflow: auto;
      padding: 8px;
    }
    .history-select-all {
      display: grid;
      grid-template-columns: 44px 1fr;
      gap: 8px;
      align-items: center;
      padding: 8px 8px 4px;
    }
    .history-select-all-box {
      height: 36px;
      display: flex;
      align-items: center;
      justify-content: center;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 8px;
      cursor: pointer;
    }
    .history-select-all-box:hover {
      background: #f7fbff;
      border-color: #9cc5ff;
    }
    .history-select-all-text {
      color: var(--muted);
      font-size: 12px;
      font-weight: 800;
    }
    .history-row {
      display: grid;
      grid-template-columns: 44px 1fr;
      gap: 8px;
      align-items: stretch;
      margin: 0 0 6px;
    }
    .history-check-wrap {
      display: flex;
      align-items: center;
      justify-content: center;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 8px;
      cursor: pointer;
    }
    .history-check-wrap:hover {
      background: #f7fbff;
      border-color: #9cc5ff;
    }
    .history-item {
      width: 100%;
      height: auto;
      display: grid;
      grid-template-columns: 1fr auto;
      gap: 12px;
      align-items: center;
      padding: 12px 14px;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 8px;
      text-align: left;
      font-weight: 600;
      cursor: pointer;
    }
    .history-item > .history-select {
      display: none;
    }
    .history-select {
      width: 16px;
      height: 16px;
      cursor: pointer;
      margin: 0;
    }
    .history-actions {
      display: inline-flex;
      gap: 8px;
      align-items: center;
      justify-content: flex-end;
      flex-wrap: wrap;
    }
    .history-item:hover {
      background: #f7fbff;
      border-color: #9cc5ff;
    }
    .history-name {
      font-size: 14px;
      font-weight: 800;
    }
    .history-meta {
      margin-top: 4px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 500;
    }
    .history-tags {
      display: inline-flex;
      gap: 6px;
      align-items: center;
      justify-content: flex-end;
      flex-wrap: wrap;
    }
    .history-tag {
      display: inline-flex;
      min-height: 24px;
      align-items: center;
      padding: 3px 8px;
      border-radius: 999px;
      background: #eef4ff;
      color: #1d4ed8;
      font-size: 12px;
      font-weight: 800;
    }
    .history-tag.report {
      background: #ecfdf3;
      color: #0f7b3d;
    }
    .history-delete {
      height: 30px;
      padding: 0 10px;
      background: #fff;
      color: #b42318;
      border: 1px solid #f1b9b4;
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .history-delete:hover {
      background: #fff1f0;
      border-color: #e0776e;
    }
    .history-bulk-actions {
      display: flex;
      padding: 14px 18px 18px;
      border-top: 1px solid #eef2f7;
      justify-content: flex-end;
    }
    .danger-outline {
      background: #fff;
      color: #b42318;
      border: 1px solid #f1b9b4;
    }
    .danger-outline:hover {
      background: #fff1f0;
      border-color: #e0776e;
      color: #b42318;
    }
    .keyword-modal {
      width: min(720px, calc(100vw - 40px));
    }
    .keyword-body {
      padding: 20px 22px 22px;
    }
    .keyword-editor {
      display: grid;
      grid-template-columns: minmax(0, 1fr) 86px;
      gap: 10px;
      margin-bottom: 14px;
    }
    .keyword-editor input {
      width: 100%;
      height: 38px;
      padding: 0 12px;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
    }
    .keyword-tabs {
      display: flex;
      gap: 8px;
      margin-bottom: 14px;
    }
    .keyword-tab {
      height: 32px;
      padding: 0 10px;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .keyword-tab.active {
      color: #fff;
      border-color: var(--brand);
      background: var(--brand);
    }
    .keyword-list {
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      max-height: 300px;
      overflow: auto;
      padding: 12px;
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #f8fafc;
    }
    .keyword-item {
      width: auto;
      min-height: 32px;
      padding: 0 11px;
      text-align: center;
      color: var(--text);
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 999px;
      font-size: 13px;
      font-weight: 700;
    }
    .keyword-item.selected {
      background: #eef4ff;
      border-color: #1f6feb;
      color: #174ea6;
    }
    .source-list {
      display: grid;
      gap: 8px;
      max-height: 340px;
      overflow: auto;
    }
    .source-item {
      min-height: 42px;
      display: flex;
      align-items: center;
      padding: 10px 12px;
      color: var(--text);
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
      font-weight: 700;
    }
    .modal-actions {
      display: flex;
      justify-content: space-between;
      gap: 8px;
      margin-top: 16px;
    }
    .modal-actions.history-bulk-actions {
      justify-content: flex-end;
    }
    .modal-actions.history-bulk-actions #historyBulkDeleteBtn {
      margin-left: auto;
    }
    .empty-state {
      padding: 36px 18px;
      color: var(--muted);
      text-align: center;
      font-size: 13px;
    }
    @media (max-width: 1100px) {
      .controls { grid-template-columns: 1fr; }
    }
  </style>
</head>
<body>
  <header>
    <div class="brand">
      <div class="brand-mark">JB</div>
      <div class="brand-copy">
        <h1>議곗궗?곌뎄 ?꾩슦誘?/h1>
      </div>
    </div>
  </header>
  <main>
    <section class="controls">
      <div class="group">
        <div class="group-head">
          <div class="group-title">?댁뒪</div>
          <button class="jb header-action" id="keywordBtn" type="button">寃???ㅼ썙??/button>
        </div>
        <div class="date-row">
          <label>?쒖옉??input id="newsStart" type="date"></label>
          <label>醫낅즺??input id="newsEnd" type="date"></label>
        </div>
        <div class="date-row limit-row">
          <label>???吏二쇱궗 理쒕? 嫄댁닔<input id="newsBankMax" type="number" min="1" max="100" value="10"></label>
          <label>洹몄쇅 理쒕? 嫄댁닔<input id="newsOtherMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="button-row">
          <button class="jb" id="collectNewsBtn" type="button">?먮즺 ?섏쭛</button>
        </div>
      </div>
      <div class="group">
        <div class="group-head">
          <div class="group-title">援??湲곌?</div>
          <button class="jb header-action" id="agencyListBtn" type="button">湲곌? 紐⑸줉</button>
        </div>
        <div class="date-row">
          <label>?쒖옉??input id="agencyStart" type="date"></label>
          <label>醫낅즺??input id="agencyEnd" type="date"></label>
        </div>
        <div class="limit-row">
          <label>湲곌?蹂?理쒕? 嫄댁닔<input id="agencyMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="card-spacer"></div>
        <div class="limit-row collect-row">
          <div class="button-row">
            <button class="jb" id="collectAgencyBtn" type="button">?먮즺 ?섏쭛</button>
          </div>
        </div>
      </div>
      <div class="group">
        <div class="group-head">
          <div class="group-title">湲덉쑖?곌뎄??/div>
          <button class="jb header-action" id="researchListBtn" type="button">?곌뎄??紐⑸줉</button>
        </div>
        <div class="date-row">
          <label>?쒖옉??input id="researchStart" type="date"></label>
          <label>醫낅즺??input id="researchEnd" type="date"></label>
        </div>
        <div class="limit-row">
          <label>?곌뎄?뚮퀎 理쒕? 嫄댁닔<input id="researchMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="card-spacer"></div>
        <div class="limit-row collect-row">
          <div class="button-row">
            <button class="jb" id="collectResearchBtn" type="button">?먮즺 ?섏쭛</button>
          </div>
        </div>
      </div>
    </section>

    <div class="side">
      <button class="secondary" id="collectBtn">?쇨큵 ?섏쭛</button>
    </div>

    <div class="status" id="status">湲곌컙???좏깮?????섏쭛???쒖옉?섏꽭??</div>

    <div class="toolbar">
      <div class="summary" id="summary">?좏깮 0嫄?/ ?꾩껜 0嫄?/div>
      <div>
        <button class="ghost" id="historyBtn">?ㅽ뻾 湲곕줉</button>
        <button class="secondary" id="reportBtn">蹂닿퀬???앹꽦</button>
        <button class="ghost" id="openReportBtn" disabled>蹂닿퀬???닿린</button>
      </div>
    </div>

    <div class="tabs" id="tabs"></div>

    <div class="table-wrap">
      <table>
        <thead>
          <tr>
            <th><input id="headCheck" type="checkbox"></th>
            <th>援щ텇</th>
            <th>?몃줎/湲곌?/?곌뎄??/th>
            <th>?쇱옄</th>
            <th>?쒕ぉ</th>
            <th>?뺤떇</th>
            <th></th>
          </tr>
        </thead>
        <tbody id="tbody"></tbody>
      </table>
    </div>
  </main>

  <div class="modal-backdrop" id="historyModal" aria-hidden="true">
    <section class="modal" role="dialog" aria-modal="true" aria-labelledby="historyTitle">
      <div class="modal-head">
        <div class="modal-title" id="historyTitle">?ㅽ뻾 湲곕줉</div>
        <button class="modal-close" id="historyCloseBtn" type="button" title="?リ린">횞</button>
      </div>
      <div class="modal-help">遺덈윭???ㅽ뻾 湲곕줉???좏깮?섏꽭?? 蹂닿퀬?쒓? ?앹꽦??湲곕줉? 蹂꾨룄濡??쒖떆?⑸땲??</div>
      <div class="history-select-all">
        <label class="history-select-all-box">
          <input id="historySelectAll" class="history-select" type="checkbox" aria-label="?꾩껜 ?좏깮">
        </label>
        <div class="history-select-all-text">?꾩껜 ?좏깮</div>
      </div>
      <div class="history-list" id="historyList"></div>
      <div class="modal-actions history-bulk-actions">
        <button class="danger-outline" id="historyBulkDeleteBtn" type="button" disabled>?좏깮 ??젣</button>
      </div>
    </section>
  </div>

  <div class="modal-backdrop" id="keywordModal" aria-hidden="true">
    <section class="modal keyword-modal" role="dialog" aria-modal="true" aria-labelledby="keywordTitle">
      <div class="modal-head">
        <div class="modal-title" id="keywordTitle">?댁뒪 寃???ㅼ썙??/div>
        <button class="modal-close" id="keywordCloseBtn" type="button" title="?リ린">횞</button>
      </div>
      <div class="modal-help">寃?됲븷 臾멸뎄瑜?異붽??섍굅????젣?섏꽭?? ??λ맂 ?ㅼ썙?쒕뒗 ?ㅼ쓬 ?섏쭛怨????ъ떎????洹몃?濡??ъ슜?⑸땲??</div>
      <div class="keyword-body">
        <div class="keyword-tabs">
          <button class="keyword-tab active" id="keywordBankTab" type="button">???吏二쇱궗</button>
          <button class="keyword-tab" id="keywordOtherTab" type="button">洹몄쇅</button>
        </div>
        <div class="keyword-editor">
          <input id="keywordInput" type="text" placeholder="?? ?꾨턿???>
          <button id="keywordAddBtn" type="button">異붽?</button>
        </div>
        <div class="keyword-list" id="keywordList"></div>
        <div class="modal-actions">
          <button class="jb" id="keywordGroupAddBtn" type="button" disabled>洹몃９ 異붽?</button>
          <button class="danger-outline" id="keywordDeleteBtn" type="button">??젣</button>
          <button class="jb" id="keywordSaveBtn" type="button">???/button>
        </div>
        <div class="keyword-list" id="keywordGroupList"></div>
      </div>
    </section>
  </div>

  <div class="modal-backdrop" id="sourceModal" aria-hidden="true">
    <section class="modal keyword-modal" role="dialog" aria-modal="true" aria-labelledby="sourceTitle">
      <div class="modal-head">
        <div class="modal-title" id="sourceTitle">?섏쭛 ???/div>
        <button class="modal-close" id="sourceCloseBtn" type="button" title="?リ린">횞</button>
      </div>
      <div class="modal-help" id="sourceHelp">?꾩옱 ?ㅼ젙???섏쭛 ???紐⑸줉?낅땲??</div>
      <div class="keyword-body">
        <div class="source-list" id="sourceList"></div>
      </div>
    </section>
  </div>

  <script>
    let rows = [];
    let activeTab = "all";
    let currentReportPath = "";
    let keywordState = {bank: [], other: []};
    let activeKeywordKind = "bank";
    let selectedKeywordIndex = -1;
    let selectedKeywordIndexes = new Set();
    let selectedHistoryRuns = new Set();
    const sourceLists = {
      agency: ["湲덉쑖媛먮룆??, "湲덉쑖?꾩썝??, "?쒓뎅???],
      research: [
        "?쒓뎅湲덉쑖?곌뎄??,
        "?섎굹湲덉쑖?곌뎄???곌뎄蹂닿퀬??,
        "?섎굹湲덉쑖?곌뎄???뺢린蹂닿퀬??,
        "KB寃쎌쁺?곌뎄??,
        "KDB誘몃옒?꾨왂?곌뎄??,
        "?곕━湲덉쑖寃쎌쁺?곌뎄???곌뎄蹂닿퀬??,
        "?곕━湲덉쑖寃쎌쁺?곌뎄???숇궓??Review"
      ]
    };

    function isoDate(d) {
      return d.toISOString().slice(0, 10);
    }
    function setDefaults() {
      const end = new Date();
      const start = new Date();
      start.setDate(end.getDate() - 7);
      for (const id of ["newsStart", "agencyStart", "researchStart"]) {
        document.getElementById(id).value = isoDate(start);
      }
      for (const id of ["newsEnd", "agencyEnd", "researchEnd"]) {
        document.getElementById(id).value = isoDate(end);
      }
    }
    function payload(targets = {news: true, agency: true, research: true}) {
      return {
        news_start: document.getElementById("newsStart").value,
        news_end: document.getElementById("newsEnd").value,
        agency_start: document.getElementById("agencyStart").value,
        agency_end: document.getElementById("agencyEnd").value,
        research_start: document.getElementById("researchStart").value,
        research_end: document.getElementById("researchEnd").value,
        news_bank_max: Number(document.getElementById("newsBankMax").value || 10),
        news_other_max: Number(document.getElementById("newsOtherMax").value || 10),
        agency_max: Number(document.getElementById("agencyMax").value || 10),
        research_max: Number(document.getElementById("researchMax").value || 10),
        include_news: Boolean(targets.news),
        include_agency: Boolean(targets.agency),
        include_research: Boolean(targets.research)
      };
    }
    function updateSummary() {
      const selected = rows.filter(r => r.checked).length;
      const visible = visibleRows();
      const visibleSelected = visible.filter(r => r.checked).length;
      document.getElementById("summary").textContent = `?좏깮 ${selected}嫄?/ ?꾩껜 ${rows.length}嫄?쨌 ?꾩옱 ??${visibleSelected}/${visible.length}嫄?;
      document.getElementById("headCheck").checked = visible.length > 0 && visibleSelected === visible.length;
      updateReportControls();
    }
    function setReportPath(path) {
      currentReportPath = path || "";
      updateReportControls();
    }
    function updateReportControls() {
      const selected = rows.filter(r => r.checked).length;
      const reportBtn = document.getElementById("reportBtn");
      reportBtn.disabled = selected === 0;
      reportBtn.textContent = currentReportPath ? "蹂닿퀬???ъ깮?? : "蹂닿퀬???앹꽦";
      document.getElementById("openReportBtn").disabled = !currentReportPath;
    }
    function tabKey(row) {
      if (row.category === "그외") return "뉴스 그외";
      if (row.category === "은행/지주사") return "뉴스 은행/지주사";
      if (row.category === "국가기관") return "국가기관";
      if (row.category === "금융연구소") return "금융연구소";
      return row.category || "湲고?";
    }
    function visibleRows() {
      if (activeTab === "all") return rows;
      return rows.filter(row => tabKey(row) === activeTab);
    }
    function renderTabs() {
      const tabs = document.getElementById("tabs");
      const counts = new Map();
      for (const row of rows) {
        const key = tabKey(row);
        counts.set(key, (counts.get(key) || 0) + 1);
      }
      const tabOrder = ["뉴스 그외", "뉴스 은행/지주사", "국가기관", "금융연구소"];
      const ordered = ["all", ...tabOrder.filter(key => counts.has(key))];
      if (activeTab !== "all" && !ordered.includes(activeTab)) activeTab = "all";
      tabs.innerHTML = "";
      for (const key of ordered) {
        const btn = document.createElement("button");
        btn.className = "tab" + (key === activeTab ? " active" : "");
        btn.textContent = key === "all" ? `?꾩껜 (${rows.length})` : `${key} (${counts.get(key)})`;
        btn.addEventListener("click", () => {
          activeTab = key;
          render();
        });
        tabs.appendChild(btn);
      }
    }
    function render() {
      renderTabs();
      const tbody = document.getElementById("tbody");
      tbody.innerHTML = "";
      visibleRows().forEach((row) => {
        const index = rows.findIndex(r => r.id === row.id);
        const tr = document.createElement("tr");
        tr.dataset.id = row.id;
        tr.innerHTML = `
          <td><input type="checkbox" ${row.checked ? "checked" : ""}></td>
          <td>${escapeHtml(row.category || "")}</td>
          <td>${escapeHtml(row.source_name || "")}</td>
          <td>${escapeHtml(row.published_date || "")}</td>
          <td class="title">${escapeHtml(row.title || "")}</td>
          <td><span class="kind">${escapeHtml(row.file_type || "")}</span></td>
          <td class="action-cell"><button class="open-item-btn" type="button" title="${row.file_type === "article" ? "湲곗궗 留곹겕 ?닿린" : "PDF ?닿린"}">${openItemIcon(row)}</button></td>
        `;
        tr.querySelector("input").addEventListener("change", e => {
          rows[index].checked = e.target.checked;
          updateReportControls();
          updateSummary();
        });
        tr.querySelector(".open-item-btn").addEventListener("click", () => openItem(row));
        tbody.appendChild(tr);
      });
      updateSummary();
    }
    function openItemIcon(row) {
      if (row.file_type === "article") {
        return `<svg viewBox="0 0 24 24" aria-hidden="true"><path d="M7 17 17 7"></path><path d="M9 7h8v8"></path><path d="M5 5v14h14"></path></svg>`;
      }
      return `<svg viewBox="0 0 24 24" aria-hidden="true"><path d="M7 3h7l5 5v13H7z"></path><path d="M14 3v6h5"></path><path d="M9 14h6"></path><path d="M9 17h6"></path></svg>`;
    }
    function escapeHtml(value) {
      return String(value)
        .replaceAll("&", "&amp;")
        .replaceAll("<", "&lt;")
        .replaceAll(">", "&gt;")
        .replaceAll('"', "&quot;");
    }
    function setCollectButtonsDisabled(disabled) {
      for (const id of ["collectBtn", "collectNewsBtn", "collectAgencyBtn", "collectResearchBtn"]) {
        document.getElementById(id).disabled = disabled;
      }
    }
    async function collect(targets = {news: true, agency: true, research: true}) {
      setCollectButtonsDisabled(true);
      setReportPath("");
      updateReportControls();
      let streaming = false;
      document.getElementById("status").textContent = "?섏쭛 以묒엯?덈떎. ?ъ씠???묐떟???곕씪 1~3遺??뺣룄 嫄몃┫ ???덉뒿?덈떎.";
      try {
        const res = await fetch("/collect-start", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify(payload(targets))
        });
        const data = await res.json();
        if (!res.ok || !data.ok) throw new Error(data.error || "failed to start collection");
        const events = new EventSource(`/collect-events?id=${encodeURIComponent(data.job_id)}`);
        events.onmessage = (event) => {
          const msg = JSON.parse(event.data);
          if (msg.type === "source_start") {
            document.getElementById("status").textContent = `Collecting: ${msg.source}`;
          } else if (msg.type === "source_done") {
            document.getElementById("status").textContent = `${msg.source}: ${msg.count} item(s) found`;
          } else if (msg.type === "source_error") {
            document.getElementById("status").textContent = `${msg.source}: ${msg.error}`;
          } else if (msg.type === "decode_start") {
            document.getElementById("status").textContent = `Decoding article URL (${msg.index}/${msg.total}): ${msg.title || ""}`;
          } else if (msg.type === "decode_done") {
            document.getElementById("status").textContent = `Decoded article URL (${msg.index}/${msg.total})`;
          } else if (msg.type === "download_start") {
            document.getElementById("status").textContent = `Saving PDF (${msg.index}/${msg.total}): ${msg.source}`;
          } else if (msg.type === "complete") {
            rows = msg.items.map(x => ({...x, checked: true}));
            setReportPath("");
            activeTab = "all";
            render();
            document.getElementById("status").textContent = `Collection complete: ${rows.length} item(s) 쨌 ${msg.run_dir}`;
            events.close();
            setCollectButtonsDisabled(false);
          } else if (msg.type === "fatal") {
            document.getElementById("status").textContent = "Error: " + msg.error;
            events.close();
            setCollectButtonsDisabled(false);
          }
        };
        events.onerror = () => {
          document.getElementById("status").textContent = "Progress connection closed.";
          events.close();
          setCollectButtonsDisabled(false);
        };
        streaming = true;
        return;
        if (!res.ok || !data.ok) throw new Error(data.error || "?섏쭛 ?ㅽ뙣");
        rows = data.items.map(x => ({...x, checked: true}));
        activeTab = "all";
        render();
        document.getElementById("status").textContent = `?섏쭛 ?꾨즺: ${rows.length}嫄?쨌 ????대뜑: ${data.run_dir}`;
      } catch (err) {
        document.getElementById("status").textContent = "?ㅻ쪟: " + err.message;
      } finally {
        if (!streaming) setCollectButtonsDisabled(false);
      }
    }
    async function openItem(row) {
      const res = await fetch("/open", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({id: row.id})
      });
      const data = await res.json();
      if (data.open_in_client && data.target) {
        window.open(data.target, "_blank", "noopener");
        return;
      }
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "?먮즺 ?닿린 ?ㅻ쪟: " + (data.error || "?????놁뒿?덈떎.");
      }
    }
    async function openReport() {
      if (!currentReportPath) return;
      const res = await fetch("/open-report", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "蹂닿퀬???닿린 ?ㅻ쪟: " + (data.error || "?????놁뒿?덈떎.");
      }
    }
    async function saveSelection() {
      const ids = rows.filter(r => r.checked).map(r => r.id);
      const res = await fetch("/save-selection", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({ids})
      });
      const data = await res.json();
      if (data.ok) {
        document.getElementById("status").textContent = `?좏깮 ??ぉ ????꾨즺: ${ids.length}嫄?쨌 ${data.csv}`;
      } else {
        document.getElementById("status").textContent = "????ㅻ쪟: " + data.error;
      }
    }
    async function ensureApiKey() {
      const statusRes = await fetch("/api-key-status");
      const status = await statusRes.json();
      if (status.has_key) {
        const useExisting = confirm("??λ맂 OpenAI API key瑜??ъ슜?좉퉴?? 痍⑥냼瑜??꾨Ⅴ硫???key瑜??낅젰?⑸땲??");
        if (useExisting) return true;
      }
      const key = prompt("OpenAI API key瑜??낅젰?섏꽭?? 濡쒖뺄 ?뚯씪???뷀샇????λ맗?덈떎.");
      if (!key) return false;
      const res = await fetch("/api-key", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({api_key: key})
      });
      const data = await res.json();
      if (!data.ok) {
        alert("API key ????ㅽ뙣: " + data.error);
        return false;
      }
      return true;
    }
    async function generateReport() {
      const ids = rows.filter(r => r.checked).map(r => r.id);
      if (!ids.length) {
        alert("蹂닿퀬?쒖뿉 ?ы븿???먮즺瑜??섎굹 ?댁긽 ?좏깮?섏꽭??");
        return;
      }
      const selectedRows = rows.filter(r => r.checked);
      const missingSections = [];
      if (!selectedRows.some(r => r.category === "은행/지주사" || r.category === "그외")) {
        missingSections.push("뉴스");
      }
      if (!selectedRows.some(r => r.category === "국가기관")) {
        missingSections.push("국가기관");
      }
      if (!selectedRows.some(r => r.category === "금융연구소")) {
        missingSections.push("금융연구소");
      }
      if (missingSections.length) {
        const okToContinue = confirm(`${missingSections.join(", ")} ?섏쭛 ?먮즺媛 ?놁뒿?덈떎.\n?먮즺媛 ?녿뒗 ?뱀뀡? 鍮꾩슦怨?蹂닿퀬?쒕? ?앹꽦?섏떆寃좎뒿?덇퉴?`);
        if (!okToContinue) return;
      }
      const ok = await ensureApiKey();
      if (!ok) return;
      const btn = document.getElementById("reportBtn");
      btn.disabled = true;
      const isRegenerate = Boolean(currentReportPath);
      document.getElementById("status").textContent = isRegenerate
        ? "蹂닿퀬???ъ깮??以묒엯?덈떎. 泥댄겕???먮즺 湲곗??쇰줈 湲곗〈 寃곌낵臾쇱쓣 ?泥댄빀?덈떎."
        : "蹂닿퀬???앹꽦 以묒엯?덈떎. ?좏깮?섏? ?딆? PDF瑜??뺣━?섍퀬 OpenAI API瑜??몄텧?⑸땲??";
      try {
        const res = await fetch("/generate-report-start", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({ids})
        });
        const data = await res.json();
        if (!res.ok || !data.ok) throw new Error(data.error || "蹂닿퀬???앹꽦 ?쒖옉 ?ㅽ뙣");
        const events = new EventSource(`/collect-events?id=${encodeURIComponent(data.job_id)}`);
        events.onmessage = (event) => {
          const msg = JSON.parse(event.data);
          if (msg.type === "generate_status") {
            document.getElementById("status").textContent = msg.message || "蹂닿퀬???앹꽦 以묒엯?덈떎.";
          } else if (msg.type === "generate_upload") {
            document.getElementById("status").textContent = `?먮즺 ?낅줈??以?(${msg.index}/${msg.total}): ${msg.title || ""}`;
          } else if (msg.type === "complete") {
            rows = msg.items.map(x => ({...x, checked: true}));
            setReportPath(msg.docx || "");
            activeTab = "all";
            render();
            document.getElementById("status").textContent = `${isRegenerate ? "蹂닿퀬???ъ깮?? : "蹂닿퀬???앹꽦"} ?꾨즺. 蹂닿퀬???닿린 踰꾪듉???뚮윭 ?뺤씤?섏꽭??`;
            events.close();
            updateReportControls();
          } else if (msg.type === "fatal") {
            document.getElementById("status").textContent = "蹂닿퀬???앹꽦 ?ㅻ쪟: " + msg.error;
            events.close();
            updateReportControls();
          }
        };
        events.onerror = () => {
          document.getElementById("status").textContent = "蹂닿퀬???앹꽦 ?곹깭 ?곌껐??醫낅즺?섏뿀?듬땲??";
          events.close();
          updateReportControls();
        };
      } catch (err) {
        document.getElementById("status").textContent = "蹂닿퀬???앹꽦 ?ㅻ쪟: " + err.message;
        updateReportControls();
      }
    }
    async function loadHistory() {
      const res = await fetch("/runs");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "?ㅽ뻾 湲곕줉 ?ㅻ쪟: " + data.error;
        return;
      }
      const label = data.runs.map((r, i) => `${i + 1}. ${r.name} (${r.count}嫄?`).join("\n");
      const picked = prompt("遺덈윭???ㅽ뻾 湲곕줉 踰덊샇瑜??낅젰?섏꽭??\n\n" + label);
      const index = Number(picked) - 1;
      if (!Number.isInteger(index) || index < 0 || index >= data.runs.length) return;
      const run = data.runs[index];
      const detail = await fetch(`/run?id=${encodeURIComponent(run.name)}`);
      const loaded = await detail.json();
      if (!loaded.ok) {
        document.getElementById("status").textContent = "?ㅽ뻾 湲곕줉 遺덈윭?ㅺ린 ?ㅻ쪟: " + loaded.error;
        return;
      }
      rows = loaded.items.map(x => ({...x, checked: true}));
      setReportPath(loaded.docx || "");
      activeTab = "all";
      render();
      document.getElementById("status").textContent = `?ㅽ뻾 湲곕줉 遺덈윭?ㅺ린 ?꾨즺: ${run.name}`;
    }
    async function loadHistoryModal() {
      const res = await fetch("/runs");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "?ㅽ뻾 湲곕줉 ?ㅻ쪟: " + data.error;
        return;
      }
      showHistoryModal(data.runs);
    }
    function showHistoryModal(runs) {
      const modal = document.getElementById("historyModal");
      const list = document.getElementById("historyList");
      selectedHistoryRuns = new Set();
      updateHistoryBulkDelete();
      list.innerHTML = "";
      if (!runs.length) {
        list.innerHTML = `<div class="empty-state">??λ맂 ?ㅽ뻾 湲곕줉???놁뒿?덈떎.</div>`;
      }
      runs.forEach((run) => {
        const row = document.createElement("div");
        row.className = "history-row";
        const checkWrap = document.createElement("label");
        checkWrap.className = "history-check-wrap";
        checkWrap.innerHTML = `<input class="history-select" type="checkbox" aria-label="?ㅽ뻾 湲곕줉 ?좏깮">`;
        const item = document.createElement("div");
        item.className = "history-item";
        item.innerHTML = `
          <input class="history-select" type="checkbox" aria-label="?ㅽ뻾 湲곕줉 ?좏깮">
          <div>
            <div class="history-name">${escapeHtml(formatRunName(run.name))}</div>
            <div class="history-meta">${escapeHtml(run.path || "")}</div>
          </div>
          <div class="history-actions">
            <div class="history-tags">
              <span class="history-tag">${Number(run.count || 0)}嫄?/span>
              ${run.has_report ? `<span class="history-tag report">蹂닿퀬??/span>` : ""}
            </div>
            <button class="history-delete" type="button" title="?ㅽ뻾 湲곕줉 ??젣">??젣</button>
          </div>
        `;
        item.addEventListener("click", () => openRunFromModal(run.name));
        const checkbox = checkWrap.querySelector(".history-select");
        checkbox.addEventListener("change", () => {
          if (checkbox.checked) {
            selectedHistoryRuns.add(run.name);
          } else {
            selectedHistoryRuns.delete(run.name);
          }
          updateHistoryBulkDelete();
        });
        item.querySelector(".history-delete").addEventListener("click", (event) => {
          event.stopPropagation();
          deleteRunFromModal(run.name);
        });
        row.appendChild(checkWrap);
        row.appendChild(item);
        list.appendChild(row);
      });
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
    }
    function updateHistoryBulkDelete() {
      const btn = document.getElementById("historyBulkDeleteBtn");
      if (!btn) return;
      const count = selectedHistoryRuns.size;
      btn.disabled = count === 0;
      const all = document.getElementById("historySelectAll");
      if (all) {
        const checkboxes = [...document.querySelectorAll("#historyList .history-check-wrap .history-select")];
        all.checked = checkboxes.length > 0 && checkboxes.every(box => box.checked);
        all.indeterminate = count > 0 && !all.checked;
      }
      btn.textContent = count ? `?좏깮 ??젣 (${count})` : "?좏깮 ??젣";
    }
    function setAllHistorySelected(checked) {
      document.querySelectorAll("#historyList .history-check-wrap .history-select").forEach((box) => {
        if (box.checked !== checked) {
          box.checked = checked;
          box.dispatchEvent(new Event("change"));
        }
      });
      updateHistoryBulkDelete();
    }
    function closeHistoryModal() {
      const modal = document.getElementById("historyModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }
    function formatRunName(name) {
      const m = String(name).match(/^(\d{2})(\d{2})(\d{2})_(\d{2})(\d{2})(?:_(\d+))?$/);
      if (!m) return name;
      const suffix = m[6] ? ` #${m[6]}` : "";
      return `20${m[1]}.${m[2]}.${m[3]} ${m[4]}:${m[5]}${suffix}`;
    }
    async function openRunFromModal(runName) {
      closeHistoryModal();
      const detail = await fetch(`/run?id=${encodeURIComponent(runName)}`);
      const loaded = await detail.json ();
      if (!loaded.ok) {
        document.getElementById("status").textContent = "?ㅽ뻾 湲곕줉 遺덈윭?ㅺ린 ?ㅻ쪟: " + loaded.error;
        return;
      }
      rows = loaded.items.map(x => ({...x, checked: true}));
      setReportPath(loaded.docx || "");
      activeTab = "all";
      render();
      document.getElementById("status").textContent = `?ㅽ뻾 湲곕줉 遺덈윭?ㅺ린 ?꾨즺: ${runName}`;
    }
    async function deleteRunFromModal(runName) {
      if (!confirm(`${formatRunName(runName)} ?ㅽ뻾 湲곕줉????젣?좉퉴??\n\n?대떦 ?대뜑??PDF, JSON, ?앹꽦 蹂닿퀬?쒓? 紐⑤몢 ??젣?⑸땲??`)) {
        return;
      }
      const res = await fetch("/delete-run", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({id: runName})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "?ㅽ뻾 湲곕줉 ??젣 ?ㅻ쪟: " + (data.error || "??젣?????놁뒿?덈떎.");
        return;
      }
      if (data.cleared_current) {
        rows = [];
        setReportPath("");
        activeTab = "all";
        render();
      }
      document.getElementById("status").textContent = `?ㅽ뻾 湲곕줉 ??젣 ?꾨즺: ${runName}`;
      await loadHistoryModal();
    }

    async function deleteSelectedHistoryRuns() {
      const runNames = [...selectedHistoryRuns];
      if (!runNames.length) return;
      if (!confirm(`?좏깮???ㅽ뻾 湲곕줉 ${runNames.length}媛쒕? ??젣?좉퉴??\n\n?대떦 ?대뜑??PDF, JSON, ?앹꽦 蹂닿퀬?쒓? 紐⑤몢 ??젣?⑸땲??`)) {
        return;
      }
      const btn = document.getElementById("historyBulkDeleteBtn");
      btn.disabled = true;
      let deleted = 0;
      let clearedCurrent = false;
      for (const runName of runNames) {
        const res = await fetch("/delete-run", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({id: runName})
        });
        const data = await res.json();
        if (!res.ok || !data.ok) {
          document.getElementById("status").textContent = `?ㅽ뻾 湲곕줉 ??젣 ?ㅻ쪟: ${runName} 쨌 ${data.error || "??젣?????놁뒿?덈떎."}`;
          await loadHistoryModal();
          return;
        }
        deleted += 1;
        clearedCurrent = clearedCurrent || Boolean(data.cleared_current);
      }
      if (clearedCurrent) {
        rows = [];
        setReportPath("");
        activeTab = "all";
        render();
      }
      document.getElementById("status").textContent = `?ㅽ뻾 湲곕줉 ${deleted}媛???젣 ?꾨즺`;
      await loadHistoryModal();
    }

    async function openKeywordModal() {
      const res = await fetch("/news-keywords");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "키워드 불러오기 오류: " + data.error;
        return;
      }
      keywordState = data.keywords;
      activeKeywordKind = "bank";
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      renderKeywordModal();
      const modal = document.getElementById("keywordModal");
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
      document.getElementById("keywordInput").focus();
    }

    function closeKeywordModal() {
      const modal = document.getElementById("keywordModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }

    function setKeywordKind(kind) {
      activeKeywordKind = kind;
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      renderKeywordModal();
    }

    function currentKeywordSection() {
      const section = keywordState[activeKeywordKind] || {keywords: [], groups: []};
      section.keywords = section.keywords || [];
      section.groups = section.groups || [];
      keywordState[activeKeywordKind] = section;
      return section;
    }

    function renderKeywordModal() {
      document.getElementById("keywordBankTab").className = "keyword-tab" + (activeKeywordKind === "bank" ? " active" : "");
      document.getElementById("keywordOtherTab").className = "keyword-tab" + (activeKeywordKind === "other" ? " active" : "");
      const section = currentKeywordSection();
      const list = document.getElementById("keywordList");
      const items = section.keywords;
      list.innerHTML = "";
      if (!items.length) {
        list.innerHTML = `<div class="empty-state">등록된 키워드가 없습니다.</div>`;
      }
      items.forEach((keyword, index) => {
        const btn = document.createElement("button");
        btn.type = "button";
        btn.className = "keyword-item" + (selectedKeywordIndexes.has(index) ? " selected" : "");
        btn.textContent = keyword;
        btn.addEventListener("click", () => {
          if (selectedKeywordIndexes.has(index)) selectedKeywordIndexes.delete(index);
          else selectedKeywordIndexes.add(index);
          selectedKeywordIndex = index;
          renderKeywordModal();
        });
        btn.addEventListener("dblclick", () => addKeywordGroup([keyword]));
        list.appendChild(btn);
      });
      document.getElementById("keywordGroupAddBtn").disabled = selectedKeywordIndexes.size === 0;
      const groupList = document.getElementById("keywordGroupList");
      groupList.innerHTML = "";
      if (!section.groups.length) {
        groupList.innerHTML = `<div class="empty-state">등록된 키워드 그룹이 없습니다.</div>`;
      }
      section.groups.forEach((group, index) => {
        const btn = document.createElement("button");
        btn.type = "button";
        btn.className = "keyword-item";
        btn.textContent = `${group.join(" AND ")}  ×`;
        btn.title = "클릭하면 그룹이 삭제됩니다.";
        btn.addEventListener("click", () => {
          section.groups.splice(index, 1);
          renderKeywordModal();
        });
        groupList.appendChild(btn);
      });
    }

    function addKeyword() {
      const input = document.getElementById("keywordInput");
      const value = input.value.trim();
      if (!value) return;
      const section = currentKeywordSection();
      if (!section.keywords.includes(value)) {
        section.keywords.push(value);
        selectedKeywordIndex = section.keywords.length - 1;
        selectedKeywordIndexes = new Set([selectedKeywordIndex]);
      }
      input.value = "";
      renderKeywordModal();
      input.focus();
    }

    function deleteKeyword() {
      const section = currentKeywordSection();
      const selected = [...selectedKeywordIndexes].sort((a, b) => b - a);
      if (!selected.length) return;
      const deleted = new Set();
      for (const index of selected) {
        if (index >= 0 && index < section.keywords.length) {
          deleted.add(section.keywords[index]);
          section.keywords.splice(index, 1);
        }
      }
      section.groups = section.groups.filter(group => !group.some(keyword => deleted.has(keyword)));
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      renderKeywordModal();
    }

    function addSelectedKeywordGroup() {
      const section = currentKeywordSection();
      const group = [...selectedKeywordIndexes].sort((a, b) => a - b).map(index => section.keywords[index]).filter(Boolean);
      addKeywordGroup(group);
    }

    function addKeywordGroup(group) {
      if (!group.length) return;
      const section = currentKeywordSection();
      const key = JSON.stringify(group);
      if (!section.groups.some(existing => JSON.stringify(existing) === key)) {
        section.groups.push(group);
      }
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      renderKeywordModal();
    }

    async function saveKeywords() {
      const res = await fetch("/news-keywords", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({keywords: keywordState})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "키워드 저장 오류: " + (data.error || "저장할 수 없습니다.");
        return;
      }
      keywordState = data.keywords;
      closeKeywordModal();
      document.getElementById("status").textContent = "뉴스 검색 키워드 저장 완료";
    }
    function openSourceModal(kind) {
      const title = kind === "agency" ? "湲곌? 紐⑸줉" : "?곌뎄??紐⑸줉";
      const help = kind === "agency"
        ? "援??湲곌? ?먮즺 ?섏쭛 ???꾨옒 湲곌??먯꽌 ?먮즺瑜?媛?몄샃?덈떎."
        : "湲덉쑖?곌뎄???먮즺 ?섏쭛 ???꾨옒 ?곌뎄?뚯뿉???먮즺瑜?媛?몄샃?덈떎.";
      document.getElementById("sourceTitle").textContent = title;
      document.getElementById("sourceHelp").textContent = help;
      const list = document.getElementById("sourceList");
      list.innerHTML = "";
      for (const source of sourceLists[kind] || []) {
        const item = document.createElement("div");
        item.className = "source-item";
        item.textContent = source;
        list.appendChild(item);
      }
      const modal = document.getElementById("sourceModal");
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
    }

    function closeSourceModal() {
      const modal = document.getElementById("sourceModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }

    document.getElementById("collectBtn").addEventListener("click", () => collect({news: true, agency: true, research: true}));
    document.getElementById("collectNewsBtn").addEventListener("click", () => collect({news: true, agency: false, research: false}));
    document.getElementById("collectAgencyBtn").addEventListener("click", () => collect({news: false, agency: true, research: false}));
    document.getElementById("collectResearchBtn").addEventListener("click", () => collect({news: false, agency: false, research: true}));
    document.getElementById("reportBtn").addEventListener("click", generateReport);
    document.getElementById("openReportBtn").addEventListener("click", openReport);
    document.getElementById("historyBtn").addEventListener("click", loadHistoryModal);
    document.getElementById("keywordBtn").addEventListener("click", openKeywordModal);
    document.getElementById("agencyListBtn").addEventListener("click", () => openSourceModal("agency"));
    document.getElementById("researchListBtn").addEventListener("click", () => openSourceModal("research"));
    document.getElementById("historyCloseBtn").addEventListener("click", closeHistoryModal);
    document.getElementById("historyBulkDeleteBtn").addEventListener("click", deleteSelectedHistoryRuns);
    document.getElementById("historySelectAll").addEventListener("change", e => setAllHistorySelected(e.target.checked));
    document.getElementById("keywordCloseBtn").addEventListener("click", closeKeywordModal);
    document.getElementById("sourceCloseBtn").addEventListener("click", closeSourceModal);
    document.getElementById("keywordBankTab").addEventListener("click", () => setKeywordKind("bank"));
    document.getElementById("keywordOtherTab").addEventListener("click", () => setKeywordKind("other"));
    document.getElementById("keywordAddBtn").addEventListener("click", addKeyword);
    document.getElementById("keywordDeleteBtn").addEventListener("click", deleteKeyword);
    document.getElementById("keywordGroupAddBtn").addEventListener("click", addSelectedKeywordGroup);
    document.getElementById("keywordSaveBtn").addEventListener("click", saveKeywords);
    document.getElementById("keywordInput").addEventListener("keydown", e => {
      if (e.key === "Enter") addKeyword();
    });
    document.getElementById("historyModal").addEventListener("click", e => {
      if (e.target.id === "historyModal") closeHistoryModal();
    });
    document.getElementById("keywordModal").addEventListener("click", e => {
      if (e.target.id === "keywordModal") closeKeywordModal();
    });
    document.getElementById("sourceModal").addEventListener("click", e => {
      if (e.target.id === "sourceModal") closeSourceModal();
    });
    document.addEventListener("keydown", e => {
      if (e.key === "Escape") {
        closeHistoryModal();
        closeKeywordModal();
        closeSourceModal();
      }
    });
    document.getElementById("headCheck").addEventListener("change", e => {
      visibleRows().forEach(r => r.checked = e.target.checked);
      render();
    });
    setDefaults();
    updateReportControls();
  </script>
</body>
</html>
"""


class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt: str, *args) -> None:
        print(f"[app] {self.address_string()} {fmt % args}")

    def do_GET(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/":
            self.send_text(INDEX_HTML, "text/html; charset=utf-8")
            return
        if parsed.path == "/collect-events":
            self.handle_collect_events(parsed)
            return
        if parsed.path == "/api-key-status":
            self.send_json({"ok": True, "has_key": KEY_PATH.exists()})
            return
        if parsed.path == "/news-keywords":
            self.send_json({"ok": True, "keywords": load_news_keywords()})
            return
        if parsed.path == "/runs":
            self.handle_runs()
            return
        if parsed.path == "/run":
            self.handle_run(parsed)
            return
        self.send_json({"ok": False, "error": "not found"}, status=404)

    def do_POST(self) -> None:
        try:
            if self.path == "/collect-start":
                self.handle_collect_start()
            elif self.path == "/collect":
                self.handle_collect()
            elif self.path == "/open":
                self.handle_open()
            elif self.path == "/open-report":
                self.handle_open_report()
            elif self.path == "/save-selection":
                self.handle_save_selection()
            elif self.path == "/api-key":
                self.handle_api_key()
            elif self.path == "/news-keywords":
                self.handle_news_keywords()
            elif self.path == "/generate-report-start":
                self.handle_generate_report_start()
            elif self.path == "/generate-report":
                self.handle_generate_report()
            elif self.path == "/delete-run":
                self.handle_delete_run()
            else:
                self.send_json({"ok": False, "error": "not found"}, status=404)
        except Exception as exc:
            traceback.print_exc()
            self.send_json({"ok": False, "error": str(exc)}, status=500)

    def handle_collect(self) -> None:
        data = self.read_json()
        news_start = parse_iso_date(data["news_start"])
        news_end = parse_iso_date(data["news_end"])
        agency_start = parse_iso_date(data["agency_start"])
        agency_end = parse_iso_date(data["agency_end"])
        research_start = parse_iso_date(data["research_start"])
        research_end = parse_iso_date(data["research_end"])
        news_bank_max = int(data.get("news_bank_max") or 10)
        news_other_max = int(data.get("news_other_max") or 10)
        agency_max = int(data.get("agency_max") or 10)
        research_max = int(data.get("research_max") or 10)
        news_keywords = load_news_keywords()
        include_news = bool(data.get("include_news", True))
        include_agency = bool(data.get("include_agency", True))
        include_research = bool(data.get("include_research", True))
        run_dir = collector.make_run_dir()

        items = collector.collect_by_ranges(
            news_start,
            news_end,
            agency_start,
            agency_end,
            research_start,
            research_end,
            max_per_source=None,
            dry_run=False,
            news_bank_max=news_bank_max,
            news_other_max=news_other_max,
            agency_max=agency_max,
            research_max=research_max,
            output_dir=run_dir,
            news_keywords=news_keywords,
            include_news=include_news,
            include_agency=include_agency,
            include_research=include_research,
        )
        collector.write_outputs(items, output_dir=run_dir)
        payload = []
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = ""
            for idx, item in enumerate(items):
                row = row_for_client(item)
                row["id"] = idx
                payload.append(row)
        self.send_json({"ok": True, "items": payload, "run_dir": str(run_dir)})

    def handle_collect_start(self) -> None:
        data = self.read_json()
        job_id = uuid.uuid4().hex
        job_queue: queue.Queue = queue.Queue()
        job = {"queue": job_queue, "done": False}
        with STATE_LOCK:
            JOBS[job_id] = job
            STATE["report_path"] = ""
        thread = threading.Thread(target=run_collect_job, args=(job_id, data), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_collect_events(self, parsed) -> None:
        params = urllib.parse.parse_qs(parsed.query)
        job_id = params.get("id", [""])[0]
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if not job:
            self.send_json({"ok": False, "error": "invalid job id"}, status=404)
            return

        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Connection", "keep-alive")
        self.end_headers()

        job_queue = job["queue"]
        while True:
            try:
                event = job_queue.get(timeout=15)
            except queue.Empty:
                event = {"type": "ping"}
            data = f"data: {json.dumps(event, ensure_ascii=False)}\n\n".encode("utf-8")
            try:
                self.wfile.write(data)
                self.wfile.flush()
            except (BrokenPipeError, ConnectionResetError):
                break
            if event.get("type") in {"complete", "fatal"}:
                with STATE_LOCK:
                    JOBS.pop(job_id, None)
                break

    def handle_open(self) -> None:
        try:
            data = self.read_json()
            item_id = int(data["id"])
            with STATE_LOCK:
                items = STATE.get("items", [])
                run_dir_value = str(STATE.get("run_dir") or "")
                if item_id < 0 or item_id >= len(items):
                    self.send_json({"ok": False, "error": "invalid id"}, status=400)
                    return
                item = items[item_id]
            if item.file_type == "article":
                target_url = resolve_article_open_url(item)
                if target_url:
                    item.original_url = target_url
                    self.send_json({"ok": True, "target": target_url, "open_in_client": True})
                    return
                self.send_json({"ok": False, "error": "湲곗궗 URL???놁뒿?덈떎."}, status=404)
                return
            run_dir = Path(run_dir_value) if run_dir_value else None
            path = resolve_item_local_path(item, run_dir)
            if path:
                os.startfile(str(path))
                self.send_json({"ok": True, "target": str(path)})
                return
            self.send_json({"ok": False, "error": "?ㅽ뻾 湲곕줉 ?대뜑?먯꽌 PDF ?뚯씪??李얠쓣 ???놁뒿?덈떎."}, status=404)
        except Exception as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=500)

    def handle_open_report(self) -> None:
        with STATE_LOCK:
            report_path = str(STATE.get("report_path") or "")
            run_dir_value = str(STATE.get("run_dir") or "")
        candidates = []
        if report_path:
            candidates.append(Path(report_path))
        if run_dir_value:
            candidates.append(Path(run_dir_value) / "generated_report.docx")
        for path in candidates:
            if path.exists() and path.is_file():
                os.startfile(path)
                self.send_json({"ok": True, "docx": str(path)})
                return
        self.send_json({"ok": False, "error": "?앹꽦??蹂닿퀬???뚯씪??李얠쓣 ???놁뒿?덈떎."}, status=404)

    def handle_save_selection(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        with STATE_LOCK:
            selected = [item for i, item in enumerate(STATE.get("items", [])) if i in ids]
            run_dir_value = STATE.get("run_dir")
        run_dir = Path(run_dir_value) if run_dir_value else collector.OUT_DIR
        selected_csv = run_dir / "selected_metadata.csv"
        selected_xlsx = run_dir / "selected_metadata.xlsx"
        write_items(selected, selected_csv, selected_xlsx)
        self.send_json(
            {
                "ok": True,
                "count": len(selected),
                "csv": str(selected_csv),
                "xlsx": str(selected_xlsx),
            }
        )

    def handle_api_key(self) -> None:
        data = self.read_json()
        api_key = str(data.get("api_key") or "").strip()
        if not api_key:
            self.send_json({"ok": False, "error": "api_key is required"}, status=400)
            return
        save_api_key(api_key)
        self.send_json({"ok": True})

    def handle_news_keywords(self) -> None:
        data = self.read_json()
        try:
            keywords = normalize_news_keywords(data.get("keywords", data))
        except ValueError as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        save_news_keywords(keywords)
        self.send_json({"ok": True, "keywords": keywords})

    def handle_generate_report(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        result = perform_generate_report(ids)
        self.send_json(result)

    def handle_generate_report_start(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        job_id = uuid.uuid4().hex
        job_queue: queue.Queue = queue.Queue()
        job = {"queue": job_queue, "done": False}
        with STATE_LOCK:
            JOBS[job_id] = job
        thread = threading.Thread(target=run_generate_report_job, args=(job_id, ids), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_runs(self) -> None:
        runs = []
        if collector.RUNS_DIR.exists():
            for path in sorted((p for p in collector.RUNS_DIR.iterdir() if p.is_dir()), reverse=True):
                metadata = path / "metadata.csv"
                count = 0
                if metadata.exists():
                    with metadata.open("r", encoding="utf-8-sig", newline="") as f:
                        count = sum(1 for _ in csv.DictReader(f))
                runs.append(
                    {
                        "name": path.name,
                        "path": str(path),
                        "count": count,
                        "has_report": (path / "generated_report.docx").exists(),
                    }
                )
        self.send_json({"ok": True, "runs": runs})

    def handle_run(self, parsed) -> None:
        run_id = urllib.parse.parse_qs(parsed.query).get("id", [""])[0]
        run_dir = resolve_run_dir(run_id)
        if run_dir is None:
            self.send_json({"ok": False, "error": "invalid run id"}, status=400)
            return
        metadata = run_dir / "metadata.csv"
        if not metadata.exists():
            self.send_json({"ok": False, "error": "metadata.csv not found"}, status=404)
            return
        items: list[collector.Item] = []
        with metadata.open("r", encoding="utf-8-sig", newline="") as f:
            for idx, row in enumerate(csv.DictReader(f)):
                item = item_from_row(row)
                items.append(item)
        payload = []
        for idx, item in enumerate(items):
            out = row_for_client(item)
            out["id"] = idx
            payload.append(out)
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = str(run_dir / "generated_report.docx") if (run_dir / "generated_report.docx").exists() else ""
        self.send_json(
            {
                "ok": True,
                "items": payload,
                "run_dir": str(run_dir),
                "docx": str(run_dir / "generated_report.docx") if (run_dir / "generated_report.docx").exists() else "",
            }
        )

    def handle_delete_run(self) -> None:
        data = self.read_json()
        run_id = str(data.get("id") or "")
        run_dir = resolve_run_dir(run_id)
        if run_dir is None:
            self.send_json({"ok": False, "error": "invalid run id"}, status=400)
            return
        if not run_dir.exists() or not run_dir.is_dir():
            self.send_json({"ok": False, "error": "run not found"}, status=404)
            return
        cleared_current = False
        with STATE_LOCK:
            current = str(STATE.get("run_dir") or "")
            if current and same_path(current, run_dir):
                STATE["items"] = []
                STATE["run_dir"] = None
                STATE["report_path"] = ""
                cleared_current = True
        shutil.rmtree(run_dir)
        self.send_json({"ok": True, "deleted": str(run_dir), "cleared_current": cleared_current})

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0"))
        raw = self.rfile.read(length).decode("utf-8")
        return json.loads(raw or "{}")

    def send_text(self, body: str, content_type: str) -> None:
        data = body.encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_json(self, data: dict, status: int = 200) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def perform_generate_report(ids: set[int], progress=None) -> dict:
    def emit(message: str) -> None:
        if progress:
            progress("generate_status", {"message": message})

    emit("?좏깮 ??ぉ???뺤씤?섎뒗 以묒엯?덈떎.")
    with STATE_LOCK:
        current_items = list(STATE.get("items", []))
        source_run_dir = str(STATE.get("run_dir") or "")
    if not current_items:
        raise RuntimeError("?섏쭛??紐⑸줉???놁뒿?덈떎.")
    selected = [item for i, item in enumerate(current_items) if i in ids]
    unselected = [item for i, item in enumerate(current_items) if i not in ids]
    existing_run_dir = resolve_existing_run_dir(source_run_dir)
    if not selected:
        raise RuntimeError("?좏깮????ぉ???놁뒿?덈떎.")

    emit("湲곗궗 ?먮Ц 留곹겕瑜??뺤씤?섎뒗 以묒엯?덈떎.")
    normalize_article_urls(selected, progress=progress)

    emit("OpenAI API key瑜??뺤씤?섎뒗 以묒엯?덈떎.")
    api_key = load_api_key()
    if not api_key:
        raise RuntimeError("OpenAI API key媛 ??λ릺???덉? ?딆뒿?덈떎.")

    emit("?좏깮?섏? ?딆? ?먮즺瑜??뺣━?섎뒗 以묒엯?덈떎.")
    if existing_run_dir:
        for item in unselected:
            delete_local_file_under(item.local_path, existing_run_dir)
    else:
        for item in unselected:
            delete_local_file(item.local_path)

    run_dir = existing_run_dir or collector.make_run_dir()
    emit("?ㅽ뻾 湲곕줉 ?대뜑瑜?以鍮꾪븯??以묒엯?덈떎.")
    if not existing_run_dir:
        move_selected_files(selected, run_dir, move=False)
    enrich_article_texts(selected, progress=progress)
    collector.write_outputs(selected, output_dir=run_dir)

    emit("?꾨＼?꾪듃? ?낅젰 ?먮즺瑜?以鍮꾪븯??以묒엯?덈떎.")
    prompt_path = ensure_prompt_template()
    prompt = prompt_path.read_text(encoding="utf-8")
    report_json = call_openai_report(api_key, prompt, selected, progress=progress)
    llm_output_text = str(report_json.pop("_openai_output_text", "") or "")
    llm_output_path = run_dir / "llm_output.txt"
    if llm_output_text:
        emit("LLM output text瑜???ν븯??以묒엯?덈떎.")
        llm_output_path.write_text(llm_output_text, encoding="utf-8")
    report_json_path = run_dir / "report_data.json"
    emit("LLM ?묐떟 JSON????ν븯??以묒엯?덈떎.")
    report_json_path.write_text(
        json.dumps(report_json, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    docx_path = run_dir / "generated_report.docx"
    emit("?뚮뱶 蹂닿퀬?쒕? ?앹꽦?섎뒗 以묒엯?덈떎.")
    generate_docx(report_json, selected, docx_path)

    payload = []
    with STATE_LOCK:
        STATE["items"] = selected
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = str(docx_path)
        for idx, item in enumerate(selected):
            row = row_for_client(item)
            row["id"] = idx
            payload.append(row)
    return {
        "ok": True,
        "items": payload,
        "run_dir": str(run_dir),
        "json": str(report_json_path),
        "llm_output": str(llm_output_path) if llm_output_text else "",
        "docx": str(docx_path),
    }


def enrich_article_texts(items: list[collector.Item], progress=None) -> None:
    article_items = [
        item for item in items
        if item.file_type == "article" and not item.extra.get("article_text") and (item.original_url or item.url)
    ]
    if not article_items:
        return
    normalize_article_urls(article_items)
    http = collector.Http()
    total = len(article_items)
    for index, item in enumerate(article_items, start=1):
        if progress:
            progress("generate_status", {"message": f"湲곗궗 蹂몃Ц 異붿텧 以?({index}/{total}): {item.title}"})
        text, note = collector.fetch_article_text(
            http,
            item.original_url or item.url,
            referer=item.extra.get("google_news_url", "") if isinstance(item.extra, dict) else "",
        )
        if text:
            item.extra["article_text"] = text
        if note:
            item.notes = "; ".join(x for x in [item.notes, note] if x)


def run_generate_report_job(job_id: str, ids: set[int]) -> None:
    def emit(event_type: str, payload: dict | None = None) -> None:
        payload = dict(payload or {})
        payload["type"] = event_type
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if job:
            job["queue"].put(payload)

    try:
        result = perform_generate_report(ids, progress=emit)
        emit("complete", result)
    except Exception as exc:
        traceback.print_exc()
        emit("fatal", {"error": str(exc)})


def parse_iso_date(value: str) -> dt.date:
    return dt.date.fromisoformat(value)


def ensure_prompt_template() -> Path:
    if not PROMPT_PATH.exists():
        PROMPT_PATH.write_text(DEFAULT_REPORT_PROMPT, encoding="utf-8")
    return PROMPT_PATH


def default_news_keywords() -> dict:
    return {
        "bank": {
            "keywords": list(collector.DEFAULT_NEWS_KEYWORDS["bank"]["keywords"]),
            "groups": [list(group) for group in collector.DEFAULT_NEWS_KEYWORDS["bank"]["groups"]],
        },
        "other": {
            "keywords": list(collector.DEFAULT_NEWS_KEYWORDS["other"]["keywords"]),
            "groups": [list(group) for group in collector.DEFAULT_NEWS_KEYWORDS["other"]["groups"]],
        },
    }


def normalize_keyword_section(value, default_section: dict) -> dict:
    if isinstance(value, list):
        raw_keywords = value
        raw_groups = [[item] for item in value]
    elif isinstance(value, dict):
        raw_keywords = value.get("keywords", default_section["keywords"])
        raw_groups = value.get("groups") or [[item] for item in raw_keywords]
    else:
        raw_keywords = default_section["keywords"]
        raw_groups = default_section["groups"]
    if not isinstance(raw_keywords, list):
        raw_keywords = default_section["keywords"]
    keywords = []
    seen = set()
    for raw in raw_keywords:
        text = str(raw or "").strip().strip('"').strip("'").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        keywords.append(text)
    groups = []
    seen_groups = set()
    if isinstance(raw_groups, list):
        for raw_group in raw_groups:
            raw_items = [raw_group] if isinstance(raw_group, str) else raw_group
            if not isinstance(raw_items, list):
                continue
            group = []
            group_seen = set()
            for raw in raw_items:
                text = str(raw or "").strip().strip('"').strip("'").strip()
                if not text or text in group_seen:
                    continue
                group.append(text)
                group_seen.add(text)
                if text not in seen:
                    keywords.append(text)
                    seen.add(text)
            key = tuple(group)
            if group and key not in seen_groups:
                groups.append(group)
                seen_groups.add(key)
    if not keywords:
        keywords = list(default_section["keywords"])
    if not groups:
        groups = [[item] for item in keywords]
    return {"keywords": keywords, "groups": groups}


def normalize_news_keywords(value) -> dict:
    if not isinstance(value, dict):
        raise ValueError("keywords must be an object")
    defaults = default_news_keywords()
    out = {}
    for key in ["bank", "other"]:
        out[key] = normalize_keyword_section(value.get(key, defaults[key]), defaults[key])
    return out


def load_news_keywords() -> dict:
    if NEWS_KEYWORDS_PATH.exists():
        try:
            return normalize_news_keywords(json.loads(NEWS_KEYWORDS_PATH.read_text(encoding="utf-8")))
        except Exception:
            pass
    keywords = default_news_keywords()
    save_news_keywords(keywords)
    return keywords


def save_news_keywords(keywords: dict) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    NEWS_KEYWORDS_PATH.write_text(
        json.dumps(normalize_news_keywords(keywords), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def delete_local_file(path_value: str) -> None:
    if not path_value:
        return
    path = Path(path_value)
    try:
        if path.exists() and path.is_file():
            path.unlink()
    except Exception:
        pass


def resolve_existing_run_dir(path_value: str) -> Path | None:
    if not path_value:
        return None
    try:
        path = Path(path_value).resolve()
        runs_root = collector.RUNS_DIR.resolve()
        if path.exists() and path.is_dir() and str(path).startswith(str(runs_root)):
            return path
    except Exception:
        return None
    return None


def resolve_run_dir(run_id: str) -> Path | None:
    if not re.fullmatch(r"\d{6}_\d{4}(?:_\d+)?", run_id or ""):
        return None
    try:
        runs_root = collector.RUNS_DIR.resolve()
        path = (runs_root / run_id).resolve()
        path.relative_to(runs_root)
        return path
    except Exception:
        return None


def same_path(left: str | Path, right: str | Path) -> bool:
    try:
        return Path(left).resolve() == Path(right).resolve()
    except Exception:
        return False


def delete_local_file_under(path_value: str, root: Path) -> None:
    if not path_value:
        return
    try:
        path = Path(path_value).resolve()
        root = root.resolve()
        if path.exists() and path.is_file() and str(path).startswith(str(root)):
            path.unlink()
    except Exception:
        pass


def resolve_item_local_path(item: collector.Item, run_dir: Path | None = None) -> Path | None:
    candidates: list[Path] = []
    if item.local_path:
        raw = Path(item.local_path)
        candidates.append(raw)
        if run_dir and not raw.is_absolute():
            candidates.append(run_dir / raw)
        if run_dir and raw.name:
            candidates.append(run_dir / raw.name)
            candidates.append(run_dir / "raw_pdfs" / raw.name)
    for candidate in candidates:
        try:
            path = candidate.resolve()
            if path.exists() and path.is_file():
                return path
        except Exception:
            continue
    if run_dir and item.local_path:
        name = Path(item.local_path).name
        if name:
            try:
                run_root = run_dir.resolve()
                for path in run_root.rglob(name):
                    if path.is_file():
                        path.resolve().relative_to(run_root)
                        return path
            except Exception:
                return None
    return None


def resolve_article_open_url(item: collector.Item) -> str:
    if item.original_url:
        return item.original_url
    return item.url or ""


def normalize_article_urls(items: list[collector.Item], progress=None) -> None:
    if progress:
        def decode_progress(event_type: str, payload: dict) -> None:
            if event_type == "decode_start":
                progress(
                    "generate_status",
                    {
                        "message": (
                            f"湲곗궗 ?먮Ц 留곹겕 蹂??以?({payload.get('index')}/{payload.get('total')}): "
                            f"{payload.get('title') or ''}"
                        )
                    },
                )
            elif event_type == "decode_done":
                progress(
                    "generate_status",
                    {"message": f"湲곗궗 ?먮Ц 留곹겕 蹂???꾨즺 ({payload.get('index')}/{payload.get('total')})"},
                )
        collector.populate_original_urls(items, progress=decode_progress)
    else:
        collector.populate_original_urls(items)


def move_selected_files(items: list[collector.Item], run_dir: Path, move: bool = True) -> None:
    raw_root = run_dir / "raw_pdfs"
    for item in items:
        if not item.local_path:
            continue
        src = Path(item.local_path)
        if not src.exists() or not src.is_file():
            continue
        source_dir = raw_root / collector.clean_filename(item.source_name, item.source_name).replace(".pdf", "")
        source_dir.mkdir(parents=True, exist_ok=True)
        dst = source_dir / src.name
        if dst.exists():
            dst = source_dir / f"{dst.stem}_{uuid.uuid4().hex[:6]}{dst.suffix}"
        if move:
            shutil.move(str(src), str(dst))
        else:
            shutil.copy2(str(src), str(dst))
        item.local_path = str(dst)


def item_from_row(row: dict) -> collector.Item:
    extra = {}
    try:
        extra = json.loads(row.get("extra_json") or "{}")
    except json.JSONDecodeError:
        extra = {}
    item = collector.Item(
        category=row.get("category", ""),
        source_name=row.get("source_name", ""),
        title=row.get("title", ""),
        url=row.get("url", ""),
        published_date=row.get("published_date", ""),
        file_type=row.get("file_type", "pdf"),
        download_url=row.get("download_url", ""),
        local_path=row.get("local_path", ""),
        notes=row.get("notes", ""),
        original_url=row.get("original_url", ""),
        extra=extra,
    )
    if item.file_type == "article" and (item.source_name or "").lower() in {"v.daum.net", "daum", "daum ?댁뒪"}:
        try:
            collector.normalize_article_source_name(collector.Http(), item)
        except Exception:
            pass
    return item


def protect_with_dpapi(secret: str) -> bytes:
    data = secret.encode("utf-8")
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptProtectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData)
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def unprotect_with_dpapi(data: bytes) -> str:
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptUnprotectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData).decode("utf-8")
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def save_api_key(api_key: str) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    try:
        KEY_PATH.write_bytes(b"dpapi:" + protect_with_dpapi(api_key))
    except Exception:
        KEY_PATH.write_bytes(b"b64:" + base64.b64encode(api_key.encode("utf-8")))


def load_api_key() -> str:
    if not KEY_PATH.exists():
        return ""
    data = KEY_PATH.read_bytes()
    if data.startswith(b"dpapi:"):
        return unprotect_with_dpapi(data[6:])
    if data.startswith(b"b64:"):
        return base64.b64decode(data[4:]).decode("utf-8")
    return ""


class DATA_BLOB(ctypes.Structure):
    _fields_ = [("cbData", ctypes.wintypes.DWORD), ("pbData", ctypes.POINTER(ctypes.c_char))]


def multipart_request(url: str, fields: dict, files: list[tuple[str, Path]], api_key: str) -> dict:
    boundary = "----codexboundary" + uuid.uuid4().hex
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"),
                str(value).encode("utf-8"),
                b"\r\n",
            ]
        )
    for name, path in files:
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"; filename="{path.name}"\r\n'.encode("utf-8"),
                b"Content-Type: application/pdf\r\n\r\n",
                path.read_bytes(),
                b"\r\n",
            ]
        )
    chunks.append(f"--{boundary}--\r\n".encode("utf-8"))
    body = b"".join(chunks)
    req = urllib.request.Request(
        url,
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
    )
    opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
    with opener.open(req, timeout=180) as resp:
        return json.loads(resp.read().decode("utf-8"))


def openai_json_request(path: str, payload: dict, api_key: str, timeout: int = 300) -> dict:
    req = urllib.request.Request(
        f"https://api.openai.com/v1{path}",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
    )
    try:
        opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
        with opener.open(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", "replace")
        raise RuntimeError(f"OpenAI API error {exc.code}: {detail}") from exc


def call_openai_report(api_key: str, prompt: str, items: list[collector.Item], progress=None) -> dict:
    def emit(event_type: str, payload: dict | None = None) -> None:
        if progress:
            progress(event_type, payload or {})

    emit("generate_status", {"message": "LLM ?낅젰 JSON 怨④꺽??留뚮뱶??以묒엯?덈떎."})
    report_base = build_report_base(items)
    report_items = list(iter_report_items(report_base))
    merged_bullets = {"ITEMS": []}
    response_logs = []
    response_ids = []
    batch_size = 6
    total_batches = max((len(report_items) + batch_size - 1) // batch_size, 1)

    for batch_index, start in enumerate(range(0, len(report_items), batch_size), start=1):
        batch = report_items[start:start + batch_size]
        expected_ids = [item.get("ITEM_ID", "") for item in batch]
        pending = batch
        collected: dict[str, dict] = {}
        for attempt in range(1, 3):
            if not pending:
                break
            emit(
                "generate_status",
                {
                    "message": (
                        f"OpenAI??bullet ?앹꽦???붿껌?섎뒗 以묒엯?덈떎. "
                        f"({batch_index}/{total_batches}, {len(pending)}嫄?"
                    )
                },
            )
            response, text, parsed = request_openai_bullets(api_key, prompt, pending, emit)
            response_ids.append(response.get("id", ""))
            response_logs.append({
                "batch": batch_index,
                "attempt": attempt,
                "expected_item_ids": [item.get("ITEM_ID", "") for item in pending],
                "response_id": response.get("id", ""),
                "output_text": text,
            })
            for raw in extract_bullet_items(parsed):
                if not isinstance(raw, dict):
                    continue
                item_id = clean_cell_value(raw.get("ITEM_ID"))
                if item_id in expected_ids:
                    collected[item_id] = raw
            missing = [item_id for item_id in expected_ids if item_id not in collected]
            if not missing:
                break
            pending = [item for item in batch if item.get("ITEM_ID", "") in missing]
            emit("generate_status", {"message": f"?꾨씫 bullet {len(missing)}嫄댁쓣 ?ъ슂泥?븯??以묒엯?덈떎."})
        merged_bullets["ITEMS"].extend(collected[item_id] for item_id in expected_ids if item_id in collected)

    emit("generate_status", {"message": "LLM ?묐떟??JSON?쇰줈 ?뺣━?섎뒗 以묒엯?덈떎."})
    merged = merge_bullets_into_report(report_base, merged_bullets)
    merged["_openai_response_id"] = ", ".join(x for x in response_ids if x)
    merged["_openai_output_text"] = json.dumps(
        {
            "batch_size": batch_size,
            "expected_count": len(report_items),
            "received_count": len(merged_bullets["ITEMS"]),
            "responses": response_logs,
        },
        ensure_ascii=False,
        indent=2,
    )
    return merged


def request_openai_bullets(api_key: str, prompt: str, report_items: list[dict], emit) -> tuple[dict, str, dict]:
    materials = build_bullet_request_context(report_items)
    if "{{INPUT_MATERIALS}}" in prompt:
        prompt_text = prompt.replace("{{INPUT_MATERIALS}}", materials)
    else:
        prompt_text = prompt + "\n\n[?낅젰 ?먮즺]\n" + materials
    content = [{"type": "input_text", "text": prompt_text}]
    pdf_items = [item for item in report_items if item.get("_LOCAL_PATH")]
    for index, item in enumerate(pdf_items, start=1):
        path = Path(item.get("_LOCAL_PATH", ""))
        if not path or not path.exists() or path.suffix.lower() != ".pdf":
            continue
        emit("generate_upload", {"index": index, "total": len(pdf_items), "title": item.get("TITLE", "")})
        uploaded = multipart_request(
            "https://api.openai.com/v1/files",
            {"purpose": "user_data"},
            [("file", path)],
            api_key,
        )
        content.append({"type": "input_file", "file_id": uploaded["id"]})
    response = openai_json_request(
        "/responses",
        {
            "model": OPENAI_MODEL,
            "input": [{"role": "user", "content": content}],
            "text": {"format": {"type": "json_object"}},
        },
        api_key,
    )
    text = extract_response_text(response)
    parsed = extract_json_object(text)
    return response, text, parsed


def build_bullet_request_context(report_base_or_items) -> str:
    rows = []
    source_items = report_base_or_items if isinstance(report_base_or_items, list) else iter_report_items(report_base_or_items)
    for item in source_items:
        rows.append({
            "ITEM_ID": item.get("ITEM_ID", ""),
            "SECTION": item.get("_SECTION", ""),
            "NO": item.get("NO", ""),
            "SOURCE_NAME": item.get("SOURCE_NAME", ""),
            "TITLE": item.get("TITLE", ""),
            "URL": item.get("URL", ""),
            "PUBLISHED_MM_DD": item.get("PUBLISHED_MM_DD", ""),
            "LOCAL_PATH": item.get("_LOCAL_PATH", ""),
            "CONTENT_TEXT": item.get("_CONTENT_TEXT", ""),
        })
    return json.dumps(rows, ensure_ascii=False, indent=2)


def build_report_base(items: list[collector.Item]) -> dict:
    bank_items = [item for item in items if item.category == "은행/지주사"]
    other_items = [item for item in items if item.category == "그외"]
    agency_items = [item for item in items if item.category == "국가기관"]
    research_items = [item for item in items if item.category == "금융연구소"]
    return {
        "BANK_SECTION": {
            "ITEMS": metadata_items(bank_items, "bank", "1", "BANK_SECTION.ITEMS"),
            "OTHER_ITEMS": metadata_items(other_items, "other", "", "BANK_SECTION.OTHER_ITEMS"),
        },
        "GOVERNMENT_SECTION": {
            "ITEMS": metadata_items(agency_items, "agency", "2", "GOVERNMENT_SECTION.ITEMS"),
        },
        "FINANCIAL_RESEARCH_SECTION": {
            "ITEMS": metadata_items(research_items, "research", "3", "FINANCIAL_RESEARCH_SECTION.ITEMS"),
        },
    }


def metadata_items(items: list[collector.Item], id_prefix: str, section_no: str, section_name: str) -> list[dict]:
    out = []
    for idx, item in enumerate(items, start=1):
        url = item.original_url or item.url
        out.append({
            "ITEM_ID": f"{id_prefix}_{idx}",
            "_SECTION": section_name,
            "_LOCAL_PATH": item.local_path,
            "NO": f"{section_no}.{idx}" if section_no else "",
            "SOURCE_NAME": clean_cell_value(item.source_name),
            "TITLE": clean_title(item.title, item.source_name),
            "URL": clean_cell_value(url),
            "PUBLISHED_MM_DD": clean_mm_dd(item.published_mm_dd),
            "_CONTENT_TEXT": content_text_for_item(item),
            "SUMMARY_BULLET_1": "",
            "SUMMARY_BULLET_2": "",
            "SUMMARY_BULLET_3": "",
        })
    return out


def content_text_for_item(item: collector.Item, limit: int = 5000) -> str:
    text = ""
    if isinstance(item.extra, dict):
        text = item.extra.get("article_text") or item.extra.get("rss_description") or ""
    text = clean_cell_value(text)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + " ..."


def merge_bullets_into_report(report_base: dict, bullet_json: dict) -> dict:
    by_id = {item.get("ITEM_ID", ""): item for item in iter_report_items(report_base)}
    by_url = {item.get("URL", ""): item for item in iter_report_items(report_base) if item.get("URL")}
    by_title = {item.get("TITLE", ""): item for item in iter_report_items(report_base) if item.get("TITLE")}
    for raw in extract_bullet_items(bullet_json):
        if not isinstance(raw, dict):
            continue
        item_id = clean_cell_value(raw.get("ITEM_ID"))
        target = by_id.get(item_id)
        if not target:
            target = by_url.get(clean_cell_value(raw.get("URL")))
        if not target:
            target = by_title.get(clean_title(raw.get("TITLE"), raw.get("SOURCE_NAME")))
        if not target:
            continue
        for key in ["SUMMARY_BULLET_1", "SUMMARY_BULLET_2", "SUMMARY_BULLET_3"]:
            target[key] = clean_cell_value(raw.get(key))
    return strip_internal_report_fields(report_base)


def extract_bullet_items(data: dict) -> list:
    if isinstance(data.get("ITEMS"), list):
        return data["ITEMS"]
    if isinstance(data.get("BULLETS"), list):
        return data["BULLETS"]
    out = []
    for section in ["BANK_SECTION", "GOVERNMENT_SECTION", "FINANCIAL_RESEARCH_SECTION"]:
        section_data = data.get(section, {})
        if not isinstance(section_data, dict):
            continue
        if isinstance(section_data.get("ITEMS"), list):
            out.extend(section_data["ITEMS"])
        if isinstance(section_data.get("OTHER_ITEMS"), list):
            out.extend(section_data["OTHER_ITEMS"])
    return out


def iter_report_items(report_data: dict):
    yield from report_data.get("BANK_SECTION", {}).get("ITEMS", [])
    yield from report_data.get("BANK_SECTION", {}).get("OTHER_ITEMS", [])
    yield from report_data.get("GOVERNMENT_SECTION", {}).get("ITEMS", [])
    yield from report_data.get("FINANCIAL_RESEARCH_SECTION", {}).get("ITEMS", [])


def strip_internal_report_fields(report_data: dict) -> dict:
    cleaned = deepcopy(report_data)
    for item in iter_report_items(cleaned):
        for key in ["ITEM_ID", "_SECTION", "_LOCAL_PATH", "_CONTENT_TEXT"]:
            item.pop(key, None)
    return cleaned


def extract_response_text(response: dict) -> str:
    if response.get("output_text"):
        return str(response["output_text"])
    parts = []
    for output in response.get("output", []):
        for content in output.get("content", []):
            if "text" in content:
                parts.append(str(content["text"]))
    return "\n".join(parts)


def extract_json_object(text: str) -> dict:
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            return json.loads(match.group(0))
        raise


def generate_docx(report_json: dict, items: list[collector.Item], output_path: Path) -> None:
    template = TEMPLATE_PATH
    doc = Document(template) if template.exists() else Document()
    add_section_bookmarks(doc)
    request_word_field_update(doc)
    report_data = normalize_report_json(report_json, items)
    populate_report_tables(doc, report_data)
    replace_docx_placeholders(
        doc,
        {
            "BANK_SECTION_PAGE": "__PAGEREF_BANK_SECTION__",
            "AGENCY_SECTION_PAGE": "__PAGEREF_AGENCY_SECTION__",
            "RESEARCH_SECTION_PAGE": "__PAGEREF_RESEARCH_SECTION__",
        },
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(f".{output_path.stem}.{uuid.uuid4().hex[:8]}.tmp{output_path.suffix}")
    try:
        doc.save(tmp_path)
        replace_file_with_retry(tmp_path, output_path)
    finally:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


def replace_file_with_retry(src: Path, dst: Path, attempts: int = 5) -> None:
    last_exc: Exception | None = None
    for attempt in range(attempts):
        try:
            os.replace(src, dst)
            return
        except PermissionError as exc:
            last_exc = exc
            time.sleep(0.5 * (attempt + 1))
    raise PermissionError(
        f"{dst} ?뚯씪????뼱?????놁뒿?덈떎. 蹂닿퀬?쒓? ?대젮 ?덈떎硫??レ? ???ㅼ떆 ?ъ깮?깊븯?몄슂."
    ) from last_exc


def build_template_values(report_json: dict, items: list[collector.Item]) -> dict:
    values = {k: stringify_template_value(v) for k, v in report_json.items()}
    values.setdefault("REPORT_WEEK", dt.date.today().strftime("%Y-%m-%d"))
    values.setdefault("DEPARTMENT_NAME", "議곗궗?곌뎄")
    groups = {
        "BANK_DETAIL_ITEMS": [x for x in items if x.category in {"은행/지주사", "그외"}],
        "AGENCY_DETAIL_ITEMS": [x for x in items if x.category == "국가기관"],
        "RESEARCH_DETAIL_ITEMS": [x for x in items if x.category == "금융연구소"],
    }
    for key, group_items in groups.items():
        values.setdefault(key, format_detail_items(report_json.get(key), group_items))
    values["BANK_SECTION_PAGE"] = "__PAGEREF_BANK_SECTION__"
    values["AGENCY_SECTION_PAGE"] = "__PAGEREF_AGENCY_SECTION__"
    values["RESEARCH_SECTION_PAGE"] = "__PAGEREF_RESEARCH_SECTION__"
    values.setdefault("ITEM_NO", "")
    values.setdefault("TITLE", "")
    values.setdefault("URL", "")
    values.setdefault("SOURCE_NAME", "")
    values.setdefault("PUBLISHED_MM_DD", "")
    values.setdefault("SUMMARY_BULLET_1", "")
    values.setdefault("SUMMARY_BULLET_2", "")
    values.setdefault("SUMMARY_BULLET_3", "")
    values.setdefault("NO", "")
    return values


def normalize_report_json(report_json: dict, items: list[collector.Item]) -> dict:
    data = {
        "BANK_SECTION": {
            "ITEMS": normalize_items(report_json.get("BANK_SECTION", {}).get("ITEMS", []), "1"),
            "OTHER_ITEMS": normalize_items(report_json.get("BANK_SECTION", {}).get("OTHER_ITEMS", []), ""),
        },
        "GOVERNMENT_SECTION": {
            "ITEMS": normalize_items(report_json.get("GOVERNMENT_SECTION", {}).get("ITEMS", []), "2"),
        },
        "FINANCIAL_RESEARCH_SECTION": {
            "ITEMS": normalize_items(report_json.get("FINANCIAL_RESEARCH_SECTION", {}).get("ITEMS", []), "3"),
        },
    }
    if not any(data[section]["ITEMS"] for section in ["BANK_SECTION", "GOVERNMENT_SECTION", "FINANCIAL_RESEARCH_SECTION"]):
        data = fallback_report_data(items)
    return data


def normalize_items(raw_items, section_no: str) -> list[dict]:
    if not isinstance(raw_items, list):
        return []
    out = []
    for idx, raw in enumerate(raw_items, start=1):
        if not isinstance(raw, dict):
            continue
        item = {
            "NO": clean_cell_value(raw.get("NO") or (f"{section_no}.{idx}" if section_no else "")),
            "SOURCE_NAME": clean_cell_value(raw.get("SOURCE_NAME")),
            "TITLE": clean_title(raw.get("TITLE"), raw.get("SOURCE_NAME")),
            "URL": clean_cell_value(raw.get("URL")),
            "PUBLISHED_MM_DD": clean_mm_dd(raw.get("PUBLISHED_MM_DD")),
            "SUMMARY_BULLET_1": clean_cell_value(raw.get("SUMMARY_BULLET_1")),
            "SUMMARY_BULLET_2": clean_cell_value(raw.get("SUMMARY_BULLET_2")),
            "SUMMARY_BULLET_3": clean_cell_value(raw.get("SUMMARY_BULLET_3")),
        }
        out.append(item)
    return out


def fallback_report_data(items: list[collector.Item]) -> dict:
    groups = {
        "BANK_SECTION": [x for x in items if x.category in {"은행/지주사", "그외"}],
        "GOVERNMENT_SECTION": [x for x in items if x.category == "국가기관"],
        "FINANCIAL_RESEARCH_SECTION": [x for x in items if x.category == "금융연구소"],
    }
    return {
        "BANK_SECTION": {
            "ITEMS": fallback_items(groups["BANK_SECTION"], "1"),
            "OTHER_ITEMS": [],
        },
        "GOVERNMENT_SECTION": {"ITEMS": fallback_items(groups["GOVERNMENT_SECTION"], "2")},
        "FINANCIAL_RESEARCH_SECTION": {"ITEMS": fallback_items(groups["FINANCIAL_RESEARCH_SECTION"], "3")},
    }


def fallback_items(items: list[collector.Item], section_no: str) -> list[dict]:
    out = []
    for idx, item in enumerate(items, start=1):
        url = item.original_url or item.url
        out.append(
            {
                "NO": f"{section_no}.{idx}",
                "SOURCE_NAME": item.source_name,
                "TITLE": clean_title(item.title, item.source_name),
                "URL": url,
                "PUBLISHED_MM_DD": item.published_mm_dd,
                "SUMMARY_BULLET_1": "",
                "SUMMARY_BULLET_2": "",
                "SUMMARY_BULLET_3": "",
            }
        )
    return out


def clean_cell_value(value) -> str:
    return "" if value is None else re.sub(r"\s+", " ", str(value)).strip()


def clean_mm_dd(value) -> str:
    text = clean_cell_value(value)
    match = re.fullmatch(r"0?(\d{1,2})[.\-/]0?(\d{1,2})", text)
    if not match:
        return ""
    month = int(match.group(1))
    day = int(match.group(2))
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return ""
    return f"{month}.{day}"


def clean_title(value, source_name: str = "") -> str:
    text = clean_cell_value(value)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s+[媛-??{2,4}(?:\s*,\s*[媛-??{2,4}){0,5}\s+20\d{2}[.\-/]\d{1,2}[.\-/]\d{1,2}\s+議고쉶??s*\d+\s*$", "", text)
    text = re.sub(r"\s+20\d{2}[.\-/]\d{1,2}[.\-/]\d{1,2}\s+議고쉶??s*\d+\s*$", "", text)
    text = re.sub(r"\s+(?묒꽦???깅줉???묒꽦??議고쉶??\s*[:竊??\s*.*$", "", text)
    source_name = clean_cell_value(source_name)
    if source_name:
        suffix = f" - {source_name}"
        while text.endswith(suffix):
            text = text[: -len(suffix)].strip()
    return text.strip()


def row_for_client(item: collector.Item) -> dict:
    row = item.row()
    row["title"] = clean_title(row.get("title", ""), row.get("source_name", ""))
    row["published_date"] = collector.normalize_date(row.get("published_date", "")) or row.get("published_date", "")
    return row


def populate_report_tables(doc: Document, data: dict) -> None:
    if len(doc.tables) < 10:
        return
    fill_summary_table(doc.tables[2], data["BANK_SECTION"]["ITEMS"], keep_min=1)
    fill_other_summary_table(doc.tables[3], data["BANK_SECTION"]["OTHER_ITEMS"])
    fill_summary_table(doc.tables[4], data["GOVERNMENT_SECTION"]["ITEMS"], keep_min=1)
    fill_summary_table(doc.tables[5], data["FINANCIAL_RESEARCH_SECTION"]["ITEMS"], keep_min=1)
    fill_detail_table(doc.tables[6], data["BANK_SECTION"]["ITEMS"], keep_min=1)
    fill_other_detail_table(doc.tables[7], data["BANK_SECTION"]["OTHER_ITEMS"])
    fill_detail_table(doc.tables[8], data["GOVERNMENT_SECTION"]["ITEMS"], keep_min=1)
    fill_detail_table(doc.tables[9], data["FINANCIAL_RESEARCH_SECTION"]["ITEMS"], keep_min=1)


def fill_summary_table(table, items: list[dict], keep_min: int = 1) -> None:
    desired = max(len(items), keep_min)
    adjust_table_rows(table, header_rows=2, desired_data_rows=desired)
    set_fixed_table_widths(table, [650, 1600, 7550])
    rows = items + [blank_item()] * (desired - len(items))
    for row, item in zip(table.rows[2:], rows):
        set_cell_text(row.cells[0], item.get("NO", ""))
        set_cell_text(row.cells[1], item.get("SOURCE_NAME", ""))
        set_cell_text(row.cells[2], summary_title(item))


def fill_other_summary_table(table, items: list[dict]) -> None:
    adjust_table_rows(table, header_rows=0, desired_data_rows=1)
    set_fixed_table_widths(table, [650, 1600, 7550])
    row = table.rows[0]
    set_cell_text(row.cells[0], "그외")
    set_cell_text(row.cells[1], "\n".join(item.get("SOURCE_NAME", "") for item in items))
    set_cell_text(row.cells[2], "\n".join(summary_title(item) for item in items))


def fill_detail_table(table, items: list[dict], keep_min: int = 1) -> None:
    desired = max(len(items), keep_min)
    adjust_detail_rows(table, desired)
    set_fixed_table_widths(table, [650, 9150])
    rows = items + [blank_item()] * (desired - len(items))
    for idx, item in enumerate(rows):
        title_row = table.rows[1 + idx * 2]
        bullet_row = table.rows[2 + idx * 2]
        no = item.get("NO", "")
        set_cell_text(title_row.cells[0], no)
        set_cell_text(title_row.cells[1], detail_title(item))
        set_cell_text(bullet_row.cells[0], "")
        set_cell_text(bullet_row.cells[1], detail_summary(item))


def fill_other_detail_table(table, items: list[dict]) -> None:
    desired = max(len(items), 1)
    adjust_other_detail_rows(table, desired)
    set_fixed_table_widths(table, [650, 9150])
    rows = items + [blank_item()] * (desired - len(items))
    for idx, item in enumerate(rows):
        title_row = table.rows[idx * 2]
        bullet_row = table.rows[idx * 2 + 1]
        set_cell_text(title_row.cells[0], "그외" if idx == 0 else "")
        set_cell_text(title_row.cells[1], detail_title(item))
        set_cell_text(bullet_row.cells[0], "")
        set_cell_text(bullet_row.cells[1], detail_summary(item))


def adjust_table_rows(table, header_rows: int, desired_data_rows: int) -> None:
    while len(table.rows) < header_rows + desired_data_rows:
        clone_row(table, table.rows[-1])
    while len(table.rows) > header_rows + desired_data_rows:
        delete_row(table, table.rows[-1])


def adjust_detail_rows(table, desired_items: int) -> None:
    if len(table.rows) < 3:
        return
    title_row_xml = deepcopy(table.rows[1]._tr)
    bullet_row_xml = deepcopy(table.rows[2]._tr)
    remove_vertical_merges(title_row_xml)
    remove_vertical_merges(bullet_row_xml)
    while len(table.rows) > 1:
        delete_row(table, table.rows[-1])
    for _ in range(desired_items):
        table._tbl.append(deepcopy(title_row_xml))
        table._tbl.append(deepcopy(bullet_row_xml))


def adjust_other_detail_rows(table, desired_items: int) -> None:
    if len(table.rows) < 2:
        return
    title_row_xml = deepcopy(table.rows[0]._tr)
    bullet_row_xml = deepcopy(table.rows[1]._tr)
    remove_vertical_merges(title_row_xml)
    remove_vertical_merges(bullet_row_xml)
    while len(table.rows) > 0:
        delete_row(table, table.rows[-1])
    for _ in range(desired_items):
        table._tbl.append(deepcopy(title_row_xml))
        table._tbl.append(deepcopy(bullet_row_xml))


def clone_row(table, row) -> None:
    table._tbl.append(deepcopy(row._tr))


def delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def remove_vertical_merges(row_xml) -> None:
    for tc_pr in row_xml.iter(qn("w:tcPr")):
        for v_merge in list(tc_pr.findall(qn("w:vMerge"))):
            tc_pr.remove(v_merge)


def set_fixed_table_widths(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        seen = set()
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths) or id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


URL_RE = re.compile(r"https?://[^\s)]+")


def add_hyperlink(paragraph, url: str, text: str | None = None) -> None:
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_hyperlinks(paragraph, line: str) -> None:
    pos = 0
    for match in URL_RE.finditer(line):
        if match.start() > pos:
            paragraph.add_run(line[pos:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        pos = match.end()
    if pos < len(line):
        paragraph.add_run(line[pos:])


def set_cell_text(cell, text: str) -> None:
    text = text or ""
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    remove_paragraph_numbering(first)
    for child in list(first._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            first._p.remove(child)
    lines = str(text).split("\n")
    for index, line in enumerate(lines):
        if index:
            first.add_run().add_break()
        add_text_with_hyperlinks(first, line)


def remove_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def blank_item() -> dict:
    return {
        "NO": "",
        "SOURCE_NAME": "",
        "TITLE": "",
        "URL": "",
        "PUBLISHED_MM_DD": "",
        "SUMMARY_BULLET_1": "",
        "SUMMARY_BULLET_2": "",
        "SUMMARY_BULLET_3": "",
    }


def summary_title(item: dict) -> str:
    title = item.get("TITLE", "")
    if not title:
        return ""
    mmdd = item.get("PUBLISHED_MM_DD", "")
    return f"??{title}" + (f" ({mmdd})" if mmdd else "")


def detail_title(item: dict) -> str:
    title = item.get("TITLE", "")
    if not title:
        return ""
    url = item.get("URL", "")
    meta = " / ".join(x for x in [item.get("SOURCE_NAME", ""), item.get("PUBLISHED_MM_DD", "")] if x)
    second = f"{url} ({meta})" if url and meta else url or (f"({meta})" if meta else "")
    return f"??{title}" + (f"\n{second}" if second else "")


def detail_summary(item: dict) -> str:
    bullets = [
        item.get("SUMMARY_BULLET_1", ""),
        item.get("SUMMARY_BULLET_2", ""),
        item.get("SUMMARY_BULLET_3", ""),
    ]
    return "\n".join(format_detail_bullet(b) for b in bullets if clean_cell_value(b))


def format_detail_bullet(value: str) -> str:
    text = clean_cell_value(value)
    text = re.sub(r"^(?:[\-?◈룐볛?\s*)+", "", text).strip()
    return f"- {text}" if text else ""


def detail_block(item: dict) -> str:
    parts = [detail_title(item), detail_summary(item)]
    return "\n".join(part for part in parts if part)


def compact_url(url: str, limit: int = 76) -> str:
    url = clean_cell_value(url)
    if not url:
        return ""
    parsed = urllib.parse.urlparse(url)
    if parsed.netloc:
        value = parsed.netloc + parsed.path
        if parsed.query:
            value += "?..."
    else:
        value = url
    if len(value) <= limit:
        return value
    return value[: max(0, limit - 3)].rstrip("/") + "..."


def stringify_template_value(value) -> str:
    if isinstance(value, (list, dict)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return "" if value is None else str(value)


def format_detail_items(generated, fallback_items: list[collector.Item]) -> str:
    source = generated if isinstance(generated, list) else []
    lines = []
    if source:
        for idx, item in enumerate(source, start=1):
            if not isinstance(item, dict):
                lines.append(str(item))
                continue
            lines.extend(
                [
                    f"{idx}. {item.get('TITLE') or item.get('title') or ''}",
                    f"{item.get('URL') or item.get('url') or ''} ({item.get('SOURCE_NAME') or item.get('source_name') or ''} / {item.get('PUBLISHED_MM_DD') or ''})",
                    f"- {item.get('SUMMARY_BULLET_1') or ''}",
                    f"- {item.get('SUMMARY_BULLET_2') or ''}",
                    f"- {item.get('SUMMARY_BULLET_3') or ''}",
                    "",
                ]
            )
    else:
        for idx, item in enumerate(fallback_items, start=1):
            lines.extend([
                f"{idx}. {item.title}",
                f"{item.original_url or item.url} ({item.source_name} / {item.published_mm_dd})",
                "",
            ])
    return "\n".join(lines).strip()


def replace_docx_placeholders(doc: Document, values: dict) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, values)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, values)


def replace_in_paragraph(paragraph, values: dict) -> None:
    text = paragraph.text
    if "{{" not in text:
        return
    for key, value in values.items():
        text = text.replace("{{" + key + "}}", str(value))
    paragraph.clear()
    add_text_with_pageref_fields(paragraph, text)


def add_text_with_pageref_fields(paragraph, text: str) -> None:
    markers = {
        "__PAGEREF_BANK_SECTION__": "bank_section",
        "__PAGEREF_AGENCY_SECTION__": "agency_section",
        "__PAGEREF_RESEARCH_SECTION__": "research_section",
    }
    pattern = "(" + "|".join(re.escape(marker) for marker in markers) + ")"
    for part in re.split(pattern, text):
        if not part:
            continue
        bookmark = markers.get(part)
        if bookmark:
            add_pageref_field(paragraph, bookmark)
        else:
            paragraph.add_run(part)


def add_pageref_field(paragraph, bookmark_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    placeholder = OxmlElement("w:t")
    placeholder.text = "0"
    run._r.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_section_bookmarks(doc: Document) -> None:
    targets = {
        "1. 은행/지주사": "bank_section",
        "2. 국가기관": "agency_section",
        "3. 금융연구소": "research_section",
    }
    candidates: dict[str, list] = {bookmark: [] for bookmark in targets.values()}
    for paragraph in iter_doc_paragraphs(doc):
        text = paragraph.text.strip()
        for label, bookmark in targets.items():
            if text == label:
                candidates[bookmark].append(paragraph)
    bookmark_id = 100
    for bookmark, paragraphs in candidates.items():
        if not paragraphs:
            continue
        add_bookmark_to_paragraph(paragraphs[-1], bookmark, bookmark_id)
        bookmark_id += 1


def iter_doc_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def add_bookmark_to_paragraph(paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._p
    if not paragraph.runs:
        paragraph.add_run("")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def request_word_field_update(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def run_collect_job(job_id: str, data: dict) -> None:
    def emit(event_type: str, payload: dict | None = None) -> None:
        payload = dict(payload or {})
        payload["type"] = event_type
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if job:
            job["queue"].put(payload)

    try:
        news_start = parse_iso_date(data["news_start"])
        news_end = parse_iso_date(data["news_end"])
        agency_start = parse_iso_date(data["agency_start"])
        agency_end = parse_iso_date(data["agency_end"])
        research_start = parse_iso_date(data["research_start"])
        research_end = parse_iso_date(data["research_end"])
        news_bank_max = int(data.get("news_bank_max") or 10)
        news_other_max = int(data.get("news_other_max") or 10)
        agency_max = int(data.get("agency_max") or 10)
        research_max = int(data.get("research_max") or 10)
        news_keywords = load_news_keywords()
        include_news = bool(data.get("include_news", True))
        include_agency = bool(data.get("include_agency", True))
        include_research = bool(data.get("include_research", True))
        run_dir = collector.make_run_dir()
        emit("started", {"run_dir": str(run_dir)})

        items = collector.collect_by_ranges(
            news_start,
            news_end,
            agency_start,
            agency_end,
            research_start,
            research_end,
            max_per_source=None,
            dry_run=False,
            news_bank_max=news_bank_max,
            news_other_max=news_other_max,
            agency_max=agency_max,
            research_max=research_max,
            output_dir=run_dir,
            progress=emit,
            news_keywords=news_keywords,
            include_news=include_news,
            include_agency=include_agency,
            include_research=include_research,
        )
        collector.write_outputs(items, output_dir=run_dir)
        payload = []
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = ""
            for idx, item in enumerate(items):
                row = row_for_client(item)
                row["id"] = idx
                payload.append(row)
        emit("complete", {"items": payload, "run_dir": str(run_dir)})
    except Exception as exc:
        traceback.print_exc()
        emit("fatal", {"error": str(exc)})


def write_items(items: list[collector.Item], csv_path: Path, xlsx_path: Path) -> None:
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = list(collector.Item("", "", "", "").row().keys())
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for item in items:
            writer.writerow(row_for_client(item))

    wb = Workbook()
    ws = wb.active
    ws.title = "selected_metadata"
    ws.append(fieldnames)
    for item in items:
        row = row_for_client(item)
        ws.append([row[k] for k in fieldnames])
    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col[:100])
        ws.column_dimensions[col[0].column_letter].width = min(max(max_len + 2, 12), 45)
    wb.save(xlsx_path)


def main() -> int:
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    print(f"議곗궗?곌뎄 ?꾩슦誘? http://{HOST}:{PORT}")
    if os.environ.get("RRA_NO_BROWSER") != "1":
        threading.Timer(0.8, lambda: webbrowser.open(f"http://{HOST}:{PORT}/")).start()
    server.serve_forever()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

