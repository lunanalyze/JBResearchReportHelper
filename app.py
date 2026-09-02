from __future__ import annotations

import datetime as dt
import base64
import ctypes
import ctypes.wintypes
import json
import os
import re
import shutil
import sys
import threading
import time
import traceback
import urllib.parse
import urllib.request
import webbrowser
import csv
import queue
import uuid
from copy import deepcopy
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

sys.dont_write_bytecode = True

import collector
import paths
import updater
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from openpyxl import Workbook


HOST = "127.0.0.1"
PORT = 8765
SERVER_IDLE_TIMEOUT_SECONDS = 15 * 60
SERVER_IDLE_CHECK_SECONDS = 15
APP_DIR = paths.APP_DIR
APP_DATA_DIR = paths.APP_DATA_DIR
CONFIG_DIR = paths.CONFIG_DIR
PROMPT_PATH = paths.RESOURCES_DIR / "report_prompt.md"
TEMPLATE_PATH = paths.RESOURCES_DIR / "report_template.docx"
KEY_PATH = CONFIG_DIR / "openai_key.bin"
NEWS_KEYWORDS_PATH = CONFIG_DIR / "news_keywords.json"
OPENAI_MODEL = "gpt-5.6-luna"
DEFAULT_REPORT_PROMPT = """Return only one JSON object for filling the weekly report template.
Use the provided metadata and attached PDFs.

Required top-level keys:
- REPORT_WEEK
- DEPARTMENT_NAME
- BANK_SECTION_PAGE
- AGENCY_SECTION_PAGE
- RESEARCH_SECTION_PAGE
- BANK_DETAIL_ITEMS
- AGENCY_DETAIL_ITEMS
- RESEARCH_DETAIL_ITEMS

Each DETAIL_ITEMS value should be an array of objects with ITEM_NO, TITLE, URL, SOURCE_NAME, PUBLISHED_MM_DD, SUMMARY_BULLET_1, SUMMARY_BULLET_2, SUMMARY_BULLET_3.
"""

paths.ensure_app_dirs()
paths.migrate_legacy_data()
paths.copy_default_resource("report_prompt.md", DEFAULT_REPORT_PROMPT, overwrite=True)
paths.copy_default_resource("report_template.docx", overwrite=True)

STATE_LOCK = threading.Lock()
STATE: dict = {
    "items": [],
    "run_dir": None,
    "report_path": "",
    "pending_run_dir": None,
    "cancelled_run_dirs": set(),
    "active_collect_job_id": "",
    "cancelled_collect_jobs": set(),
    "active_generate_job_id": "",
    "cancelled_generate_jobs": set(),
    "last_activity": time.time(),
}
JOBS: dict[str, dict] = {}


INDEX_HTML = r"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>조사연구 도우미</title>
  <style>
    :root {
      color-scheme: light;
      --bg: #f6f7f9;
      --panel: #ffffff;
      --text: #17202a;
      --muted: #6b7280;
      --line: #d8dee8;
      --brand: #1f6feb;
      --brand-dark: #1557bd;
      --ok: #0f7b3d;
      --warn: #b45309;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: "Malgun Gothic", "Segoe UI", sans-serif;
      background: var(--bg);
      color: var(--text);
    }
    header {
      padding: 18px 24px 14px;
      background: var(--panel);
      border-bottom: 1px solid var(--line);
    }
    .brand {
      display: flex;
      align-items: center;
      gap: 12px;
    }
    .brand-mark {
      width: 38px;
      height: 38px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      border-radius: 8px;
      background: #eef4ff;
      color: #1557bd;
      border: 1px solid #c7ddff;
      font-size: 16px;
      font-weight: 900;
    }
    .brand-copy {
      min-width: 0;
    }
    h1 {
      margin: 0;
      font-size: 21px;
      font-weight: 700;
      letter-spacing: 0;
    }
    /* 업데이트 알림 — 켜자마자 모달로 막으면 "업무 보러 켰는데 방해받는다"가 된다.
       배너는 눈에 들어오지만 무시할 수 있고, 닫아도 다음 버전에는 다시 뜬다. */
    .update-bar {
      display: flex;
      align-items: center;
      gap: 10px;
      padding: 10px 24px;
      background: #eef4ff;
      border-bottom: 1px solid #c7ddff;
      font-size: 13px;
    }
    .update-bar[hidden] { display: none; }
    .update-mark {
      width: 20px;
      height: 20px;
      flex: 0 0 auto;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      border-radius: 999px;
      background: var(--brand);
      color: #fff;
      font-size: 12px;
      font-weight: 900;
    }
    .update-copy {
      flex: 1 1 auto;
      min-width: 0;
      overflow: hidden;
      text-overflow: ellipsis;
      white-space: nowrap;
    }
    .update-copy strong { font-weight: 800; }
    .update-muted { color: var(--muted); }
    .update-dismiss {
      width: 28px;
      height: 28px;
      padding: 0;
      background: transparent;
      border-color: transparent;
      color: var(--muted);
      font-size: 16px;
      line-height: 1;
    }
    .update-dismiss:hover { background: #dbe7ff; }
    .update-modal {
      width: min(520px, 100%);
      grid-template-rows: auto auto minmax(0, 1fr);
    }
    .update-body {
      padding: 0 18px 18px;
      overflow: auto;
    }
    .update-notes {
      white-space: pre-wrap;
      max-height: 190px;
      overflow: auto;
      background: var(--bg);
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 10px 12px;
      font-size: 13px;
    }
    .update-notes[hidden] { display: none; }
    .update-guide {
      margin: 12px 0 0;
      color: var(--muted);
      font-size: 13px;
      line-height: 1.65;
    }
    .update-guide[hidden] { display: none; }
    .update-progress {
      display: flex;
      align-items: center;
      gap: 10px;
      margin-top: 12px;
      font-size: 13px;
    }
    .update-progress[hidden] { display: none; }
    .update-spinner {
      width: 16px;
      height: 16px;
      flex: 0 0 auto;
      border-radius: 999px;
      border: 2px solid #c7ddff;
      border-top-color: var(--brand);
      animation: update-spin 0.8s linear infinite;
    }
    @keyframes update-spin { to { transform: rotate(360deg); } }
    .update-error {
      margin-top: 12px;
      color: #b42318;
      font-size: 13px;
      white-space: pre-wrap;
    }
    .update-error[hidden] { display: none; }
    /* .modal-actions 가 display:flex 라 [hidden] 만으로는 안 숨는다. */
    .update-actions { justify-content: flex-end; }
    .update-actions[hidden] { display: none; }
    main { padding: 18px 24px 28px; }
    .controls {
      display: grid;
      grid-template-columns: repeat(3, minmax(260px, 1fr));
      gap: 12px;
      align-items: stretch;
      margin-bottom: 14px;
    }
    .group {
      min-height: 238px;
      display: flex;
      flex-direction: column;
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 12px;
    }
    .group-title {
      font-size: 13px;
      font-weight: 700;
      margin-bottom: 10px;
    }
    .group-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      margin-bottom: 10px;
      min-height: 38px;
    }
    .group-head .group-title {
      margin-bottom: 0;
    }
    .header-action {
      height: 38px;
      min-width: 96px;
      padding: 0 10px;
      font-size: 13px;
    }
    .header-placeholder {
      width: 96px;
      height: 38px;
      flex: 0 0 auto;
    }
    .date-row {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 8px;
    }
    .limit-row {
      margin-top: 10px;
    }
    .card-spacer {
      display: none;
    }
    .button-row {
      display: flex;
      align-items: center;
      justify-content: flex-start;
      gap: 8px;
      margin-top: auto;
      padding-top: 10px;
    }
    .button-row button {
      width: auto;
      min-width: 96px;
      padding: 0 10px;
    }
    .news-ranges {
      display: grid;
      gap: 8px;
    }
    .news-range-row {
      display: grid;
      grid-template-columns: minmax(118px, 1fr) minmax(118px, 1fr) minmax(138px, .9fr) 30px;
      gap: 8px;
      align-items: end;
    }
    .news-range-row label {
      min-width: 0;
      font-size: 11px;
    }
    .news-range-row input {
      width: 100%;
    }
    .news-range-remove {
      width: 30px;
      min-width: 30px;
      height: 32px;
      padding: 0;
      border-radius: 6px;
      font-size: 15px;
      font-weight: 900;
    }
    .news-keyword-detail {
      width: 100%;
      height: 32px;
      min-width: 0;
      padding: 0 10px;
      font-size: 12px;
      white-space: nowrap;
    }
    .news-range-tools {
      margin-top: 8px;
    }
    .news-range-add {
      width: auto;
      min-width: 0;
      height: 30px;
      padding: 0 10px;
      font-size: 12px;
    }
    .collect-row {
      margin-top: auto;
      padding-top: 10px;
    }
    .collect-row .button-row {
      margin-top: 0;
      padding-top: 0;
    }
    label {
      display: block;
      font-size: 12px;
      color: var(--muted);
    }
    input[type="date"], input[type="number"] {
      width: 100%;
      margin-top: 4px;
      padding: 8px 9px;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
      background: #fff;
    }
    .side {
      display: flex;
      justify-content: flex-end;
      margin: -2px 0 14px;
    }
    .side button {
      min-width: 104px;
      padding: 0 12px;
    }
    button {
      height: 38px;
      border: 1px solid var(--brand);
      border-radius: 6px;
      background: var(--brand);
      color: #fff;
      font-weight: 700;
      cursor: pointer;
    }
    button:hover { background: var(--brand-dark); }
    button.secondary {
      background: var(--brand);
      color: #fff;
      border-color: var(--brand);
    }
    button.secondary:hover { background: var(--brand-dark); }
    button.jb {
      background: #eef4ff;
      color: #1557bd;
      border-color: #c7ddff;
    }
    button.jb:hover {
      background: #dceaff;
      border-color: #9cc5ff;
    }
    button.ghost {
      background: #fff;
      color: var(--text);
      border-color: var(--line);
    }
    button.ghost:hover {
      background: #eef2f7;
    }
    button.danger {
      background: var(--brand);
      color: #fff;
      border-color: var(--brand);
    }
    button.danger:hover { background: var(--brand-dark); }
    button:disabled {
      opacity: 0.55;
      cursor: not-allowed;
    }
    .status {
      min-height: 28px;
      margin: 8px 0 12px;
      color: var(--muted);
      font-size: 13px;
    }
    .toolbar {
      display: flex;
      gap: 8px;
      align-items: center;
      justify-content: space-between;
      margin: 10px 0;
    }
    .summary {
      color: var(--muted);
      font-size: 13px;
    }
    .table-wrap {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: auto;
      max-height: calc(100vh - 270px);
      cursor: default;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
      cursor: default;
    }
    th, td {
      border-bottom: 1px solid var(--line);
      padding: 8px 10px;
      vertical-align: middle;
      white-space: nowrap;
      cursor: default;
    }
    th {
      position: sticky;
      top: 0;
      background: #f1f4f8;
      text-align: left;
      z-index: 1;
    }
    .sort-button {
      height: 26px;
      display: inline-flex;
      align-items: center;
      gap: 5px;
      padding: 0 7px;
      border: 1px solid transparent;
      background: transparent;
      color: var(--text);
      font: inherit;
      font-weight: 800;
    }
    .sort-button:hover,
    .sort-button.active {
      background: #fff;
      border-color: #cfd8e3;
    }
    .sort-indicator {
      width: 12px;
      color: var(--brand);
      font-size: 10px;
      line-height: 1;
    }
    .sort-menu {
      position: fixed;
      z-index: 30;
      width: 240px;
      padding: 6px;
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #fff;
      box-shadow: 0 14px 32px rgba(15, 23, 42, 0.18);
    }
    .sort-menu button {
      width: 100%;
      height: 30px;
      display: flex;
      align-items: center;
      justify-content: flex-start;
      padding: 0 9px;
      border: 0;
      border-radius: 6px;
      background: #fff;
      color: var(--text);
      font-size: 12px;
      font-weight: 700;
    }
    .sort-menu button:hover {
      background: #eef4ff;
      color: #1557bd;
    }
    .sort-menu-divider {
      height: 1px;
      margin: 6px 2px;
      background: var(--line);
    }
    .filter-actions {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 6px;
      margin: 6px 0;
    }
    .filter-actions button {
      justify-content: center;
      border: 1px solid var(--line);
      font-size: 11px;
    }
    .filter-values {
      max-height: 220px;
      overflow: auto;
      padding: 4px 2px;
    }
    .filter-value {
      min-height: 28px;
      display: flex;
      align-items: center;
      gap: 7px;
      padding: 4px 6px;
      border-radius: 6px;
      color: var(--text);
      font-size: 12px;
      cursor: pointer;
    }
    .filter-value:hover {
      background: #f1f5fb;
    }
    .filter-value input {
      width: 14px;
      height: 14px;
      accent-color: var(--brand);
    }
    .filter-value span {
      min-width: 0;
      overflow: hidden;
      text-overflow: ellipsis;
      white-space: nowrap;
    }
    td.keyword-cell {
      max-width: 260px;
      overflow: hidden;
      text-overflow: ellipsis;
    }
    td.title {
      min-width: 420px;
      max-width: 760px;
      white-space: normal;
      line-height: 1.35;
    }
    .action-cell {
      width: 64px;
      text-align: center;
    }
    .open-item-btn {
      width: 34px;
      height: 30px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      padding: 0;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .open-item-btn:hover {
      background: #eef4ff;
      border-color: #9cc5ff;
      color: #1557bd;
    }
    .open-item-btn svg {
      width: 17px;
      height: 17px;
      stroke: currentColor;
      stroke-width: 2.2;
      fill: none;
    }
    tr:hover { background: #f7fbff; }
    .tabs {
      display: flex;
      flex-wrap: wrap;
      gap: 6px;
      margin: 8px 0 10px;
    }
    .tab {
      height: 32px;
      padding: 0 11px;
      border: 1px solid var(--line);
      border-radius: 999px;
      background: #fff;
      color: var(--text);
      font-weight: 600;
      cursor: pointer;
    }
    .tab.active {
      background: #e8f1ff;
      border-color: #9cc5ff;
      color: #1557bd;
    }
    .badge {
      display: inline-flex;
      min-width: 24px;
      height: 22px;
      align-items: center;
      justify-content: center;
      border-radius: 999px;
      background: #eef4ff;
      color: #1d4ed8;
      font-weight: 700;
      font-size: 12px;
    }
    .kind {
      color: var(--muted);
      font-size: 12px;
    }
    .modal-backdrop {
      position: fixed;
      inset: 0;
      display: none;
      align-items: center;
      justify-content: center;
      padding: 24px;
      background: rgba(15, 23, 42, 0.42);
      z-index: 20;
    }
    .modal-backdrop.open { display: flex; }
    .modal {
      width: min(760px, 100%);
      max-height: min(720px, calc(100vh - 48px));
      display: grid;
      grid-template-rows: auto auto 1fr;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: 0 24px 70px rgba(15, 23, 42, 0.28);
      overflow: hidden;
    }
    .modal-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 16px 18px;
      border-bottom: 1px solid var(--line);
    }
    .modal-title {
      font-size: 17px;
      font-weight: 800;
    }
    .modal-close {
      width: 34px;
      height: 34px;
      padding: 0;
      background: #fff;
      color: var(--text);
      border-color: var(--line);
      font-size: 18px;
      line-height: 1;
    }
    .modal-close:hover { background: #eef2f7; }
    .history-modal {
      grid-template-rows: auto auto auto auto minmax(0, 1fr);
    }
    .modal-help {
      padding: 10px 18px;
      color: var(--muted);
      font-size: 13px;
    }
    .history-list {
      overflow: auto;
      padding: 8px;
    }
    .history-bulk {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 10px 12px 0;
    }
    .history-bulk-actions {
      display: inline-flex;
      align-items: center;
      justify-content: flex-end;
      gap: 8px;
      flex-wrap: wrap;
    }
    .history-bulk-left {
      display: grid;
      grid-template-columns: 34px auto;
      align-items: center;
      gap: 10px;
    }
    .history-check-label {
      display: inline-flex;
      align-items: center;
      color: var(--muted);
      font-size: 12px;
      font-weight: 800;
    }
    .history-check-box {
      width: 34px;
      height: 34px;
      display: flex;
      align-items: center;
      justify-content: center;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 7px;
      cursor: pointer;
    }
    .history-check-box:hover {
      border-color: var(--brand);
      background: #eef4ff;
    }
    .history-check-box input,
    .history-check {
      width: 16px;
      height: 16px;
      accent-color: var(--brand);
    }
    .history-item {
      width: 100%;
      height: auto;
      display: grid;
      grid-template-columns: 34px minmax(0, 1fr);
      gap: 12px;
      align-items: center;
      padding: 0;
      margin: 0 0 6px;
    }
    .history-card {
      min-height: 46px;
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      gap: 12px;
      align-items: center;
      padding: 12px 14px;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 8px;
      text-align: left;
      font-weight: 600;
      cursor: pointer;
    }
    .history-actions {
      display: inline-flex;
      gap: 8px;
      align-items: center;
      justify-content: flex-end;
      flex-wrap: wrap;
    }
    .history-card:hover {
      background: #f7fbff;
      border-color: #9cc5ff;
    }
    .history-name {
      font-size: 14px;
      font-weight: 800;
    }
    .history-meta {
      margin-top: 4px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 500;
    }
    .history-tags {
      display: inline-flex;
      gap: 6px;
      align-items: center;
      justify-content: flex-end;
      flex-wrap: wrap;
    }
    .history-tag {
      display: inline-flex;
      min-height: 24px;
      align-items: center;
      padding: 3px 8px;
      border-radius: 999px;
      background: #eef4ff;
      color: #1d4ed8;
      font-size: 12px;
      font-weight: 800;
    }
    .history-tag.report {
      background: #ecfdf3;
      color: #0f7b3d;
    }
    .history-delete {
      height: 30px;
      padding: 0 10px;
      background: #fff;
      color: #b42318;
      border: 1px solid #f1b9b4;
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .history-delete:hover {
      background: #fff1f0;
      border-color: #e0776e;
    }
    .history-merge-panel {
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto auto;
      gap: 8px;
      align-items: center;
      padding: 10px 12px 0;
    }
    .history-merge-panel[hidden] {
      display: none;
    }
    .history-merge-input {
      width: 100%;
      height: 34px;
      padding: 0 10px;
      border: 1px solid var(--line);
      border-radius: 7px;
      font-size: 13px;
      font-weight: 700;
    }
    .danger-outline {
      background: #fff;
      color: #b42318;
      border: 1px solid #f1b9b4;
    }
    .danger-outline:hover {
      background: #fff1f0;
      border-color: #e0776e;
      color: #b42318;
    }
    .keyword-modal {
      width: min(720px, calc(100vw - 40px));
    }
    .keyword-body {
      min-height: 0;
      overflow: auto;
      padding: 20px 22px 22px;
    }
    .keyword-editor {
      display: grid;
      grid-template-columns: minmax(0, 1fr) 86px;
      gap: 10px;
      margin-bottom: 14px;
    }
    .keyword-editor input {
      width: 100%;
      height: 38px;
      padding: 0 12px;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
    }
    .keyword-tabs {
      display: flex;
      gap: 8px;
      margin-bottom: 14px;
    }
    .keyword-tab {
      height: 32px;
      padding: 0 10px;
      background: #fff;
      color: var(--text);
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .keyword-tab.active {
      color: #fff;
      border-color: var(--brand);
      background: var(--brand);
    }
    .keyword-list {
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      max-height: 300px;
      overflow: auto;
      padding: 12px;
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #f8fafc;
    }
    .keyword-section-title {
      margin: 14px 0 8px;
      color: #475467;
      font-size: 12px;
      font-weight: 800;
    }
    .keyword-section-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      margin: 14px 0 8px;
    }
    .keyword-section-head .keyword-section-title {
      margin: 0;
    }
    .keyword-head-actions {
      display: flex;
      align-items: center;
      gap: 8px;
    }
    .keyword-select-all {
      height: 28px;
      padding: 0 9px;
      color: var(--brand);
      background: #fff;
      border: 1px solid #b8c7e8;
      border-radius: 6px;
      font-size: 12px;
      font-weight: 800;
    }
    .keyword-select-all:hover {
      background: #eef4ff;
      border-color: var(--brand);
    }
    .keyword-delete-btn {
      color: #b42318;
      border-color: #f3b8b1;
      background: #fff;
    }
    .keyword-delete-btn:hover {
      color: #fff;
      border-color: #d92d20;
      background: #d92d20;
    }
    .keyword-group-list {
      margin-top: 8px;
    }
    .keyword-item {
      width: auto;
      min-height: 32px;
      padding: 0 11px;
      text-align: center;
      color: var(--text);
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 999px;
      font-size: 13px;
      font-weight: 700;
    }
    .keyword-item[draggable="true"] {
      cursor: grab;
    }
    .keyword-item.dragging {
      opacity: .45;
      cursor: grabbing;
    }
    .keyword-drop-marker {
      width: 0;
      min-height: 32px;
      border-left: 3px solid var(--brand);
      border-radius: 999px;
      align-self: center;
      pointer-events: none;
      box-shadow: 0 0 0 2px rgba(31, 111, 235, .10);
    }
    .keyword-drag-ghost {
      position: fixed;
      left: -9999px;
      top: -9999px;
      display: flex;
      flex-wrap: wrap;
      gap: 6px;
      max-width: 360px;
      padding: 8px;
      background: rgba(255, 255, 255, .96);
      border: 1px solid var(--line);
      border-radius: 10px;
      box-shadow: var(--shadow);
    }
    .keyword-item.selected {
      background: #eef4ff;
      border-color: #1f6feb;
      color: #174ea6;
    }
    .keyword-filter-item {
      position: relative;
      display: inline-flex;
      flex-direction: column;
      align-items: center;
      justify-content: center;
      gap: 2px;
      min-height: 44px;
      line-height: 1.15;
      transition: color .15s ease, transform .15s ease, background .15s ease, border-color .15s ease;
    }
    .keyword-limit {
      color: var(--brand);
      font-size: 11px;
      font-weight: 900;
    }
    .keyword-filter-label {
      display: block;
    }
    .keyword-limit-input {
      width: 28px;
      height: 13px;
      padding: 0;
      text-align: center;
      color: var(--brand);
      background: transparent;
      border: 0;
      border-radius: 0;
      font-size: 11px;
      font-weight: 900;
      line-height: 1;
      font-family: inherit;
      appearance: textfield;
      -moz-appearance: textfield;
    }
    .keyword-limit-input::-webkit-outer-spin-button,
    .keyword-limit-input::-webkit-inner-spin-button {
      margin: 0;
      appearance: none;
      -webkit-appearance: none;
    }
    .keyword-limit-input:focus {
      outline: 1px solid #9cc5ff;
      outline-offset: 1px;
    }
    .keyword-limit-editing .keyword-filter-item {
      border-color: var(--line);
      background: #fff;
    }
    .keyword-filter-item::after {
      content: "↑";
      position: absolute;
      left: 50%;
      top: 50%;
      transform: translate(-50%, -38%);
      opacity: 0;
      color: #fff;
      font-size: 15px;
      font-weight: 900;
      transition: opacity .15s ease, transform .15s ease;
    }
    .keyword-filter-item:hover {
      color: rgba(255, 255, 255, .28);
      background: var(--brand);
      border-color: var(--brand);
      transform: translateY(-1px);
    }
    .keyword-filter-item:hover::after {
      opacity: 1;
      transform: translate(-50%, -62%);
    }
    .keyword-limit-editing .keyword-filter-item:hover {
      color: var(--text);
      background: #f8fbff;
      border-color: var(--brand);
      transform: none;
    }
    .keyword-limit-editing .keyword-filter-item:hover::after {
      opacity: 0;
    }
    .source-list {
      display: grid;
      gap: 8px;
      max-height: 340px;
      overflow: auto;
    }
    .source-item {
      min-height: 42px;
      display: flex;
      align-items: center;
      padding: 10px 12px;
      color: var(--text);
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
      font-weight: 700;
    }
    .modal-actions {
      display: flex;
      justify-content: space-between;
      gap: 8px;
      margin-top: 16px;
    }
    .keyword-action-left {
      display: flex;
      gap: 8px;
      flex-wrap: wrap;
    }
    .keyword-action-right {
      display: flex;
      gap: 8px;
      flex-wrap: wrap;
      justify-content: flex-end;
      margin-left: auto;
    }
    .keyword-bulk-limit {
      display: none;
      align-items: center;
      gap: 6px;
    }
    .keyword-bulk-limit.open {
      display: inline-flex;
    }
    .keyword-bulk-limit input {
      width: 64px;
      height: 38px;
      padding: 0 8px;
      text-align: center;
      border: 1px solid var(--line);
      border-radius: 6px;
      font-size: 13px;
      font-weight: 800;
    }
    .empty-state {
      padding: 36px 18px;
      color: var(--muted);
      text-align: center;
      font-size: 13px;
    }
    @media (max-width: 1100px) {
      .controls { grid-template-columns: 1fr; }
    }
  </style>
</head>
<body>
  <header>
    <div class="brand">
      <div class="brand-mark">JB</div>
      <div class="brand-copy">
        <h1>조사연구 도우미</h1>
      </div>
    </div>
  </header>
  <div class="update-bar" id="updateBar" hidden>
    <span class="update-mark" aria-hidden="true">↑</span>
    <div class="update-copy" id="updateBarCopy"></div>
    <button class="jb" id="updateNowBtn" type="button">지금 업데이트</button>
    <button class="update-dismiss" id="updateDismissBtn" type="button" title="이 버전 알림 닫기" aria-label="이 버전 알림 닫기">×</button>
  </div>
  <main>
    <section class="controls">
      <div class="group">
        <div class="group-head">
          <div class="group-title">뉴스</div>
          <button class="jb header-action" id="keywordBtn" type="button">검색 키워드</button>
        </div>
        <div class="news-ranges" id="newsRanges"></div>
        <div class="news-range-tools">
          <button class="ghost news-range-add" id="addNewsRangeBtn" type="button">+ 수집 기간 추가</button>
        </div>
        <div class="button-row">
          <button class="jb" id="collectNewsBtn" type="button">자료 수집</button>
        </div>
      </div>
      <div class="group">
        <div class="group-head">
          <div class="group-title">국가기관</div>
          <button class="jb header-action" id="agencyListBtn" type="button">기관 목록</button>
        </div>
        <div class="date-row">
          <label>시작일<input id="agencyStart" type="date"></label>
          <label>종료일<input id="agencyEnd" type="date"></label>
        </div>
        <div class="limit-row">
          <label>기관별 최대 건수<input id="agencyMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="card-spacer"></div>
        <div class="limit-row collect-row">
          <div class="button-row">
            <button class="jb" id="collectAgencyBtn" type="button">자료 수집</button>
          </div>
        </div>
      </div>
      <div class="group">
        <div class="group-head">
          <div class="group-title">금융연구소</div>
          <button class="jb header-action" id="researchListBtn" type="button">연구소 목록</button>
        </div>
        <div class="date-row">
          <label>시작일<input id="researchStart" type="date"></label>
          <label>종료일<input id="researchEnd" type="date"></label>
        </div>
        <div class="limit-row">
          <label>연구소별 최대 건수<input id="researchMax" type="number" min="1" max="100" value="10"></label>
        </div>
        <div class="card-spacer"></div>
        <div class="limit-row collect-row">
          <div class="button-row">
            <button class="jb" id="collectResearchBtn" type="button">자료 수집</button>
          </div>
        </div>
      </div>
    </section>

    <div class="side">
      <button class="secondary" id="collectBtn">일괄 수집</button>
    </div>

    <div class="status" id="status">기간을 선택한 뒤 수집을 시작하세요.</div>

    <div class="toolbar">
      <div class="summary" id="summary">선택 0건 / 전체 0건</div>
      <div>
        <button class="ghost" id="resetBtn">초기화</button>
        <button class="ghost" id="historyBtn">실행 기록</button>
        <button class="secondary" id="reportBtn">보고서 생성</button>
        <button class="ghost" id="openReportBtn" disabled>보고서 열기</button>
      </div>
    </div>

    <div class="tabs" id="tabs"></div>

    <div class="table-wrap">
      <table>
        <thead>
          <tr>
            <th><input id="headCheck" type="checkbox"></th>
            <th>구분</th>
            <th><button class="sort-button" type="button" data-sort-key="source_name"><span>언론/기관/연구소</span><span class="sort-indicator"></span></button></th>
            <th><button class="sort-button" type="button" data-sort-key="keywords"><span>키워드</span><span class="sort-indicator"></span></button></th>
            <th><button class="sort-button" type="button" data-sort-key="published_date"><span>일자</span><span class="sort-indicator"></span></button></th>
            <th>제목</th>
            <th>형식</th>
            <th></th>
          </tr>
        </thead>
        <tbody id="tbody"></tbody>
      </table>
    </div>
    <div class="sort-menu" id="sortMenu" hidden>
      <button type="button" data-sort-direction="asc">오름차순 정렬</button>
      <button type="button" data-sort-direction="desc">내림차순 정렬</button>
      <div class="sort-menu-divider"></div>
      <div class="filter-actions">
        <button type="button" id="filterSelectAllBtn">전체 선택</button>
        <button type="button" id="filterClearAllBtn">전체 해제</button>
      </div>
      <div class="filter-values" id="filterValues"></div>
    </div>
  </main>

  <div class="modal-backdrop" id="historyModal" aria-hidden="true">
    <section class="modal history-modal" role="dialog" aria-modal="true" aria-labelledby="historyTitle">
      <div class="modal-head">
        <div class="modal-title" id="historyTitle">실행 기록</div>
        <button class="modal-close" id="historyCloseBtn" type="button" title="닫기">×</button>
      </div>
      <div class="modal-help">불러올 실행 기록을 선택하세요. 보고서가 생성된 기록은 별도로 표시됩니다.</div>
      <div class="history-bulk">
        <div class="history-bulk-left">
          <label class="history-check-box" title="전체 선택"><input id="historySelectAll" type="checkbox"></label>
          <span class="history-check-label">전체 선택</span>
        </div>
        <div class="history-bulk-actions">
          <button class="jb" id="historyMergeSelectedBtn" type="button" disabled>수집 자료 병합</button>
          <button class="history-delete" id="historyDeleteSelectedBtn" type="button" disabled>삭제</button>
        </div>
      </div>
      <div class="history-merge-panel" id="historyMergePanel" hidden>
        <input class="history-merge-input" id="historyMergeTitleInput" type="text" aria-label="병합 실행 기록 제목">
        <button class="jb" id="historyMergeSaveBtn" type="button">저장</button>
        <button class="ghost" id="historyMergeCancelBtn" type="button">취소</button>
      </div>
      <div class="history-list" id="historyList"></div>
    </section>
  </div>

  <div class="modal-backdrop" id="keywordModal" aria-hidden="true">
    <section class="modal keyword-modal" role="dialog" aria-modal="true" aria-labelledby="keywordTitle">
      <div class="modal-head">
        <div class="modal-title" id="keywordTitle">뉴스 검색 키워드</div>
        <button class="modal-close" id="keywordCloseBtn" type="button" title="닫기">×</button>
      </div>
      <div class="modal-help">검색할 문구를 추가하거나 삭제하세요. 저장된 키워드는 다음 수집과 앱 재실행 때 그대로 사용됩니다.</div>
      <div class="keyword-body">
        <div class="keyword-tabs">
          <button class="keyword-tab active" id="keywordBankTab" type="button">은행/지주사</button>
          <button class="keyword-tab" id="keywordOtherTab" type="button">그외</button>
        </div>
        <div class="keyword-editor">
          <input id="keywordInput" type="text" placeholder="예: 전북은행">
          <button id="keywordAddBtn" type="button">추가</button>
        </div>
        <div class="keyword-section-head">
          <div class="keyword-section-title">키워드 목록</div>
          <div class="keyword-head-actions">
            <button class="keyword-select-all" id="keywordSelectAllBtn" type="button">전체 선택</button>
            <button class="keyword-select-all keyword-delete-btn" id="keywordDeleteBtn" type="button">삭제</button>
          </div>
        </div>
        <div class="keyword-list" id="keywordList"></div>
        <div class="modal-actions">
          <div class="keyword-action-left">
            <button class="jb" id="keywordFilterAddBtn" type="button" disabled>필터 추가</button>
            <button class="jb" id="keywordGroupAddBtn" type="button" disabled>필터 그룹 추가</button>
          </div>
          <div class="keyword-action-right">
            <div class="keyword-bulk-limit" id="keywordBulkLimit">
              <input id="keywordBulkLimitInput" type="number" min="1" max="100" value="0" aria-label="일괄 수집 건수">
              <button class="ghost" id="keywordBulkLimitBtn" type="button">일괄 적용</button>
            </div>
            <button class="ghost" id="keywordLimitEditBtn" type="button">수집 건수 편집</button>
          </div>
        </div>
        <div class="keyword-section-title">적용 필터</div>
        <div class="keyword-list keyword-group-list" id="keywordGroupList"></div>
      </div>
    </section>
  </div>

  <div class="modal-backdrop" id="sourceModal" aria-hidden="true">
    <section class="modal keyword-modal" role="dialog" aria-modal="true" aria-labelledby="sourceTitle">
      <div class="modal-head">
        <div class="modal-title" id="sourceTitle">수집 대상</div>
        <button class="modal-close" id="sourceCloseBtn" type="button" title="닫기">×</button>
      </div>
      <div class="modal-help" id="sourceHelp">현재 설정된 수집 대상 목록입니다.</div>
      <div class="keyword-body">
        <div class="source-list" id="sourceList"></div>
      </div>
    </section>
  </div>

  <div class="modal-backdrop" id="updateModal" aria-hidden="true">
    <section class="modal update-modal" role="dialog" aria-modal="true" aria-labelledby="updateTitle">
      <div class="modal-head">
        <div class="modal-title" id="updateTitle">업데이트</div>
        <button class="modal-close" id="updateCloseBtn" type="button" title="닫기">×</button>
      </div>
      <div class="modal-help" id="updateHelp"></div>
      <div class="update-body">
        <div class="update-notes" id="updateNotes" hidden></div>
        <p class="update-guide" id="updateGuide" hidden></p>
        <div class="update-progress" id="updateProgress" hidden>
          <span class="update-spinner" aria-hidden="true"></span>
          <span id="updatePhase"></span>
        </div>
        <div class="update-error" id="updateError" hidden></div>
        <div class="modal-actions update-actions" id="updateActions">
          <button class="ghost" id="updateLaterBtn" type="button">나중에</button>
          <button class="jb" id="updateApplyBtn" type="button">업데이트</button>
        </div>
      </div>
    </section>
  </div>

  <script>
    let rows = [];
    let activeTab = "all";
    let sortState = {key: "published_date", direction: "asc"};
    let activeSortKey = "";
    let columnValueFilters = {};
    let currentReportPath = "";
    let keywordState = {bank: {keywords: [], groups: []}, other: {keywords: [], groups: []}};
    let activeKeywordKind = "bank";
    let selectedKeywordIndex = -1;
    let selectedKeywordIndexes = new Set();
    let keywordBulkSelected = false;
    let draggedKeywordIndex = -1;
    let keywordDropIndex = -1;
    let keywordPersistSeq = 0;
    let keywordLimitEditMode = false;
    let activeKeywordRangeRow = null;
    let activeKeywordRunId = "";
    let activeKeywordScope = "global";
    let newsRangeSeq = 0;
    let historyRuns = [];
    let selectedHistoryRuns = new Set();
    let collectGeneration = 0;
    let activeCollectEvents = null;
    let generateGeneration = 0;
    let activeGenerateEvents = null;
    let heartbeatTimer = null;
    let updateInfo = null;
    let updateApplying = false;
    const UPDATE_DISMISS_KEY = "rrh-update-dismissed";
    const UPDATE_CHECK_INTERVAL_MS = 6 * 60 * 60 * 1000;
    const sourceLists = {
      agency: ["금융감독원", "금융위원회", "한국은행", "한국금융연구원"],
      research: [
        "하나금융연구소 연구보고서",
        "하나금융연구소 정기보고서",
        "KB경영연구소",
        "KDB미래전략연구소",
        "우리금융경영연구소 연구보고서",
        "우리금융경영연구소 동남아 Review"
      ]
    };

    function isoDate(d) {
      const year = d.getFullYear();
      const month = String(d.getMonth() + 1).padStart(2, "0");
      const day = String(d.getDate()).padStart(2, "0");
      return `${year}-${month}-${day}`;
    }
    function cloneKeywordState(value) {
      return normalizeKeywordState(JSON.parse(JSON.stringify(value || keywordState || {})));
    }
    function normalizeKeywordGroup(group) {
      let keywords = [];
      let limit = 3;
      if (Array.isArray(group)) {
        keywords = group;
      } else if (group && typeof group === "object") {
        keywords = Array.isArray(group.keywords) ? group.keywords : [];
        limit = Number(group.limit || group.max || 3);
      } else if (typeof group === "string") {
        keywords = [group];
      }
      const cleaned = [];
      for (const raw of keywords) {
        const text = String(raw || "").trim();
        if (text && !cleaned.includes(text)) cleaned.push(text);
      }
      return {keywords: cleaned, limit: Math.max(1, Math.min(100, Number.isFinite(limit) ? Math.floor(limit) : 3))};
    }
    function normalizeKeywordSectionClient(section = {}) {
      const keywords = [];
      for (const raw of Array.isArray(section.keywords) ? section.keywords : []) {
        const text = String(raw || "").trim();
        if (text && !keywords.includes(text)) keywords.push(text);
      }
      const rawGroups = Array.isArray(section.groups) ? section.groups : keywords.map(keyword => [keyword]);
      const groups = [];
      const seen = new Set();
      for (const rawGroup of rawGroups) {
        const group = normalizeKeywordGroup(rawGroup);
        const key = JSON.stringify(group.keywords);
        if (!group.keywords.length || seen.has(key)) continue;
        for (const keyword of group.keywords) {
          if (!keywords.includes(keyword)) keywords.push(keyword);
        }
        groups.push(group);
        seen.add(key);
      }
      return {keywords, groups};
    }
    function normalizeKeywordState(value = {}) {
      return {
        bank: normalizeKeywordSectionClient(value.bank || {}),
        other: normalizeKeywordSectionClient(value.other || {}),
      };
    }
    function groupKeywords(group) {
      return normalizeKeywordGroup(group).keywords;
    }
    function groupLimit(group) {
      return normalizeKeywordGroup(group).limit;
    }
    function groupLabel(group) {
      return groupKeywords(group).join(" AND ");
    }
    function keywordGroupCount(state) {
      const normalized = normalizeKeywordState(state);
      return (normalized.bank.groups || []).length + (normalized.other.groups || []).length;
    }
    function runNameFromPath(path) {
      const text = String(path || "").replaceAll("\\", "/");
      return text.split("/").filter(Boolean).pop() || "";
    }
    function useGlobalKeywordScope() {
      activeKeywordRunId = "";
      activeKeywordScope = "global";
    }
    function useRunKeywordScope(runId, keywords = null) {
      activeKeywordRunId = String(runId || "");
      activeKeywordScope = activeKeywordRunId ? "run" : "global";
      if (keywords) keywordState = normalizeKeywordState(keywords);
    }
    function maybeUseRunKeywords(runName, loaded) {
      if (!loaded || loaded.is_merged || !loaded.has_run_keywords) return;
      if (confirm("해당 건의 키워드 목록을 불러오시겠습니까?")) {
        useRunKeywordScope(runName, loaded.keywords || null);
      }
    }
    function setRangeKeywords(row, keywords) {
      const cloned = cloneKeywordState(keywords);
      if (!keywordGroupCount(cloned)) {
        delete row.dataset.keywords;
        return;
      }
      row.dataset.keywords = JSON.stringify(cloned);
    }
    function getRangeKeywords(row) {
      if (!row || !row.dataset.keywords) return cloneKeywordState(keywordState);
      try {
        return cloneKeywordState(JSON.parse(row.dataset.keywords));
      } catch (e) {
        return cloneKeywordState(keywordState);
      }
    }
    function addNewsRange(values = {}) {
      const container = document.getElementById("newsRanges");
      const id = ++newsRangeSeq;
      const row = document.createElement("div");
      row.className = "news-range-row";
      row.dataset.rangeId = String(id);
      setRangeKeywords(row, values.keywords || keywordState);
      row.innerHTML = `
        <label>시작일<input class="news-start" type="date" value="${escapeHtml(values.start || "")}"></label>
        <label>종료일<input class="news-end" type="date" value="${escapeHtml(values.end || "")}"></label>
        <button class="ghost news-keyword-detail" type="button">세부 키워드 설정</button>
        <button class="danger-outline news-range-remove" type="button" title="수집 기간 삭제">×</button>
      `;
      row.querySelector(".news-keyword-detail").addEventListener("click", () => openKeywordModal(row));
      row.querySelector(".news-range-remove").addEventListener("click", () => {
        if (document.querySelectorAll(".news-range-row").length <= 1) return;
        row.remove();
      });
      container.appendChild(row);
    }
    function getNewsRanges() {
      return [...document.querySelectorAll(".news-range-row")].map(row => {
        const range = {
          news_start: row.querySelector(".news-start").value,
          news_end: row.querySelector(".news-end").value,
        };
        if (row.dataset.keywords) {
          range.news_keywords = getRangeKeywords(row);
        }
        return range;
      }).filter(range => range.news_start && range.news_end);
    }
    function setDefaults() {
      const end = new Date();
      const start = new Date();
      start.setDate(end.getDate() - 7);
      document.getElementById("newsRanges").innerHTML = "";
      newsRangeSeq = 0;
      addNewsRange({start: isoDate(end), end: isoDate(end)});
      for (const id of ["agencyStart", "researchStart"]) {
        document.getElementById(id).value = isoDate(start);
      }
      for (const id of ["agencyEnd", "researchEnd"]) {
        document.getElementById(id).value = isoDate(end);
      }
    }
    function payload(targets = {news: true, agency: true, research: true}) {
      const ranges = getNewsRanges();
      const firstRange = ranges[0] || {};
      return {
        news_start: firstRange.news_start || "",
        news_end: firstRange.news_end || "",
        news_ranges: ranges,
        agency_start: document.getElementById("agencyStart").value,
        agency_end: document.getElementById("agencyEnd").value,
        research_start: document.getElementById("researchStart").value,
        research_end: document.getElementById("researchEnd").value,
        news_bank_max: 3,
        news_other_max: 3,
        keyword_run_id: activeKeywordRunId || "",
        agency_max: Number(document.getElementById("agencyMax").value || 10),
        research_max: Number(document.getElementById("researchMax").value || 10),
        include_news: Boolean(targets.news),
        include_agency: Boolean(targets.agency),
        include_research: Boolean(targets.research)
      };
    }
    function updateSummary() {
      const selected = rows.filter(r => r.checked).length;
      const visible = visibleRows();
      const visibleSelected = visible.filter(r => r.checked).length;
      document.getElementById("summary").textContent = `선택 ${selected}건 / 전체 ${rows.length}건 · 현재 탭 ${visibleSelected}/${visible.length}건`;
      document.getElementById("headCheck").checked = visible.length > 0 && visibleSelected === visible.length;
      updateReportControls();
    }
    function setReportPath(path) {
      currentReportPath = path || "";
      updateReportControls();
    }
    function updateReportControls() {
      const selected = rows.filter(r => r.checked).length;
      const reportBtn = document.getElementById("reportBtn");
      reportBtn.disabled = selected === 0;
      reportBtn.textContent = currentReportPath ? "보고서 재생성" : "보고서 생성";
      document.getElementById("openReportBtn").disabled = !currentReportPath;
    }
    function tabKey(row) {
      if (row.category === "그외") return "뉴스 그외";
      if (row.category === "은행/지주사") return "뉴스 은행/지주사";
      if (row.category === "국가기관") return "국가기관";
      if (row.category === "금융연구소") return "금융연구소";
      return row.category || "기타";
    }
    function rowSortValue(row, key) {
      if (key === "published_date") return Date.parse(row.published_date || "") || 0;
      return String(row[key] || "").trim();
    }
    function compareRows(a, b) {
      const key = sortState.key || "published_date";
      const direction = sortState.direction === "asc" ? 1 : -1;
      const valueA = rowSortValue(a, key);
      const valueB = rowSortValue(b, key);
      let result = 0;
      if (typeof valueA === "number" && typeof valueB === "number") {
        result = valueA - valueB;
      } else {
        result = String(valueA).localeCompare(String(valueB), "ko", {numeric: true, sensitivity: "base"});
      }
      if (result === 0) {
        result = String(a.title || "").localeCompare(String(b.title || ""), "ko", {numeric: true, sensitivity: "base"});
      }
      return result * direction;
    }
    function rowPassesValueFilters(row) {
      for (const [key, selected] of Object.entries(columnValueFilters)) {
        if (!selected || selected.size === 0) return false;
        const value = String(row[key] || "");
        if (!selected.has(value)) return false;
      }
      return true;
    }
    function visibleRows() {
      return rows
        .filter(row => activeTab === "all" || tabKey(row) === activeTab)
        .filter(rowPassesValueFilters)
        .slice()
        .sort(compareRows);
    }
    function updateSortHeaders() {
      document.querySelectorAll(".sort-button").forEach(button => {
        const key = button.dataset.sortKey;
        const active = key === sortState.key || Boolean(columnValueFilters[key]);
        button.classList.toggle("active", active);
        const indicator = button.querySelector(".sort-indicator");
        if (indicator) indicator.textContent = key === sortState.key ? (sortState.direction === "asc" ? "\u25B2" : "\u25BC") : "";
      });
    }
    function closeSortMenu() {
      const menu = document.getElementById("sortMenu");
      menu.hidden = true;
      activeSortKey = "";
    }
    function uniqueColumnValues(key) {
      const values = [];
      const seen = new Set();
      rows
        .filter(row => activeTab === "all" || tabKey(row) === activeTab)
        .forEach(row => {
          const value = String(row[key] || "");
          if (seen.has(value)) return;
          seen.add(value);
          values.push(value);
        });
      return values.sort((a, b) => {
        if (key === "published_date") return (Date.parse(a) || 0) - (Date.parse(b) || 0);
        return a.localeCompare(b, "ko", {numeric: true, sensitivity: "base"});
      });
    }
    function selectedValuesForKey(key, values) {
      if (!columnValueFilters[key]) return new Set(values);
      return new Set(columnValueFilters[key]);
    }
    function setColumnValueFilter(key, selected, values) {
      if (selected.size === values.length) delete columnValueFilters[key];
      else columnValueFilters[key] = selected;
    }
    function renderFilterValues(key) {
      const list = document.getElementById("filterValues");
      const values = uniqueColumnValues(key);
      const selected = selectedValuesForKey(key, values);
      list.innerHTML = "";
      for (const value of values) {
        const label = document.createElement("label");
        label.className = "filter-value";
        label.title = value || "(\uBE48 \uAC12)";
        label.innerHTML = `
          <input type="checkbox" value="${escapeHtml(value)}" ${selected.has(value) ? "checked" : ""}>
          <span>${escapeHtml(value || "(\uBE48 \uAC12)")}</span>
        `;
        label.querySelector("input").addEventListener("change", () => {
          const current = selectedValuesForKey(key, values);
          if (label.querySelector("input").checked) current.add(value);
          else current.delete(value);
          setColumnValueFilter(key, current, values);
          render();
          renderFilterValues(key);
        });
        list.appendChild(label);
      }
    }
    function openSortMenu(button) {
      const menu = document.getElementById("sortMenu");
      activeSortKey = button.dataset.sortKey || "";
      const rect = button.getBoundingClientRect();
      renderFilterValues(activeSortKey);
      menu.style.left = `${Math.min(rect.left, window.innerWidth - 252)}px`;
      menu.style.top = `${rect.bottom + 4}px`;
      menu.hidden = false;
    }
    function renderTabs() {
      const tabs = document.getElementById("tabs");
      const counts = new Map();
      for (const row of rows) {
        const key = tabKey(row);
        counts.set(key, (counts.get(key) || 0) + 1);
      }
      const tabOrder = ["뉴스 그외", "뉴스 은행/지주사", "국가기관", "금융연구소"];
      const ordered = ["all", ...tabOrder.filter(key => counts.has(key))];
      if (activeTab !== "all" && !ordered.includes(activeTab)) activeTab = "all";
      tabs.innerHTML = "";
      for (const key of ordered) {
        const btn = document.createElement("button");
        btn.className = "tab" + (key === activeTab ? " active" : "");
        btn.textContent = key === "all" ? `전체 (${rows.length})` : `${key} (${counts.get(key)})`;
        btn.addEventListener("click", () => {
          activeTab = key;
          render();
        });
        tabs.appendChild(btn);
      }
    }
    function render() {
      renderTabs();
      updateSortHeaders();
      const tbody = document.getElementById("tbody");
      tbody.innerHTML = "";
      visibleRows().forEach((row) => {
        const index = rows.findIndex(r => r.id === row.id);
        const tr = document.createElement("tr");
        tr.dataset.id = row.id;
        tr.innerHTML = `
          <td><input type="checkbox" ${row.checked ? "checked" : ""}></td>
          <td>${escapeHtml(row.category || "")}</td>
          <td>${escapeHtml(row.source_name || "")}</td>
          <td class="keyword-cell" title="${escapeHtml(row.keywords || "")}">${escapeHtml(row.keywords || "")}</td>
          <td>${escapeHtml(row.published_date || "")}</td>
          <td class="title">${escapeHtml(row.title || "")}</td>
          <td><span class="kind">${escapeHtml(row.file_type || "")}</span></td>
          <td class="action-cell"><button class="open-item-btn" type="button" title="${row.file_type === "article" ? "기사 링크 열기" : "PDF 열기"}">${openItemIcon(row)}</button></td>
        `;
        tr.querySelector("input").addEventListener("change", e => {
          rows[index].checked = e.target.checked;
          updateReportControls();
          updateSummary();
        });
        tr.querySelector(".open-item-btn").addEventListener("click", () => openItem(row));
        tbody.appendChild(tr);
      });
      updateSummary();
    }
    function openItemIcon(row) {
      if (row.file_type === "article") {
        return `<svg viewBox="0 0 24 24" aria-hidden="true"><path d="M7 17 17 7"></path><path d="M9 7h8v8"></path><path d="M5 5v14h14"></path></svg>`;
      }
      return `<svg viewBox="0 0 24 24" aria-hidden="true"><path d="M7 3h7l5 5v13H7z"></path><path d="M14 3v6h5"></path><path d="M9 14h6"></path><path d="M9 17h6"></path></svg>`;
    }
    function escapeHtml(value) {
      return String(value)
        .replaceAll("&", "&amp;")
        .replaceAll("<", "&lt;")
        .replaceAll(">", "&gt;")
        .replaceAll('"', "&quot;");
    }
    function isNetworkErrorText(value) {
      const text = String(value || "").toLowerCase();
      return [
        "failed to fetch",
        "networkerror",
        "network error",
        "load failed",
        "name resolution",
        "getaddrinfo",
        "connection",
        "timed out",
        "timeout",
        "ssl",
        "urlopen",
        "winerror 100",
        "네트워크",
        "인터넷",
      ].some(part => text.includes(part));
    }
    function friendlyError(prefix, error) {
      const message = error instanceof Error ? error.message : String(error || "");
      if (isNetworkErrorText(message)) return "네트워크가 연결되어 있지 않습니다";
      return `${prefix}: ${message || "알 수 없는 오류"}`;
    }
    function isOffline() {
      return typeof navigator !== "undefined" && navigator.onLine === false;
    }
    function showNetworkUnavailable() {
      document.getElementById("status").textContent = "네트워크가 연결되어 있지 않습니다";
    }
    function setCollectButtonsDisabled(disabled) {
      for (const id of ["collectBtn", "collectNewsBtn", "collectAgencyBtn", "collectResearchBtn"]) {
        document.getElementById(id).disabled = disabled;
      }
    }
    async function collect(targets = {news: true, agency: true, research: true}) {
      if (isOffline()) {
        showNetworkUnavailable();
        return;
      }
      setCollectButtonsDisabled(true);
      setReportPath("");
      updateReportControls();
      let streaming = false;
      const generation = ++collectGeneration;
      document.getElementById("status").textContent = "수집 중입니다. 사이트 응답에 따라 1~3분 정도 걸릴 수 있습니다.";
      try {
        const res = await fetch("/collect-start", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify(payload(targets))
        });
        const data = await res.json();
        if (!res.ok || !data.ok) throw new Error(data.error || "failed to start collection");
        const events = new EventSource(`/collect-events?id=${encodeURIComponent(data.job_id)}`);
        activeCollectEvents = events;
        events.onmessage = (event) => {
          if (generation !== collectGeneration) return;
          const msg = JSON.parse(event.data);
          if (msg.type === "source_start") {
            document.getElementById("status").textContent = `Collecting: ${msg.source}`;
          } else if (msg.type === "source_done") {
            document.getElementById("status").textContent = `${msg.source}: ${msg.count} item(s) found`;
          } else if (msg.type === "source_error") {
            document.getElementById("status").textContent = isNetworkErrorText(msg.error)
              ? "네트워크가 연결되어 있지 않습니다"
              : `${msg.source}: ${msg.error}`;
          } else if (msg.type === "decode_start") {
            document.getElementById("status").textContent = `Decoding article URL (${msg.index}/${msg.total}): ${msg.title || ""}`;
          } else if (msg.type === "decode_done") {
            document.getElementById("status").textContent = `Decoded article URL (${msg.index}/${msg.total})`;
          } else if (msg.type === "download_start") {
            document.getElementById("status").textContent = `Saving PDF (${msg.index}/${msg.total}): ${msg.source}`;
          } else if (msg.type === "complete") {
            rows = msg.items.map(x => ({...x, checked: true}));
            useRunKeywordScope(msg.run_name || runNameFromPath(msg.run_dir), msg.keywords || null);
            setReportPath("");
            activeTab = "all";
            render();
            document.getElementById("status").textContent = `Collection complete: ${rows.length} item(s) · ${msg.run_dir}`;
            events.close();
            if (activeCollectEvents === events) activeCollectEvents = null;
            setCollectButtonsDisabled(false);
          } else if (msg.type === "fatal") {
            document.getElementById("status").textContent = friendlyError("Error", msg.error);
            events.close();
            if (activeCollectEvents === events) activeCollectEvents = null;
            setCollectButtonsDisabled(false);
          }
        };
        events.onerror = () => {
          if (generation !== collectGeneration) return;
          document.getElementById("status").textContent = "Progress connection closed.";
          events.close();
          if (activeCollectEvents === events) activeCollectEvents = null;
          setCollectButtonsDisabled(false);
        };
        streaming = true;
        return;
        if (!res.ok || !data.ok) throw new Error(data.error || "수집 실패");
        rows = data.items.map(x => ({...x, checked: true}));
        activeTab = "all";
        render();
        document.getElementById("status").textContent = `수집 완료: ${rows.length}건 · 저장 폴더: ${data.run_dir}`;
      } catch (err) {
        document.getElementById("status").textContent = friendlyError("오류", err);
      } finally {
        if (!streaming) setCollectButtonsDisabled(false);
      }
    }
    async function resetWorkspace() {
      collectGeneration += 1;
      generateGeneration += 1;
      if (activeCollectEvents) {
        activeCollectEvents.close();
        activeCollectEvents = null;
      }
      if (activeGenerateEvents) {
        activeGenerateEvents.close();
        activeGenerateEvents = null;
      }
      const res = await fetch("/reset-state", {method: "POST"});
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "초기화 오류: " + (data.error || "초기화할 수 없습니다.");
        return;
      }
      rows = [];
      activeTab = "all";
      useGlobalKeywordScope();
      setReportPath("");
      render();
      setCollectButtonsDisabled(false);
      document.getElementById("status").textContent = "기간을 선택한 뒤 수집을 시작하세요.";
    }
    async function shutdownApp() {
      const ok = confirm("조사연구 도우미를 종료할까요?");
      if (!ok) return;
      try {
        await fetch("/shutdown", {method: "POST"});
        document.getElementById("status").textContent = "앱을 종료합니다. 브라우저 창을 닫아도 됩니다.";
      } catch (err) {
        document.getElementById("status").textContent = "종료 요청을 보냈습니다. 브라우저 창을 닫아도 됩니다.";
      }
    }
    function startHeartbeat() {
      if (heartbeatTimer) clearInterval(heartbeatTimer);
      const ping = () => fetch("/heartbeat", {method: "POST"}).catch(() => {});
      ping();
      heartbeatTimer = setInterval(ping, 60 * 1000);
    }

    /* ── 업데이트 ──────────────────────────────────────────────────────────
       배너 → 확인 창 → POST /update/apply. 서버는 팩을 받아 SHA-256 을 검증하고 업데이터를
       띄운 뒤 스스로 종료한다. 그때부터 이 화면이 /heartbeat 를 두드리며 기다리다, 서버가
       새 버전으로 응답하면 스스로 새로고침한다 — 사용자가 앱을 닫거나 설치 파일을 받을 일이 없다. */
    function readUpdateDismissed() {
      try {
        return localStorage.getItem(UPDATE_DISMISS_KEY) || "";
      } catch (err) {
        return "";
      }
    }
    function writeUpdateDismissed(version) {
      try {
        localStorage.setItem(UPDATE_DISMISS_KEY, version || "");
      } catch (err) {
        /* 저장 못 해도 알림만 다시 뜰 뿐이다 */
      }
    }
    function formatUpdateSize(bytes) {
      /* 실제 팩은 수십 MB 지만, 반올림해서 "0MB" 로 보이는 것보다는 단위를 낮추는 편이 낫다. */
      const value = Number(bytes || 0);
      if (value <= 0) return "";
      if (value >= 10 * 1024 * 1024) return `${(value / 1024 / 1024).toFixed(0)}MB`;
      if (value >= 1024 * 1024) return `${(value / 1024 / 1024).toFixed(1)}MB`;
      return `${Math.max(1, Math.round(value / 1024))}KB`;
    }
    function renderUpdateBar() {
      const bar = document.getElementById("updateBar");
      const show = Boolean(updateInfo && updateInfo.available)
        && readUpdateDismissed() !== updateInfo.latest
        && !updateApplying;
      bar.hidden = !show;
      if (!show) return;
      const detail = [`현재 ${updateInfo.current || ""}`];
      const size = formatUpdateSize(updateInfo.size);
      if (size) detail.push(size);
      if (updateInfo.restart_required !== false) detail.push("적용 시 앱이 자동으로 재시작됩니다");
      const firstNote = String(updateInfo.notes || "").split("\n")[0].trim();
      if (firstNote) detail.push(firstNote);
      document.getElementById("updateBarCopy").innerHTML =
        `<strong>새 버전 ${escapeHtml(updateInfo.latest || "")}</strong>`
        + `<span class="update-muted"> 사용 가능 · ${escapeHtml(detail.join(" · "))}</span>`;
    }
    async function checkForUpdate() {
      /* 실패해도 조용하다 — /update/check 는 망이 막혀 있어도 200 + note 로 온다.
         첫 화면에 에러를 띄우지 않는 것이 이 API 의 규격이다. */
      try {
        const res = await fetch("/update/check", {cache: "no-store"});
        if (!res.ok) return;
        updateInfo = await res.json();
      } catch (err) {
        return;
      }
      renderUpdateBar();
    }
    function setUpdateModalMode(mode) {
      /* mode: "confirm" | "progress" | "error" */
      document.getElementById("updateNotes").hidden = mode !== "confirm" || !String(updateInfo && updateInfo.notes || "").trim();
      document.getElementById("updateGuide").hidden = mode !== "confirm";
      document.getElementById("updateProgress").hidden = mode !== "progress";
      document.getElementById("updateError").hidden = mode !== "error";
      document.getElementById("updateActions").hidden = mode === "progress";
      document.getElementById("updateLaterBtn").textContent = mode === "error" ? "닫기" : "나중에";
      document.getElementById("updateApplyBtn").hidden = mode !== "confirm";
      /* 적용 중에는 닫지 못하게 한다 — 이때 다른 작업을 시작하면 재시작에 함께 끊긴다. */
      document.getElementById("updateCloseBtn").hidden = mode === "progress";
    }
    function openUpdateModal() {
      if (!updateInfo || !updateInfo.available) return;
      const size = formatUpdateSize(updateInfo.size);
      document.getElementById("updateTitle").textContent = `버전 ${updateInfo.latest} 로 업데이트`;
      document.getElementById("updateHelp").textContent =
        `현재 ${updateInfo.current} · 새 버전 ${updateInfo.latest}`
        + (updateInfo.released_at ? ` · ${updateInfo.released_at}` : "");
      document.getElementById("updateNotes").textContent = String(updateInfo.notes || "");
      document.getElementById("updateGuide").innerHTML =
        (size ? `내려받을 용량 ${escapeHtml(size)}. ` : "")
        + (updateInfo.restart_required !== false
            ? "적용하면 앱이 스스로 종료·갱신·재시작합니다. 진행 중인 작업이 있으면 업데이트가 거부됩니다.<br>"
            : "재시작 없이 반영됩니다.<br>")
        + "수집 자료·실행 기록·API 키는 그대로 유지됩니다. 실패하면 이전 버전으로 자동 복구합니다.";
      setUpdateModalMode("confirm");
      const modal = document.getElementById("updateModal");
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
    }
    function closeUpdateModal() {
      if (updateApplying) return;
      const modal = document.getElementById("updateModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }
    function showUpdatePhase(message) {
      document.getElementById("updatePhase").textContent = message;
      setUpdateModalMode("progress");
    }
    function showUpdateError(message) {
      updateApplying = false;
      document.getElementById("updateError").textContent = message;
      setUpdateModalMode("error");
      renderUpdateBar();
    }
    /* 서버가 다시 뜰 때까지 기다린다 — 종료·교체·재기동까지 수십 초 걸릴 수 있다.

       주의: /update/apply 는 응답을 먼저 내보내고 잠시 뒤에 프로세스를 끝낸다. 그래서 이 함수가
       도는 첫 1~2초 동안은 **아직 살아 있는 옛 서버**가 정상 응답한다. 그것을 '돌아왔다'로 치면
       교체가 끝나기도 전에 새로고침해 버린다. 서버가 한 번 끊긴 것을 보거나, 버전이 바뀐 것을
       확인한 뒤에만 돌아온 것으로 본다. */
    async function waitForServer(previousVersion, timeoutMs = 180000) {
      const deadline = Date.now() + timeoutMs;
      let sawDown = false;
      while (Date.now() < deadline) {
        try {
          const res = await fetch("/heartbeat", {cache: "no-store"});
          if (res.ok) {
            const data = await res.json();
            const version = String(data.version || "");
            if (sawDown || (version && version !== previousVersion)) return version;
          } else {
            sawDown = true;
          }
        } catch (err) {
          sawDown = true;   /* 아직 안 떴다 — 끊긴 것을 봤다는 표시이기도 하다 */
        }
        await new Promise(resolve => setTimeout(resolve, 2000));
      }
      return null;
    }
    async function applyUpdate() {
      if (updateApplying) return;
      updateApplying = true;
      renderUpdateBar();
      showUpdatePhase("업데이트 파일을 받고 검증하는 중입니다… (수십 초)");
      let payload = null;
      try {
        const res = await fetch("/update/apply", {method: "POST"});
        payload = await res.json().catch(() => null);
        if (!res.ok || !payload || payload.ok === false) {
          /* 409 = 진행 중인 작업이 있음, 502 = 다운로드·검증 실패. 사유를 그대로 보여준다. */
          showUpdateError((payload && payload.error) || `업데이트를 시작하지 못했습니다. (HTTP ${res.status})`);
          return;
        }
      } catch (err) {
        showUpdateError("업데이트를 시작하지 못했습니다: " + err);
        return;
      }
      /* 하트비트는 그대로 둔다 — 죽은 서버로 가는 ping 은 조용히 실패하고, 새 서버가 뜨면
         곧바로 활동 표시가 이어져 유휴 종료 타이머가 리셋된다. */
      showUpdatePhase("앱을 다시 시작하는 중입니다… 이 창을 닫지 마세요.");
      const previous = (updateInfo && updateInfo.current) || "";
      const version = await waitForServer(previous);
      if (version === null) {
        showUpdateError("앱이 다시 뜨지 않았습니다. 이전 버전으로 되돌렸을 수 있습니다 — 시작 메뉴에서 다시 실행해 주세요.");
        return;
      }
      if (version === previous) {
        /* 되돌아왔다 — 업데이터가 새 버전을 못 띄워 이전 버전을 복원한 경우다. */
        showUpdateError(`새 버전이 적용되지 않아 이전 버전(${version})으로 되돌아갔습니다. 업데이트 창에 남은 사유를 확인해 주세요.`);
        return;
      }
      showUpdatePhase(`업데이트 완료 (${version}) — 화면을 새로 불러옵니다.`);
      writeUpdateDismissed("");
      setTimeout(() => window.location.reload(), 1200);
    }
    async function openItem(row) {
      if (row.file_type === "article" && isOffline()) {
        showNetworkUnavailable();
        return;
      }
      try {
      const res = await fetch("/open", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({id: row.id})
      });
      const data = await res.json();
      if (data.open_in_client && data.target) {
        window.open(data.target, "_blank", "noopener");
        return;
      }
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "자료 열기 오류: " + (data.error || "열 수 없습니다.");
      }
      } catch (err) {
        document.getElementById("status").textContent = friendlyError("자료 열기 오류", err);
      }
    }
    async function openReport() {
      if (!currentReportPath) return;
      const res = await fetch("/open-report", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "보고서 열기 오류: " + (data.error || "열 수 없습니다.");
      }
    }
    async function saveSelection() {
      const ids = rows.filter(r => r.checked).map(r => r.id);
      const res = await fetch("/save-selection", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({ids})
      });
      const data = await res.json();
      if (data.ok) {
        document.getElementById("status").textContent = `선택 항목 저장 완료: ${ids.length}건 · ${data.csv}`;
      } else {
        document.getElementById("status").textContent = "저장 오류: " + data.error;
      }
    }
    async function ensureApiKey() {
      const statusRes = await fetch("/api-key-status");
      const status = await statusRes.json();
      if (status.has_key) {
        const useExisting = confirm("저장된 OpenAI API key를 사용할까요? 취소를 누르면 새 key를 입력합니다.");
        if (useExisting) return true;
      }
      const key = prompt("OpenAI API key를 입력하세요. 로컬 파일에 암호화 저장됩니다.");
      if (!key) return false;
      const res = await fetch("/api-key", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({api_key: key})
      });
      const data = await res.json();
      if (!data.ok) {
        alert("API key 저장 실패: " + data.error);
        return false;
      }
      return true;
    }
    async function generateReport() {
      if (isOffline()) {
        showNetworkUnavailable();
        return;
      }
      const ids = rows.filter(r => r.checked).map(r => r.id);
      if (!ids.length) {
        alert("보고서에 포함할 자료를 하나 이상 선택하세요.");
        return;
      }
      const selectedRows = rows.filter(r => r.checked);
      const missingSections = [];
      if (!selectedRows.some(r => r.category === "은행/지주사" || r.category === "그외")) {
        missingSections.push("뉴스");
      }
      if (!selectedRows.some(r => r.category === "국가기관")) {
        missingSections.push("국가기관");
      }
      if (!selectedRows.some(r => r.category === "금융연구소")) {
        missingSections.push("금융연구소");
      }
      if (missingSections.length) {
        const okToContinue = confirm(`${missingSections.join(", ")} 수집 자료가 없습니다.\n자료가 없는 섹션은 비우고 보고서를 생성하시겠습니까?`);
        if (!okToContinue) return;
      }
      const ok = await ensureApiKey();
      if (!ok) return;
      const btn = document.getElementById("reportBtn");
      btn.disabled = true;
      const isRegenerate = Boolean(currentReportPath);
      const generation = ++generateGeneration;
      document.getElementById("status").textContent = isRegenerate
        ? "보고서 재생성 중입니다. 체크된 자료 기준으로 기존 결과물을 대체합니다."
        : "보고서 생성 중입니다. 선택하지 않은 PDF를 정리하고 OpenAI API를 호출합니다.";
      try {
        const res = await fetch("/generate-report-start", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({ids})
        });
        const data = await res.json();
        if (!res.ok || !data.ok) throw new Error(data.error || "보고서 생성 시작 실패");
        const events = new EventSource(`/collect-events?id=${encodeURIComponent(data.job_id)}`);
        activeGenerateEvents = events;
        events.onmessage = (event) => {
          if (generation !== generateGeneration) return;
          const msg = JSON.parse(event.data);
          if (msg.type === "generate_status") {
            document.getElementById("status").textContent = msg.message || "보고서 생성 중입니다.";
          } else if (msg.type === "generate_upload") {
            document.getElementById("status").textContent = `자료 업로드 중 (${msg.index}/${msg.total}): ${msg.title || ""}`;
          } else if (msg.type === "complete") {
            rows = msg.items.map(x => ({...x, checked: true}));
            setReportPath(msg.docx || "");
            activeTab = "all";
            render();
            document.getElementById("status").textContent = `${isRegenerate ? "보고서 재생성" : "보고서 생성"} 완료. 보고서 열기 버튼을 눌러 확인하세요.`;
            events.close();
            if (activeGenerateEvents === events) activeGenerateEvents = null;
            updateReportControls();
          } else if (msg.type === "fatal") {
            document.getElementById("status").textContent = friendlyError("보고서 생성 오류", msg.error);
            events.close();
            if (activeGenerateEvents === events) activeGenerateEvents = null;
            updateReportControls();
          }
        };
        events.onerror = () => {
          if (generation !== generateGeneration) return;
          document.getElementById("status").textContent = "보고서 생성 상태 연결이 종료되었습니다.";
          events.close();
          if (activeGenerateEvents === events) activeGenerateEvents = null;
          updateReportControls();
        };
      } catch (err) {
        document.getElementById("status").textContent = friendlyError("보고서 생성 오류", err);
        updateReportControls();
      }
    }
    async function loadHistory() {
      const res = await fetch("/runs");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "실행 기록 오류: " + data.error;
        return;
      }
      const label = data.runs.map((r, i) => `${i + 1}. ${r.name} (${r.count}건)`).join("\n");
      const picked = prompt("불러올 실행 기록 번호를 입력하세요.\n\n" + label);
      const index = Number(picked) - 1;
      if (!Number.isInteger(index) || index < 0 || index >= data.runs.length) return;
      const run = data.runs[index];
      const detail = await fetch(`/run?id=${encodeURIComponent(run.name)}`);
      const loaded = await detail.json();
      if (!loaded.ok) {
        document.getElementById("status").textContent = "실행 기록 불러오기 오류: " + loaded.error;
        return;
      }
      rows = loaded.items.map(x => ({...x, checked: true}));
      maybeUseRunKeywords(runName, loaded);
      setReportPath(loaded.docx || "");
      activeTab = "all";
      render();
      document.getElementById("status").textContent = `실행 기록 불러오기 완료: ${run.name}`;
    }
    async function loadHistoryModal() {
      const res = await fetch("/runs");
      const data = await res.json();
      if (!data.ok) {
        document.getElementById("status").textContent = "실행 기록 오류: " + data.error;
        return;
      }
      showHistoryModal(data.runs);
    }
    function showHistoryModal(runs) {
      const modal = document.getElementById("historyModal");
      const list = document.getElementById("historyList");
      historyRuns = runs;
      selectedHistoryRuns = new Set();
      updateHistoryBulkControls();
      list.innerHTML = "";
      if (!runs.length) {
        list.innerHTML = `<div class="empty-state">저장된 실행 기록이 없습니다.</div>`;
      }
      runs.forEach((run) => {
        const item = document.createElement("div");
        item.className = "history-item";
        item.innerHTML = `
          <label class="history-check-box" title="실행 기록 선택">
            <input class="history-check" type="checkbox" aria-label="실행 기록 선택">
          </label>
          <div class="history-card">
            <div>
              <div class="history-name">${escapeHtml(displayRunName(run))}</div>
            </div>
            <div class="history-actions">
              <div class="history-tags">
                <span class="history-tag">${Number(run.count || 0)}건</span>
                ${run.has_report ? `<span class="history-tag report">보고서</span>` : ""}
              </div>
            </div>
          </div>
        `;
        item.querySelector(".history-card").addEventListener("click", () => openRunFromModal(run.name));
        item.querySelector(".history-check").addEventListener("click", (event) => {
          event.stopPropagation();
        });
        item.querySelector(".history-check").addEventListener("change", (event) => {
          event.stopPropagation();
          if (event.target.checked) selectedHistoryRuns.add(run.name);
          else selectedHistoryRuns.delete(run.name);
          hideHistoryMergePanel();
          updateHistoryBulkControls();
        });
        list.appendChild(item);
      });
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
    }
    function updateHistoryBulkControls() {
      const total = historyRuns.length;
      const selected = selectedHistoryRuns.size;
      const all = document.getElementById("historySelectAll");
      all.checked = total > 0 && selected === total;
      all.indeterminate = selected > 0 && selected < total;
      document.getElementById("historyDeleteSelectedBtn").disabled = selected === 0;
      document.getElementById("historyMergeSelectedBtn").disabled = selected < 2;
    }
    function closeHistoryModal() {
      const modal = document.getElementById("historyModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
      hideHistoryMergePanel();
    }
    function displayRunName(run) {
      return (run && run.title) ? run.title : formatRunName(run ? run.name : "");
    }
    function formatRunName(name) {
      const m = String(name).match(/^(\d{2})(\d{2})(\d{2})_(\d{2})(\d{2})(?:_(\d+))?$/);
      if (!m) return name;
      const suffix = m[6] ? ` #${m[6]}` : "";
      return `20${m[1]}.${m[2]}.${m[3]} ${m[4]}:${m[5]}${suffix}`;
    }
    function selectedHistoryRunObjects() {
      return historyRuns.filter(run => selectedHistoryRuns.has(run.name));
    }
    function dotDate(value) {
      const text = String(value || "").trim();
      const m = text.match(/^(\d{4})-(\d{2})-(\d{2})$/);
      if (m) return `${m[1]}.${m[2]}.${m[3]}`;
      const folder = text.match(/^(\d{2})(\d{2})(\d{2})_/);
      if (folder) return `20${folder[1]}.${folder[2]}.${folder[3]}`;
      const now = new Date();
      const yyyy = now.getFullYear();
      const mm = String(now.getMonth() + 1).padStart(2, "0");
      const dd = String(now.getDate()).padStart(2, "0");
      return `${yyyy}.${mm}.${dd}`;
    }
    function defaultMergeTitle() {
      const runs = selectedHistoryRunObjects();
      const startDates = runs.map(run => run.start_date).filter(Boolean).sort();
      const endDates = runs.map(run => run.end_date).filter(Boolean).sort();
      const starts = startDates.length ? startDates : runs.map(run => run.name).filter(Boolean).sort();
      const ends = endDates.length ? endDates : runs.map(run => run.name).filter(Boolean).sort();
      const start = starts[0] || "";
      const end = ends[ends.length - 1] || start;
      return `${dotDate(start)}~${dotDate(end)} 병합`;
    }
    function showHistoryMergePanel() {
      if (selectedHistoryRuns.size < 2) return;
      const panel = document.getElementById("historyMergePanel");
      const input = document.getElementById("historyMergeTitleInput");
      input.value = defaultMergeTitle();
      panel.hidden = false;
      input.focus();
      input.select();
    }
    function hideHistoryMergePanel() {
      const panel = document.getElementById("historyMergePanel");
      if (panel) panel.hidden = true;
    }
    async function openRunFromModal(runName) {
      closeHistoryModal();
      const detail = await fetch(`/run?id=${encodeURIComponent(runName)}`);
      const loaded = await detail.json();
      if (!loaded.ok) {
        document.getElementById("status").textContent = "실행 기록 불러오기 오류: " + loaded.error;
        return;
      }
      rows = loaded.items.map(x => ({...x, checked: true}));
      maybeUseRunKeywords(runName, loaded);
      setReportPath(loaded.docx || "");
      activeTab = "all";
      render();
      document.getElementById("status").textContent = `실행 기록 불러오기 완료: ${runName}`;
    }
    async function deleteRunFromModal(runName) {
      if (!confirm(`${formatRunName(runName)} 실행 기록을 삭제할까요?\n\n해당 폴더의 PDF, JSON, 생성 보고서가 모두 삭제됩니다.`)) {
        return;
      }
      const res = await fetch("/delete-run", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({id: runName})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "실행 기록 삭제 오류: " + (data.error || "삭제할 수 없습니다.");
        return;
      }
      if (data.cleared_current) {
        rows = [];
        setReportPath("");
        activeTab = "all";
        render();
      }
      document.getElementById("status").textContent = `실행 기록 삭제 완료: ${runName}`;
      await loadHistoryModal();
    }
    async function deleteSelectedHistoryRuns() {
      const names = [...selectedHistoryRuns];
      if (!names.length) return;
      if (!confirm(`선택한 실행 기록 ${names.length}개를 삭제할까요?\n\n해당 폴더의 PDF, JSON, 생성 보고서가 모두 삭제됩니다.`)) {
        return;
      }
      let clearedCurrent = false;
      for (const name of names) {
        const res = await fetch("/delete-run", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({id: name})
        });
        const data = await res.json();
        if (!res.ok || !data.ok) {
          document.getElementById("status").textContent = "실행 기록 삭제 오류: " + (data.error || "삭제할 수 없습니다.");
          await loadHistoryModal();
          return;
        }
        clearedCurrent = clearedCurrent || Boolean(data.cleared_current);
      }
      if (clearedCurrent) {
        rows = [];
        setReportPath("");
        activeTab = "all";
        render();
      }
      document.getElementById("status").textContent = `실행 기록 삭제 완료: ${names.length}개`;
      await loadHistoryModal();
    }
    async function mergeSelectedHistoryRuns() {
      const names = [...selectedHistoryRuns];
      if (names.length < 2) return;
      const input = document.getElementById("historyMergeTitleInput");
      const title = (input.value || "").trim() || defaultMergeTitle();
      const res = await fetch("/merge-runs", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({ids: names, title})
      });
      const data = await res.json();
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "수집 자료 병합 오류: " + (data.error || "병합할 수 없습니다.");
        return;
      }
      rows = data.items.map(x => ({...x, checked: true}));
      setReportPath("");
      activeTab = "all";
      render();
      closeHistoryModal();
      document.getElementById("status").textContent = `수집 자료 병합 완료: ${data.title || data.name} (${data.count}건)`;
    }

    async function openKeywordModal(rangeRow = null) {
      activeKeywordRangeRow = rangeRow;
      if (rangeRow) {
        if (!rangeRow.dataset.keywords) {
          const query = activeKeywordRunId ? `?run_id=${encodeURIComponent(activeKeywordRunId)}` : "";
          const res = await fetch(`/news-keywords${query}`);
          const data = await res.json();
          if (!data.ok) {
            document.getElementById("status").textContent = "??? ???? ??: " + data.error;
            return;
          }
          keywordState = normalizeKeywordState(data.keywords);
          setRangeKeywords(rangeRow, keywordState);
        } else {
          keywordState = getRangeKeywords(rangeRow);
        }
      } else {
        const query = activeKeywordRunId ? `?run_id=${encodeURIComponent(activeKeywordRunId)}` : "";
        const res = await fetch(`/news-keywords${query}`);
        const data = await res.json();
        if (!data.ok) {
          document.getElementById("status").textContent = "??? ???? ??: " + data.error;
          return;
        }
        keywordState = normalizeKeywordState(data.keywords);
        activeKeywordScope = data.scope || (activeKeywordRunId ? "run" : "global");
      }
      activeKeywordKind = "bank";
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      keywordBulkSelected = false;
      keywordLimitEditMode = false;
      renderKeywordModal();
      const modal = document.getElementById("keywordModal");
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
      document.getElementById("keywordInput").focus();
    }

    function closeKeywordModal() {
      const modal = document.getElementById("keywordModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
      activeKeywordRangeRow = null;
      keywordLimitEditMode = false;
    }

    function setKeywordKind(kind) {
      activeKeywordKind = kind;
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      keywordBulkSelected = false;
      renderKeywordModal();
    }

    function currentKeywordSection() {
      const section = normalizeKeywordSectionClient(keywordState[activeKeywordKind] || {keywords: [], groups: []});
      keywordState[activeKeywordKind] = section;
      return section;
    }

    function renderKeywordModal() {
      document.getElementById("keywordBankTab").className = "keyword-tab" + (activeKeywordKind === "bank" ? " active" : "");
      document.getElementById("keywordOtherTab").className = "keyword-tab" + (activeKeywordKind === "other" ? " active" : "");
      const section = currentKeywordSection();
      document.getElementById("keywordLimitEditBtn").className = keywordLimitEditMode ? "jb" : "ghost";
      document.getElementById("keywordBulkLimit").className = "keyword-bulk-limit" + (keywordLimitEditMode ? " open" : "");
      document.getElementById("keywordGroupList").className = "keyword-list keyword-group-list" + (keywordLimitEditMode ? " keyword-limit-editing" : "");
      const list = document.getElementById("keywordList");
      list.innerHTML = "";
      list.ondragover = (event) => {
        if (draggedKeywordIndex === -1) return;
        event.preventDefault();
        showKeywordDropMarkerAt(keywordInsertIndexFromPoint(event));
        event.dataTransfer.dropEffect = "move";
      };
      list.ondrop = (event) => {
        if (draggedKeywordIndex === -1) return;
        event.preventDefault();
        moveKeywordsTo(keywordDropIndex === -1 ? section.keywords.length : keywordDropIndex);
      };
      if (!section.keywords.length) {
        list.innerHTML = `<div class="empty-state">??? ???? ????.</div>`;
      }
      section.keywords.forEach((keyword, index) => {
        const btn = document.createElement("button");
        btn.type = "button";
        btn.className = "keyword-item" + (selectedKeywordIndexes.has(index) ? " selected" : "");
        btn.draggable = true;
        btn.dataset.index = String(index);
        btn.textContent = keyword;
        btn.addEventListener("dragstart", (event) => {
          draggedKeywordIndex = index;
          if (!selectedKeywordIndexes.has(index)) {
            selectedKeywordIndexes = new Set([index]);
            keywordBulkSelected = false;
            selectedKeywordIndex = index;
          }
          document.querySelectorAll("#keywordList .keyword-item.selected").forEach(el => el.classList.add("dragging"));
          btn.classList.add("dragging");
          event.dataTransfer.effectAllowed = "move";
          event.dataTransfer.setData("text/plain", String(index));
          const ghost = buildKeywordDragGhost();
          document.body.appendChild(ghost);
          event.dataTransfer.setDragImage(ghost, 12, 16);
          setTimeout(() => ghost.remove(), 0);
        });
        btn.addEventListener("dragend", () => {
          draggedKeywordIndex = -1;
          document.querySelectorAll("#keywordList .keyword-item.dragging").forEach(el => el.classList.remove("dragging"));
          clearKeywordDropMarkers();
        });
        btn.addEventListener("dragover", (event) => {
          event.preventDefault();
          if (draggedKeywordIndex !== -1) {
            showKeywordDropMarkerAt(index + (keywordDropSide(event, btn) === "after" ? 1 : 0));
            event.dataTransfer.dropEffect = "move";
          }
        });
        btn.addEventListener("drop", (event) => {
          event.preventDefault();
          moveKeywordsTo(keywordDropIndex);
        });
        btn.addEventListener("click", (event) => {
          if (selectedKeywordIndexes.has(index)) selectedKeywordIndexes.delete(index);
          else selectedKeywordIndexes.add(index);
          keywordBulkSelected = selectedKeywordIndexes.size === section.keywords.length && section.keywords.length > 0;
          selectedKeywordIndex = index;
          renderKeywordModal();
        });
        btn.addEventListener("dblclick", () => addKeywordGroup([keyword]));
        list.appendChild(btn);
      });
      document.getElementById("keywordFilterAddBtn").disabled = selectedKeywordIndexes.size === 0;
      document.getElementById("keywordGroupAddBtn").disabled = selectedKeywordIndexes.size < 2;
      const groupList = document.getElementById("keywordGroupList");
      groupList.innerHTML = "";
      if (!section.groups.length) {
        groupList.innerHTML = `<div class="empty-state">아직 추가된 그룹이 없습니다.</div>`;
      }
      section.groups.forEach((group, index) => {
        const btn = document.createElement("button");
        btn.type = "button";
        btn.className = "keyword-item keyword-filter-item";
        btn.draggable = false;
        btn.dataset.index = String(index);
        btn.innerHTML = keywordLimitEditMode
          ? `<input class="keyword-limit-input" type="number" min="1" max="100" value="${groupLimit(group)}" aria-label="최대 수집 건수"><span class="keyword-filter-label">${escapeHtml(groupLabel(group))}</span>`
          : `<span class="keyword-limit">${groupLimit(group)}</span><span class="keyword-filter-label">${escapeHtml(groupLabel(group))}</span>`;
        btn.title = keywordLimitEditMode ? "수집 건수를 편집합니다." : "클릭하면 적용 필터에서 제거됩니다.";
        btn.addEventListener("pointerdown", (event) => {
          event.stopPropagation();
          if (keywordLimitEditMode) return;
          event.preventDefault();
          removeKeywordFilter(index);
        });
        btn.addEventListener("keydown", (event) => {
          if (event.key !== "Enter" && event.key !== " ") return;
          event.preventDefault();
          event.stopPropagation();
          if (!keywordLimitEditMode) removeKeywordFilter(index);
        });
        const limitInput = btn.querySelector(".keyword-limit-input");
        if (limitInput) {
          limitInput.addEventListener("pointerdown", event => event.stopPropagation());
          limitInput.addEventListener("click", event => event.stopPropagation());
          limitInput.addEventListener("change", event => updateKeywordFilterLimit(index, event.target.value));
          limitInput.addEventListener("blur", event => updateKeywordFilterLimit(index, event.target.value));
          btn.addEventListener("click", () => {
            limitInput.focus();
            limitInput.select();
          });
        }
        groupList.appendChild(btn);
      });
    }

    function addKeyword() {
      const input = document.getElementById("keywordInput");
      const value = input.value.trim();
      if (!value) return;
      const section = currentKeywordSection();
      if (!section.keywords.includes(value)) {
        section.keywords.push(value);
        selectedKeywordIndex = section.keywords.length - 1;
        selectedKeywordIndexes = new Set([selectedKeywordIndex]);
      }
      keywordBulkSelected = false;
      input.value = "";
      renderKeywordModal();
      persistKeywords();
      input.focus();
    }

    function clearKeywordDropMarkers() {
      document.querySelectorAll(".keyword-drop-marker").forEach(el => el.remove());
      keywordDropIndex = -1;
    }

    function keywordDropSide(event, element) {
      const rect = element.getBoundingClientRect();
      return event.clientX < rect.left + rect.width / 2 ? "before" : "after";
    }

    function showKeywordDropMarkerAt(insertIndex) {
      const list = document.getElementById("keywordList");
      if (!list) return;
      const items = [...list.querySelectorAll(".keyword-item:not(.keyword-filter-item)")];
      const boundedIndex = Math.max(0, Math.min(insertIndex, items.length));
      keywordDropIndex = boundedIndex;
      const existing = list.querySelector(".keyword-drop-marker");
      if (existing) existing.remove();
      const marker = document.createElement("span");
      marker.className = "keyword-drop-marker";
      if (boundedIndex >= items.length) {
        list.appendChild(marker);
      } else {
        list.insertBefore(marker, items[boundedIndex]);
      }
    }

    function keywordInsertIndexFromPoint(event) {
      const items = [...document.querySelectorAll("#keywordList .keyword-item:not(.keyword-filter-item)")];
      for (let index = 0; index < items.length; index += 1) {
        const rect = items[index].getBoundingClientRect();
        const inRow = event.clientY >= rect.top - 6 && event.clientY <= rect.bottom + 6;
        if (inRow && event.clientX < rect.left + rect.width / 2) return index;
        if (event.clientY < rect.top) return index;
      }
      return items.length;
    }

    function buildKeywordDragGhost() {
      const section = currentKeywordSection();
      const selected = [...selectedKeywordIndexes].sort((a, b) => a - b);
      const ghost = document.createElement("div");
      ghost.className = "keyword-drag-ghost";
      for (const index of selected) {
        const chip = document.createElement("span");
        chip.className = "keyword-item selected";
        chip.textContent = section.keywords[index] || "";
        ghost.appendChild(chip);
      }
      return ghost;
    }

    function moveKeywordsTo(insertIndex) {
      const section = currentKeywordSection();
      const selected = [...selectedKeywordIndexes].sort((a, b) => a - b);
      if (!selected.length || insertIndex < 0 || insertIndex > section.keywords.length) return;
      if (selected.includes(insertIndex) || selected.includes(insertIndex - 1)) return;
      const selectedSet = new Set(selected);
      const moving = selected.map(index => section.keywords[index]);
      const remaining = section.keywords.filter((_, index) => !selectedSet.has(index));
      const adjustedInsert = insertIndex - selected.filter(index => index < insertIndex).length;
      section.keywords = [
        ...remaining.slice(0, adjustedInsert),
        ...moving,
        ...remaining.slice(adjustedInsert),
      ];
      selectedKeywordIndexes = new Set(moving.map((_, offset) => adjustedInsert + offset));
      selectedKeywordIndex = adjustedInsert + moving.length - 1;
      keywordBulkSelected = selectedKeywordIndexes.size === section.keywords.length && section.keywords.length > 0;
      draggedKeywordIndex = -1;
      clearKeywordDropMarkers();
      renderKeywordModal();
      persistKeywords();
    }

    function deleteKeyword() {
      const section = currentKeywordSection();
      const selected = [...selectedKeywordIndexes].sort((a, b) => b - a);
      if (!selected.length) return;
      const deleted = new Set();
      for (const index of selected) {
        if (index >= 0 && index < section.keywords.length) {
          deleted.add(section.keywords[index]);
          section.keywords.splice(index, 1);
        }
      }
      section.groups = section.groups.filter(group => !groupKeywords(group).some(keyword => deleted.has(keyword)));
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      keywordBulkSelected = false;
      renderKeywordModal();
      persistKeywords();
    }

    function addSelectedKeywordFilters() {
      const section = currentKeywordSection();
      const selected = [...selectedKeywordIndexes].sort((a, b) => a - b).map(index => section.keywords[index]).filter(Boolean);
      for (const keyword of selected) {
        addKeywordGroup([keyword], false);
      }
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      keywordBulkSelected = false;
      renderKeywordModal();
      persistKeywords();
    }

    function selectAllKeywords() {
      const section = currentKeywordSection();
      selectedKeywordIndex = section.keywords.length ? section.keywords.length - 1 : -1;
      selectedKeywordIndexes = new Set(section.keywords.map((_, index) => index));
      keywordBulkSelected = true;
      renderKeywordModal();
    }

    function addSelectedKeywordGroup() {
      const section = currentKeywordSection();
      const group = [...selectedKeywordIndexes].sort((a, b) => a - b).map(index => section.keywords[index]).filter(Boolean);
      addKeywordGroup(group);
    }

    function addKeywordGroup(group, shouldRender = true) {
      if (!group.length) return;
      const section = currentKeywordSection();
      const key = JSON.stringify(group);
      if (!section.groups.some(existing => JSON.stringify(groupKeywords(existing)) === key)) {
        section.groups.push({keywords: group, limit: 3});
      }
      selectedKeywordIndex = -1;
      selectedKeywordIndexes = new Set();
      keywordBulkSelected = false;
      if (shouldRender) {
        renderKeywordModal();
        persistKeywords();
      }
    }

    function removeKeywordFilter(index) {
      const section = currentKeywordSection();
      if (index < 0 || index >= section.groups.length) return;
      section.groups.splice(index, 1);
      renderKeywordModal();
      persistKeywords(true);
    }

    function updateKeywordFilterLimit(index, value) {
      const section = currentKeywordSection();
      if (index < 0 || index >= section.groups.length) return;
      const group = normalizeKeywordGroup(section.groups[index]);
      group.limit = Math.max(1, Math.min(100, Math.floor(Number(value) || 3)));
      section.groups[index] = group;
      persistKeywords(false);
    }

    function applyKeywordFilterLimitAll() {
      const section = currentKeywordSection();
      if (!section.groups.length) return;
      const input = document.getElementById("keywordBulkLimitInput");
      const limit = Math.max(1, Math.min(100, Math.floor(Number(input.value) || 3)));
      input.value = String(limit);
      section.groups = section.groups.map(group => {
        const normalized = normalizeKeywordGroup(group);
        normalized.limit = limit;
        return normalized;
      });
      renderKeywordModal();
      persistKeywords(false);
    }

    function toggleKeywordLimitEdit() {
      keywordLimitEditMode = !keywordLimitEditMode;
      renderKeywordModal();
    }

    async function persistKeywords(renderAfter = false) {
      keywordState = normalizeKeywordState(keywordState);
      if (activeKeywordRangeRow) {
        setRangeKeywords(activeKeywordRangeRow, keywordState);
        if (renderAfter) renderKeywordModal();
        document.getElementById("status").textContent = "기간별 키워드가 적용되었습니다.";
        return;
      }
      const seq = ++keywordPersistSeq;
      const snapshot = JSON.parse(JSON.stringify(keywordState));
      const res = await fetch("/news-keywords", {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({keywords: snapshot, run_id: activeKeywordRunId || ""})
      });
      const data = await res.json();
      if (seq !== keywordPersistSeq) return;
      if (!res.ok || !data.ok) {
        document.getElementById("status").textContent = "키워드 저장 오류: " + (data.error || "알 수 없는 오류.");
        return;
      }
      keywordState = normalizeKeywordState(data.keywords);
      activeKeywordScope = data.scope || (activeKeywordRunId ? "run" : "global");
      if (renderAfter) renderKeywordModal();
      document.getElementById("status").textContent = "검색 키워드가 적용되었습니다.";
    }

    function openSourceModal(kind) {
      const title = kind === "agency" ? "기관 목록" : "연구소 목록";
      const help = kind === "agency"
        ? "국가기관 자료 수집 시 아래 기관에서 자료를 가져옵니다."
        : "금융연구소 자료 수집 시 아래 연구소에서 자료를 가져옵니다.";
      document.getElementById("sourceTitle").textContent = title;
      document.getElementById("sourceHelp").textContent = help;
      const list = document.getElementById("sourceList");
      list.innerHTML = "";
      for (const source of sourceLists[kind] || []) {
        const item = document.createElement("div");
        item.className = "source-item";
        item.textContent = source;
        list.appendChild(item);
      }
      const modal = document.getElementById("sourceModal");
      modal.classList.add("open");
      modal.setAttribute("aria-hidden", "false");
    }

    function closeSourceModal() {
      const modal = document.getElementById("sourceModal");
      modal.classList.remove("open");
      modal.setAttribute("aria-hidden", "true");
    }

    document.getElementById("collectBtn").addEventListener("click", () => collect({news: true, agency: true, research: true}));
    document.getElementById("collectNewsBtn").addEventListener("click", () => collect({news: true, agency: false, research: false}));
    document.getElementById("collectAgencyBtn").addEventListener("click", () => collect({news: false, agency: true, research: false}));
    document.getElementById("collectResearchBtn").addEventListener("click", () => collect({news: false, agency: false, research: true}));
    document.getElementById("reportBtn").addEventListener("click", generateReport);
    document.getElementById("openReportBtn").addEventListener("click", openReport);
    document.getElementById("resetBtn").addEventListener("click", resetWorkspace);
    document.getElementById("historyBtn").addEventListener("click", loadHistoryModal);
    startHeartbeat();
    document.querySelectorAll(".sort-button").forEach(button => {
      button.addEventListener("click", event => {
        event.stopPropagation();
        openSortMenu(button);
      });
    });
    document.getElementById("sortMenu").addEventListener("click", event => {
      const button = event.target.closest("button[data-sort-direction]");
      if (!button || !activeSortKey) return;
      sortState = {key: activeSortKey, direction: button.dataset.sortDirection};
      render();
      renderFilterValues(activeSortKey);
    });
    document.getElementById("filterSelectAllBtn").addEventListener("click", event => {
      event.stopPropagation();
      if (!activeSortKey) return;
      delete columnValueFilters[activeSortKey];
      render();
      renderFilterValues(activeSortKey);
    });
    document.getElementById("filterClearAllBtn").addEventListener("click", event => {
      event.stopPropagation();
      if (!activeSortKey) return;
      columnValueFilters[activeSortKey] = new Set();
      render();
      renderFilterValues(activeSortKey);
    });
    document.addEventListener("click", event => {
      if (!event.target.closest("#sortMenu") && !event.target.closest(".sort-button")) {
        closeSortMenu();
      }
    });
    document.addEventListener("keydown", event => {
      if (event.key === "Escape") closeSortMenu();
    });
    document.getElementById("addNewsRangeBtn").addEventListener("click", () => {
      const ranges = getNewsRanges();
      const base = ranges[ranges.length - 1] || {};
      const previous = base.news_start ? new Date(`${base.news_start}T00:00:00`) : new Date();
      previous.setDate(previous.getDate() - 1);
      const previousIso = isoDate(previous);
      addNewsRange({
        start: previousIso,
        end: previousIso,
        keywords: cloneKeywordState(keywordState),
      });
    });
    document.getElementById("keywordBtn").addEventListener("click", () => openKeywordModal(null));
    document.getElementById("agencyListBtn").addEventListener("click", () => openSourceModal("agency"));
    document.getElementById("researchListBtn").addEventListener("click", () => openSourceModal("research"));
    document.getElementById("historyCloseBtn").addEventListener("click", closeHistoryModal);
    document.getElementById("historyMergeSelectedBtn").addEventListener("click", showHistoryMergePanel);
    document.getElementById("historyMergeSaveBtn").addEventListener("click", mergeSelectedHistoryRuns);
    document.getElementById("historyMergeCancelBtn").addEventListener("click", hideHistoryMergePanel);
    document.getElementById("historyMergeTitleInput").addEventListener("keydown", e => {
      if (e.key === "Enter") mergeSelectedHistoryRuns();
      if (e.key === "Escape") hideHistoryMergePanel();
    });
    document.getElementById("historyDeleteSelectedBtn").addEventListener("click", deleteSelectedHistoryRuns);
    document.getElementById("historySelectAll").addEventListener("change", e => {
      selectedHistoryRuns = e.target.checked ? new Set(historyRuns.map(run => run.name)) : new Set();
      document.querySelectorAll("#historyList .history-check").forEach(input => {
        const item = input.closest(".history-item");
        const index = [...document.querySelectorAll("#historyList .history-item")].indexOf(item);
        input.checked = Boolean(historyRuns[index] && selectedHistoryRuns.has(historyRuns[index].name));
      });
      hideHistoryMergePanel();
      updateHistoryBulkControls();
    });
    document.getElementById("keywordCloseBtn").addEventListener("click", closeKeywordModal);
    document.getElementById("sourceCloseBtn").addEventListener("click", closeSourceModal);
    document.getElementById("keywordBankTab").addEventListener("click", () => setKeywordKind("bank"));
    document.getElementById("keywordOtherTab").addEventListener("click", () => setKeywordKind("other"));
    document.getElementById("keywordAddBtn").addEventListener("click", addKeyword);
    document.getElementById("keywordSelectAllBtn").addEventListener("click", selectAllKeywords);
    document.getElementById("keywordDeleteBtn").addEventListener("click", deleteKeyword);
    document.getElementById("keywordFilterAddBtn").addEventListener("click", addSelectedKeywordFilters);
    document.getElementById("keywordGroupAddBtn").addEventListener("click", addSelectedKeywordGroup);
    document.getElementById("keywordLimitEditBtn").addEventListener("click", toggleKeywordLimitEdit);
    document.getElementById("keywordBulkLimitBtn").addEventListener("click", applyKeywordFilterLimitAll);
    document.getElementById("keywordBulkLimitInput").addEventListener("keydown", e => {
      if (e.key === "Enter") applyKeywordFilterLimitAll();
    });
    document.getElementById("keywordInput").addEventListener("keydown", e => {
      if (e.key === "Enter") addKeyword();
    });
    document.getElementById("historyModal").addEventListener("click", e => {
      if (e.target.id === "historyModal") closeHistoryModal();
    });
    document.getElementById("keywordModal").addEventListener("click", e => {
      if (e.target.id === "keywordModal") closeKeywordModal();
    });
    document.getElementById("sourceModal").addEventListener("click", e => {
      if (e.target.id === "sourceModal") closeSourceModal();
    });
    document.getElementById("updateNowBtn").addEventListener("click", openUpdateModal);
    document.getElementById("updateDismissBtn").addEventListener("click", () => {
      writeUpdateDismissed(updateInfo && updateInfo.latest);
      renderUpdateBar();
    });
    document.getElementById("updateCloseBtn").addEventListener("click", closeUpdateModal);
    document.getElementById("updateLaterBtn").addEventListener("click", closeUpdateModal);
    document.getElementById("updateApplyBtn").addEventListener("click", applyUpdate);
    document.getElementById("updateModal").addEventListener("click", e => {
      if (e.target.id === "updateModal") closeUpdateModal();
    });
    document.addEventListener("keydown", e => {
      if (e.key === "Escape") {
        closeHistoryModal();
        closeKeywordModal();
        closeSourceModal();
        closeUpdateModal();
      }
    });
    document.getElementById("headCheck").addEventListener("change", e => {
      visibleRows().forEach(r => r.checked = e.target.checked);
      render();
    });
    setDefaults();
    updateReportControls();
    checkForUpdate();
    setInterval(checkForUpdate, UPDATE_CHECK_INTERVAL_MS);
  </script>
</body>
</html>
"""


class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt: str, *args) -> None:
        print(f"[app] {self.address_string()} {fmt % args}")

    def do_GET(self) -> None:
        touch_server_activity()
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/":
            self.send_text(INDEX_HTML, "text/html; charset=utf-8")
            return
        if parsed.path == "/heartbeat":
            # version 을 함께 내려보낸다 — 업데이트 화면이 서버가 돌아오기를 기다리다가
            # 이 값으로 "새 버전이 떴다"를 확인한다.
            self.send_json({"ok": True, "version": updater.APP_VERSION})
            return
        if parsed.path == "/update/check":
            self.send_json(updater.check())
            return
        if parsed.path == "/update/version":
            self.send_json(
                {"ok": True, "version": updater.APP_VERSION, "installed": updater.installed()}
            )
            return
        if parsed.path == "/collect-events":
            self.handle_collect_events(parsed)
            return
        if parsed.path == "/api-key-status":
            self.send_json({"ok": True, "has_key": KEY_PATH.exists()})
            return
        if parsed.path == "/news-keywords":
            self.handle_get_news_keywords(parsed)
            return
        if parsed.path == "/runs":
            self.handle_runs()
            return
        if parsed.path == "/run":
            self.handle_run(parsed)
            return
        self.send_json({"ok": False, "error": "not found"}, status=404)

    def do_POST(self) -> None:
        touch_server_activity()
        try:
            if self.path == "/collect-start":
                self.handle_collect_start()
            elif self.path == "/collect":
                self.handle_collect()
            elif self.path == "/open":
                self.handle_open()
            elif self.path == "/open-report":
                self.handle_open_report()
            elif self.path == "/save-selection":
                self.handle_save_selection()
            elif self.path == "/api-key":
                self.handle_api_key()
            elif self.path == "/news-keywords":
                self.handle_news_keywords()
            elif self.path == "/generate-report-start":
                self.handle_generate_report_start()
            elif self.path == "/generate-report":
                self.handle_generate_report()
            elif self.path == "/delete-run":
                self.handle_delete_run()
            elif self.path == "/merge-runs":
                self.handle_merge_runs()
            elif self.path == "/reset-state":
                self.handle_reset_state()
            elif self.path == "/shutdown":
                self.handle_shutdown()
            elif self.path == "/update/apply":
                self.handle_update_apply()
            elif self.path == "/heartbeat":
                self.send_json({"ok": True, "version": updater.APP_VERSION})
            else:
                self.send_json({"ok": False, "error": "not found"}, status=404)
        except Exception as exc:
            traceback.print_exc()
            self.send_json({"ok": False, "error": str(exc)}, status=500)

    def handle_collect(self) -> None:
        data = self.read_json()
        news_start = parse_iso_date(data["news_start"])
        news_end = parse_iso_date(data["news_end"])
        agency_start = parse_iso_date(data["agency_start"])
        agency_end = parse_iso_date(data["agency_end"])
        research_start = parse_iso_date(data["research_start"])
        research_end = parse_iso_date(data["research_end"])
        news_bank_max = int(data.get("news_bank_max") or 10)
        news_other_max = int(data.get("news_other_max") or 10)
        agency_max = int(data.get("agency_max") or 10)
        research_max = int(data.get("research_max") or 10)
        news_keywords = base_keywords_from_payload(data)
        if isinstance(data.get("news_keywords"), dict):
            news_keywords = normalize_news_keywords(data.get("news_keywords"))
        include_news = bool(data.get("include_news", True))
        include_agency = bool(data.get("include_agency", True))
        include_research = bool(data.get("include_research", True))
        run_dir = collector.make_run_dir()
        save_news_keywords(news_keywords)
        save_run_keywords(run_dir, news_keywords)

        items = collector.collect_by_ranges(
            news_start,
            news_end,
            agency_start,
            agency_end,
            research_start,
            research_end,
            max_per_source=None,
            dry_run=False,
            news_bank_max=news_bank_max,
            news_other_max=news_other_max,
            agency_max=agency_max,
            research_max=research_max,
            output_dir=run_dir,
            news_keywords=news_keywords,
            include_news=include_news,
            include_agency=include_agency,
            include_research=include_research,
        )
        collector.write_outputs(items, output_dir=run_dir)
        payload = []
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = ""
            for idx, item in enumerate(items):
                row = row_for_client(item)
                row["id"] = idx
                payload.append(row)
        self.send_json({"ok": True, "items": payload, "run_dir": str(run_dir), "run_name": run_dir.name, "keywords": news_keywords})

    def handle_collect_start(self) -> None:
        data = self.read_json()
        job_id = uuid.uuid4().hex
        job_queue: queue.Queue = queue.Queue()
        job = {"queue": job_queue, "done": False}
        with STATE_LOCK:
            JOBS[job_id] = job
            STATE["report_path"] = ""
            STATE["active_collect_job_id"] = job_id
        thread = threading.Thread(target=run_collect_job, args=(job_id, data), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_collect_events(self, parsed) -> None:
        params = urllib.parse.parse_qs(parsed.query)
        job_id = params.get("id", [""])[0]
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if not job:
            self.send_json({"ok": False, "error": "invalid job id"}, status=404)
            return

        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Connection", "keep-alive")
        self.end_headers()

        job_queue = job["queue"]
        while True:
            try:
                event = job_queue.get(timeout=15)
            except queue.Empty:
                event = {"type": "ping"}
            data = f"data: {json.dumps(event, ensure_ascii=False)}\n\n".encode("utf-8")
            try:
                self.wfile.write(data)
                self.wfile.flush()
            except (BrokenPipeError, ConnectionResetError):
                break
            if event.get("type") in {"complete", "fatal"}:
                with STATE_LOCK:
                    JOBS.pop(job_id, None)
                break

    def handle_open(self) -> None:
        try:
            data = self.read_json()
            item_id = int(data["id"])
            with STATE_LOCK:
                items = STATE.get("items", [])
                run_dir_value = str(STATE.get("run_dir") or "")
                if item_id < 0 or item_id >= len(items):
                    self.send_json({"ok": False, "error": "invalid id"}, status=400)
                    return
                item = items[item_id]
            if item.file_type == "article":
                target_url = resolve_article_open_url(item)
                if target_url:
                    item.original_url = target_url
                    self.send_json({"ok": True, "target": target_url, "open_in_client": True})
                    return
                self.send_json({"ok": False, "error": "기사 URL이 없습니다."}, status=404)
                return
            run_dir = Path(run_dir_value) if run_dir_value else None
            path = resolve_item_local_path(item, run_dir)
            if path:
                os.startfile(str(path))
                self.send_json({"ok": True, "target": str(path)})
                return
            self.send_json({"ok": False, "error": "실행 기록 폴더에서 PDF 파일을 찾을 수 없습니다."}, status=404)
        except Exception as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=500)

    def handle_open_report(self) -> None:
        with STATE_LOCK:
            report_path = str(STATE.get("report_path") or "")
            run_dir_value = str(STATE.get("run_dir") or "")
        candidates = []
        if report_path:
            candidates.append(Path(report_path))
        if run_dir_value:
            candidates.append(Path(run_dir_value) / "generated_report.docx")
        for path in candidates:
            if path.exists() and path.is_file():
                os.startfile(path)
                self.send_json({"ok": True, "docx": str(path)})
                return
        self.send_json({"ok": False, "error": "생성된 보고서 파일을 찾을 수 없습니다."}, status=404)

    def handle_save_selection(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        with STATE_LOCK:
            selected = [item for i, item in enumerate(STATE.get("items", [])) if i in ids]
            run_dir_value = STATE.get("run_dir")
        run_dir = Path(run_dir_value) if run_dir_value else collector.OUT_DIR
        selected_csv = run_dir / "selected_metadata.csv"
        selected_xlsx = run_dir / "selected_metadata.xlsx"
        write_items(selected, selected_csv, selected_xlsx)
        self.send_json(
            {
                "ok": True,
                "count": len(selected),
                "csv": str(selected_csv),
                "xlsx": str(selected_xlsx),
            }
        )

    def handle_api_key(self) -> None:
        data = self.read_json()
        api_key = str(data.get("api_key") or "").strip()
        if not api_key:
            self.send_json({"ok": False, "error": "api_key is required"}, status=400)
            return
        save_api_key(api_key)
        self.send_json({"ok": True})

    def handle_get_news_keywords(self, parsed) -> None:
        run_id = urllib.parse.parse_qs(parsed.query).get("run_id", [""])[0]
        run_dir = resolve_run_dir(run_id) if run_id else None
        if run_dir is not None and run_dir.exists():
            self.send_json({"ok": True, "keywords": load_run_keywords(run_dir), "scope": "run"})
            return
        self.send_json({"ok": True, "keywords": load_news_keywords(), "scope": "global"})

    def handle_news_keywords(self) -> None:
        data = self.read_json()
        try:
            keywords = normalize_news_keywords(data.get("keywords", data))
        except ValueError as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        run_id = str(data.get("run_id") or "")
        run_dir = resolve_run_dir(run_id) if run_id else None
        if run_dir is not None and run_dir.exists():
            save_run_keywords(run_dir, keywords)
            self.send_json({"ok": True, "keywords": keywords, "scope": "run"})
            return
        save_news_keywords(keywords)
        self.send_json({"ok": True, "keywords": keywords, "scope": "global"})

    def handle_generate_report(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        result = perform_generate_report(ids)
        self.send_json(result)

    def handle_generate_report_start(self) -> None:
        data = self.read_json()
        ids = {int(x) for x in data.get("ids", [])}
        job_id = uuid.uuid4().hex
        job_queue: queue.Queue = queue.Queue()
        job = {"queue": job_queue, "done": False}
        with STATE_LOCK:
            JOBS[job_id] = job
            STATE["active_generate_job_id"] = job_id
        thread = threading.Thread(target=run_generate_report_job, args=(job_id, ids), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_runs(self) -> None:
        runs = []
        if collector.RUNS_DIR.exists():
            for path in sorted((p for p in collector.RUNS_DIR.iterdir() if p.is_dir()), reverse=True):
                summary = summarize_run_dir(path)
                runs.append(
                    {
                        "name": path.name,
                        "path": str(path),
                        "title": read_run_title(path),
                        "count": summary["count"],
                        "start_date": summary["start_date"],
                        "end_date": summary["end_date"],
                        "has_report": (path / "generated_report.docx").exists(),
                        "is_merged": is_merged_run(path),
                        "has_run_keywords": has_run_keywords(path),
                    }
                )
        self.send_json({"ok": True, "runs": runs})

    def handle_run(self, parsed) -> None:
        run_id = urllib.parse.parse_qs(parsed.query).get("id", [""])[0]
        run_dir = resolve_run_dir(run_id)
        if run_dir is None:
            self.send_json({"ok": False, "error": "invalid run id"}, status=400)
            return
        metadata = run_dir / "metadata.csv"
        if not metadata.exists():
            self.send_json({"ok": False, "error": "metadata.csv not found"}, status=404)
            return
        items: list[collector.Item] = []
        with metadata.open("r", encoding="utf-8-sig", newline="") as f:
            for idx, row in enumerate(csv.DictReader(f)):
                item = item_from_row(row)
                items.append(item)
        payload = []
        for idx, item in enumerate(items):
            out = row_for_client(item)
            out["id"] = idx
            payload.append(out)
        with STATE_LOCK:
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = str(run_dir / "generated_report.docx") if (run_dir / "generated_report.docx").exists() else ""
        self.send_json(
            {
                "ok": True,
                "items": payload,
                "run_dir": str(run_dir),
                "docx": str(run_dir / "generated_report.docx") if (run_dir / "generated_report.docx").exists() else "",
                "keywords": load_run_keywords(run_dir),
                "has_run_keywords": has_run_keywords(run_dir),
                "is_merged": is_merged_run(run_dir),
            }
        )

    def handle_merge_runs(self) -> None:
        data = self.read_json()
        run_ids = [str(value or "") for value in data.get("ids", [])]
        title = str(data.get("title") or "")
        result = merge_run_records(run_ids, title)
        payload = []
        for idx, item in enumerate(result["items"]):
            out = row_for_client(item)
            out["id"] = idx
            payload.append(out)
        with STATE_LOCK:
            STATE["items"] = result["items"]
            STATE["run_dir"] = str(result["run_dir"])
            STATE["report_path"] = ""
        self.send_json(
            {
                "ok": True,
                "name": result["run_dir"].name,
                "title": result["title"],
                "count": len(result["items"]),
                "items": payload,
                "run_dir": str(result["run_dir"]),
                "has_run_keywords": has_run_keywords(result["run_dir"]),
                "is_merged": True,
            }
        )

    def handle_delete_run(self) -> None:
        data = self.read_json()
        run_id = str(data.get("id") or "")
        run_dir = resolve_run_dir(run_id)
        if run_dir is None:
            self.send_json({"ok": False, "error": "invalid run id"}, status=400)
            return
        if not run_dir.exists() or not run_dir.is_dir():
            self.send_json({"ok": False, "error": "run not found"}, status=404)
            return
        cleared_current = False
        with STATE_LOCK:
            current = str(STATE.get("run_dir") or "")
            if current and same_path(current, run_dir):
                STATE["items"] = []
                STATE["run_dir"] = None
                STATE["report_path"] = ""
                cleared_current = True
        shutil.rmtree(run_dir)
        self.send_json({"ok": True, "deleted": str(run_dir), "cleared_current": cleared_current})

    def handle_reset_state(self) -> None:
        paths_to_delete: list[Path] = []
        with STATE_LOCK:
            active_job = str(STATE.get("active_collect_job_id") or "")
            if active_job:
                STATE.setdefault("cancelled_collect_jobs", set()).add(active_job)
            active_generate_job = str(STATE.get("active_generate_job_id") or "")
            if active_generate_job:
                STATE.setdefault("cancelled_generate_jobs", set()).add(active_generate_job)
            for key in ["pending_run_dir"]:
                value = str(STATE.get(key) or "")
                if value:
                    path = resolve_existing_run_dir(value)
                    if path is not None:
                        paths_to_delete.append(path)
                        STATE.setdefault("cancelled_run_dirs", set()).add(str(path.resolve()))
            STATE["items"] = []
            STATE["run_dir"] = None
            STATE["report_path"] = ""
            STATE["pending_run_dir"] = None
        for path in dict.fromkeys(paths_to_delete):
            try:
                if path.exists() and path.is_dir():
                    shutil.rmtree(path)
            except Exception:
                pass
        self.send_json({"ok": True})

    def handle_shutdown(self) -> None:
        self.send_json({"ok": True})
        threading.Thread(target=self.server.shutdown, daemon=True).start()

    def handle_update_apply(self) -> None:
        """업데이트 적용 — 받아서 검증하고 업데이터를 띄운 뒤 **이 프로세스를 종료**한다.

        스스로 죽는 이유: 실행 중인 exe 는 Windows 가 잠그고 있어 자기 자신을 갈아끼울 수 없다.
        """
        try:
            result = updater.apply(PORT, busy_reason=update_busy_reason)
        except updater.UpdateError as exc:
            self.send_json({"ok": False, "error": exc.message}, status=exc.status)
            return
        self.send_json(result)
        # 응답을 먼저 내보내고 종료한다 — 먼저 죽으면 화면은 "실패"로 본다.
        # os._exit 를 쓰는 것은 의도다. server.shutdown() 은 처리 중인 요청이 끝나기를 기다리고
        # 남은 스레드가 있으면 종료가 늘어지는데, 그동안 exe 잠금이 안 풀려 업데이터가 대기하다
        # 타임아웃된다. 어차피 저장할 상태는 이미 디스크에 있다.
        threading.Thread(target=_exit_soon, daemon=False).start()

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0"))
        raw = self.rfile.read(length).decode("utf-8")
        return json.loads(raw or "{}")

    def send_text(self, body: str, content_type: str) -> None:
        data = body.encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_json(self, data: dict, status: int = 200) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def perform_generate_report(ids: set[int], progress=None, is_cancelled=None, job_id: str = "") -> dict:
    def check_cancelled() -> None:
        if is_cancelled and is_cancelled():
            raise GenerateCancelled("보고서 생성이 초기화로 취소되었습니다.")

    def emit(message: str) -> None:
        check_cancelled()
        if progress:
            progress("generate_status", {"message": message})

    check_cancelled()
    emit("선택 항목을 확인하는 중입니다.")
    with STATE_LOCK:
        current_items = list(STATE.get("items", []))
        source_run_dir = str(STATE.get("run_dir") or "")
    if not current_items:
        raise RuntimeError("수집된 목록이 없습니다.")
    selected = [item for i, item in enumerate(current_items) if i in ids]
    unselected = [item for i, item in enumerate(current_items) if i not in ids]
    existing_run_dir = resolve_existing_run_dir(source_run_dir)
    if not selected:
        raise RuntimeError("선택된 항목이 없습니다.")

    check_cancelled()
    emit("기사 원문 링크를 확인하는 중입니다.")
    normalize_article_urls(selected, progress=progress)

    check_cancelled()
    emit("OpenAI API key를 확인하는 중입니다.")
    api_key = load_api_key()
    if not api_key:
        raise RuntimeError("OpenAI API key가 저장되어 있지 않습니다.")

    check_cancelled()
    emit("선택하지 않은 자료를 정리하는 중입니다.")
    if existing_run_dir:
        for item in unselected:
            delete_local_file_under(item.local_path, existing_run_dir)
    else:
        for item in unselected:
            delete_local_file(item.local_path)

    run_dir = existing_run_dir or collector.make_run_dir()
    check_cancelled()
    emit("실행 기록 폴더를 준비하는 중입니다.")
    if not existing_run_dir:
        move_selected_files(selected, run_dir, move=False)
    check_cancelled()
    enrich_article_texts(selected, progress=progress)
    check_cancelled()
    collector.write_outputs(selected, output_dir=run_dir)

    check_cancelled()
    emit("프롬프트와 입력 자료를 준비하는 중입니다.")
    prompt_path = ensure_prompt_template()
    prompt = prompt_path.read_text(encoding="utf-8")
    report_json = call_openai_report(api_key, prompt, selected, progress=progress)
    # If the user pressed reset while the LLM call was in flight, ignore the
    # returned response and avoid writing any generated report artifacts.
    check_cancelled()
    llm_output_text = str(report_json.pop("_openai_output_text", "") or "")
    llm_output_path = run_dir / "llm_output.txt"
    if llm_output_text:
        emit("LLM output text를 저장하는 중입니다.")
        check_cancelled()
        llm_output_path.write_text(llm_output_text, encoding="utf-8")
    report_json_path = run_dir / "report_data.json"
    emit("LLM 응답 JSON을 저장하는 중입니다.")
    check_cancelled()
    report_json_path.write_text(
        json.dumps(report_json, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    docx_path = run_dir / "generated_report.docx"
    emit("워드 보고서를 생성하는 중입니다.")
    check_cancelled()
    generate_docx(report_json, selected, docx_path)
    if is_cancelled and is_cancelled():
        for path in [llm_output_path, report_json_path, docx_path]:
            try:
                if path.exists():
                    path.unlink()
            except Exception:
                pass
        raise GenerateCancelled("보고서 생성이 초기화로 취소되었습니다.")

    check_cancelled()
    payload = []
    with STATE_LOCK:
        if job_id and job_id in STATE.get("cancelled_generate_jobs", set()):
            raise GenerateCancelled("보고서 생성이 초기화로 취소되었습니다.")
        STATE["items"] = selected
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = str(docx_path)
        for idx, item in enumerate(selected):
            row = row_for_client(item)
            row["id"] = idx
            payload.append(row)
    return {
        "ok": True,
        "items": payload,
        "run_dir": str(run_dir),
        "json": str(report_json_path),
        "llm_output": str(llm_output_path) if llm_output_text else "",
        "docx": str(docx_path),
    }


def enrich_article_texts(items: list[collector.Item], progress=None) -> None:
    article_items = [
        item for item in items
        if item.file_type == "article" and not item.extra.get("article_text") and (item.original_url or item.url)
    ]
    if not article_items:
        return
    normalize_article_urls(article_items)
    http = collector.Http()
    total = len(article_items)
    for index, item in enumerate(article_items, start=1):
        if progress:
            progress("generate_status", {"message": f"기사 본문 추출 중 ({index}/{total}): {item.title}"})
        text, note = collector.fetch_article_text(
            http,
            item.original_url or item.url,
            referer=item.extra.get("google_news_url", "") if isinstance(item.extra, dict) else "",
        )
        if text:
            item.extra["article_text"] = text
        if note:
            item.notes = "; ".join(x for x in [item.notes, note] if x)


def run_generate_report_job(job_id: str, ids: set[int]) -> None:
    def emit(event_type: str, payload: dict | None = None) -> None:
        payload = dict(payload or {})
        payload["type"] = event_type
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if job:
            job["queue"].put(payload)

    try:
        result = perform_generate_report(
            ids,
            progress=emit,
            is_cancelled=lambda: is_generate_job_cancelled(job_id),
            job_id=job_id,
        )
        if is_generate_job_cancelled(job_id):
            raise GenerateCancelled("보고서 생성이 초기화로 취소되었습니다.")
        emit("complete", result)
    except GenerateCancelled:
        with STATE_LOCK:
            JOBS.pop(job_id, None)
            STATE.get("cancelled_generate_jobs", set()).discard(job_id)
            if STATE.get("active_generate_job_id") == job_id:
                STATE["active_generate_job_id"] = ""
    except Exception as exc:
        traceback.print_exc()
        emit("fatal", {"error": str(exc)})
    finally:
        with STATE_LOCK:
            if (
                STATE.get("active_generate_job_id") == job_id
                and job_id not in STATE.get("cancelled_generate_jobs", set())
            ):
                STATE["active_generate_job_id"] = ""


def parse_iso_date(value: str) -> dt.date:
    return dt.date.fromisoformat(value)


def parse_news_ranges(data: dict) -> list[dict]:
    raw_ranges = data.get("news_ranges")
    ranges = []
    if isinstance(raw_ranges, list):
        for raw in raw_ranges:
            if not isinstance(raw, dict):
                continue
            start = str(raw.get("news_start") or raw.get("start") or "").strip()
            end = str(raw.get("news_end") or raw.get("end") or "").strip()
            if not start or not end:
                continue
            ranges.append(
                {
                    "start": parse_iso_date(start),
                    "end": parse_iso_date(end),
                    "bank_max": int(raw.get("news_bank_max") or raw.get("bank_max") or 10),
                    "other_max": int(raw.get("news_other_max") or raw.get("other_max") or 10),
                    "keywords": normalize_news_keywords(raw.get("news_keywords") or raw.get("keywords"))
                    if isinstance(raw.get("news_keywords") or raw.get("keywords"), dict)
                    else None,
                }
            )
    if not ranges:
        ranges.append(
            {
                "start": parse_iso_date(data["news_start"]),
                "end": parse_iso_date(data["news_end"]),
                "bank_max": int(data.get("news_bank_max") or 10),
                "other_max": int(data.get("news_other_max") or 10),
                "keywords": normalize_news_keywords(data.get("news_keywords"))
                if isinstance(data.get("news_keywords"), dict)
                else None,
            }
        )
    return ranges


def ensure_prompt_template() -> Path:
    if not PROMPT_PATH.exists():
        PROMPT_PATH.write_text(DEFAULT_REPORT_PROMPT, encoding="utf-8")
    return PROMPT_PATH


def default_news_keywords() -> dict:
    return {
        "bank": {
            "keywords": list(collector.DEFAULT_NEWS_KEYWORDS["bank"]["keywords"]),
            "groups": [
                {"keywords": list(group), "limit": 3}
                for group in collector.DEFAULT_NEWS_KEYWORDS["bank"]["groups"]
            ],
        },
        "other": {
            "keywords": list(collector.DEFAULT_NEWS_KEYWORDS["other"]["keywords"]),
            "groups": [
                {"keywords": list(group), "limit": 3}
                for group in collector.DEFAULT_NEWS_KEYWORDS["other"]["groups"]
            ],
        },
    }


def normalize_keyword_section(value, default_section: dict) -> dict:
    if isinstance(value, list):
        raw_keywords = value
        raw_groups = [[item] for item in value]
        explicit_groups = False
    elif isinstance(value, dict):
        raw_keywords = value.get("keywords", default_section["keywords"])
        explicit_groups = "groups" in value
        raw_groups = value.get("groups") if explicit_groups else [[item] for item in raw_keywords]
    else:
        explicit_groups = False
        raw_keywords = default_section["keywords"]
        raw_groups = default_section["groups"]
    if not isinstance(raw_keywords, list):
        raw_keywords = default_section["keywords"]
    keywords = []
    seen_keywords = set()
    for raw in raw_keywords:
        text = str(raw or "").strip().strip('"').strip("'").strip()
        if text and text not in seen_keywords:
            keywords.append(text)
            seen_keywords.add(text)
    groups = []
    seen_groups = set()
    if isinstance(raw_groups, list):
        for raw_group in raw_groups:
            limit = 3
            if isinstance(raw_group, dict):
                raw_items = raw_group.get("keywords") or raw_group.get("items") or []
                try:
                    limit = max(1, min(100, int(raw_group.get("limit") or raw_group.get("max") or 3)))
                except Exception:
                    limit = 3
            else:
                raw_items = [raw_group] if isinstance(raw_group, str) else raw_group
            if not isinstance(raw_items, list):
                continue
            group = []
            seen_in_group = set()
            for raw in raw_items:
                text = str(raw or "").strip().strip('"').strip("'").strip()
                if not text or text in seen_in_group:
                    continue
                group.append(text)
                seen_in_group.add(text)
                if text not in seen_keywords:
                    keywords.append(text)
                    seen_keywords.add(text)
            key = tuple(group)
            if group and key not in seen_groups:
                groups.append({"keywords": group, "limit": limit})
                seen_groups.add(key)
    if not keywords and not explicit_groups:
        keywords = list(default_section["keywords"])
    if not groups and not explicit_groups:
        groups = [{"keywords": [item], "limit": 3} for item in keywords]
    return {"keywords": keywords, "groups": groups}


def normalize_news_keywords(value) -> dict:
    if not isinstance(value, dict):
        raise ValueError("keywords must be an object")
    defaults = default_news_keywords()
    out = {}
    for key in ["bank", "other"]:
        out[key] = normalize_keyword_section(value.get(key, defaults[key]), defaults[key])
    return out


def load_news_keywords() -> dict:
    if NEWS_KEYWORDS_PATH.exists():
        try:
            keywords = normalize_news_keywords(json.loads(NEWS_KEYWORDS_PATH.read_text(encoding="utf-8")))
            save_news_keywords(keywords)
            return keywords
        except Exception:
            pass
    keywords = default_news_keywords()
    save_news_keywords(keywords)
    return keywords


def save_news_keywords(keywords: dict) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    NEWS_KEYWORDS_PATH.write_text(
        json.dumps(normalize_news_keywords(keywords), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def load_run_keywords(run_dir: Path) -> dict:
    path = run_dir / "run_keywords.json"
    try:
        if path.exists():
            return normalize_news_keywords(json.loads(path.read_text(encoding="utf-8")))
    except Exception:
        pass
    return load_news_keywords()


def save_run_keywords(run_dir: Path, keywords: dict) -> None:
    run_dir.mkdir(parents=True, exist_ok=True)
    (run_dir / "run_keywords.json").write_text(
        json.dumps(normalize_news_keywords(keywords), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def has_run_keywords(run_dir: Path) -> bool:
    return (run_dir / "run_keywords.json").exists()


def is_merged_run(run_dir: Path) -> bool:
    return (run_dir / MERGE_SOURCES_FILENAME).exists()


def base_keywords_from_payload(data: dict) -> dict:
    run_id = str(data.get("keyword_run_id") or "")
    run_dir = resolve_run_dir(run_id) if run_id else None
    if run_dir is not None and run_dir.exists():
        return load_run_keywords(run_dir)
    return load_news_keywords()


def delete_local_file(path_value: str) -> None:
    if not path_value:
        return
    path = Path(path_value)
    try:
        if path.exists() and path.is_file():
            path.unlink()
    except Exception:
        pass


def resolve_existing_run_dir(path_value: str) -> Path | None:
    if not path_value:
        return None
    try:
        path = Path(path_value).resolve()
        runs_root = collector.RUNS_DIR.resolve()
        if path.exists() and path.is_dir() and str(path).startswith(str(runs_root)):
            return path
    except Exception:
        return None
    return None


def resolve_run_dir(run_id: str) -> Path | None:
    if not re.fullmatch(r"\d{6}_\d{4}(?:_\d+)?", run_id or ""):
        return None
    try:
        runs_root = collector.RUNS_DIR.resolve()
        path = (runs_root / run_id).resolve()
        path.relative_to(runs_root)
        return path
    except Exception:
        return None


RUN_TITLE_FILENAME = "title.txt"
MERGE_SOURCES_FILENAME = "merge_sources.json"


def read_run_title(run_dir: Path) -> str:
    path = run_dir / RUN_TITLE_FILENAME
    try:
        if path.exists():
            return path.read_text(encoding="utf-8").strip()
    except Exception:
        pass
    return ""


def write_run_title(run_dir: Path, title: str) -> None:
    title = re.sub(r"\s+", " ", str(title or "")).strip()[:120]
    if title:
        (run_dir / RUN_TITLE_FILENAME).write_text(title, encoding="utf-8")


def summarize_run_dir(run_dir: Path) -> dict:
    metadata = run_dir / "metadata.csv"
    count = 0
    dates: list[str] = []
    if metadata.exists():
        with metadata.open("r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f):
                count += 1
                date_value = collector.normalize_date(row.get("published_date", "")) or ""
                if date_value:
                    dates.append(date_value)
    dates.sort()
    return {
        "count": count,
        "start_date": dates[0] if dates else "",
        "end_date": dates[-1] if dates else "",
    }


def title_from_item_dates(items: list[collector.Item]) -> str:
    dates = sorted(
        date_value
        for date_value in (collector.normalize_date(item.published_date) for item in items)
        if date_value
    )
    if not dates:
        today = dt.date.today().strftime("%Y.%m.%d")
        return f"{today}~{today} 병합"
    start = dates[0].replace("-", ".")
    end = dates[-1].replace("-", ".")
    return f"{start}~{end} 병합"


def normalize_merged_item_path(item: collector.Item, source_run_dir: Path) -> collector.Item:
    if not item.local_path:
        return item
    try:
        raw = Path(item.local_path)
        if not raw.is_absolute():
            item.local_path = str((source_run_dir / raw).resolve())
    except Exception:
        pass
    return item


def merge_run_records(run_ids: list[str], title: str = "") -> dict:
    unique_ids = list(dict.fromkeys(str(run_id or "") for run_id in run_ids))
    if len(unique_ids) < 2:
        raise ValueError("병합할 실행 기록을 2개 이상 선택하세요.")

    merged: list[collector.Item] = []
    source_names: list[str] = []
    for run_id in unique_ids:
        run_dir = resolve_run_dir(run_id)
        if run_dir is None:
            raise ValueError(f"invalid run id: {run_id}")
        metadata = run_dir / "metadata.csv"
        if not metadata.exists():
            raise FileNotFoundError(f"metadata.csv not found: {run_id}")
        source_names.append(run_id)
        with metadata.open("r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f):
                item = item_from_row(row)
                merged.append(normalize_merged_item_path(item, run_dir))

    merged = collector.dedupe(merged)
    if not merged:
        raise ValueError("병합할 수집 자료가 없습니다.")

    run_dir = collector.make_run_dir()
    collector.write_outputs(merged, output_dir=run_dir)
    title = re.sub(r"\s+", " ", str(title or "")).strip()[:120] or title_from_item_dates(merged)
    write_run_title(run_dir, title)
    (run_dir / MERGE_SOURCES_FILENAME).write_text(
        json.dumps(
            {
                "title": title,
                "sources": source_names,
                "created_at": dt.datetime.now().isoformat(timespec="seconds"),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return {"run_dir": run_dir, "title": title, "items": merged}


def same_path(left: str | Path, right: str | Path) -> bool:
    try:
        return Path(left).resolve() == Path(right).resolve()
    except Exception:
        return False


def delete_local_file_under(path_value: str, root: Path) -> None:
    if not path_value:
        return
    try:
        path = Path(path_value).resolve()
        root = root.resolve()
        if path.exists() and path.is_file() and str(path).startswith(str(root)):
            path.unlink()
    except Exception:
        pass


def resolve_item_local_path(item: collector.Item, run_dir: Path | None = None) -> Path | None:
    candidates: list[Path] = []
    if item.local_path:
        raw = Path(item.local_path)
        candidates.append(raw)
        if run_dir and not raw.is_absolute():
            candidates.append(run_dir / raw)
        if run_dir and raw.name:
            candidates.append(run_dir / raw.name)
            candidates.append(run_dir / "raw_pdfs" / raw.name)
    for candidate in candidates:
        try:
            path = candidate.resolve()
            if path.exists() and path.is_file():
                return path
        except Exception:
            continue
    if run_dir and item.local_path:
        name = Path(item.local_path).name
        if name:
            try:
                run_root = run_dir.resolve()
                for path in run_root.rglob(name):
                    if path.is_file():
                        path.resolve().relative_to(run_root)
                        return path
            except Exception:
                return None
    return None


def resolve_article_open_url(item: collector.Item) -> str:
    if item.original_url:
        return item.original_url
    if isinstance(item.extra, dict):
        google_url = str(item.extra.get("google_news_url") or "").strip()
        if google_url:
            return google_url
    return item.url or ""


def normalize_article_urls(items: list[collector.Item], progress=None) -> None:
    if progress is None:
        collector.populate_original_urls(items)
        return

    def decode_progress(event_type: str, payload: dict) -> None:
        if event_type == "decode_start":
            progress(
                "generate_status",
                {
                    "message": (
                        f"기사 원문 링크 변환 중 "
                        f"({payload.get('index', 0)}/{payload.get('total', 0)})..."
                    )
                },
            )
        elif event_type == "decode_done":
            progress(
                "generate_status",
                {
                    "message": (
                        f"기사 원문 링크 변환 완료 "
                        f"({payload.get('index', 0)}/{payload.get('total', 0)})"
                    )
                },
            )

    collector.populate_original_urls(items, progress=decode_progress)


def move_selected_files(items: list[collector.Item], run_dir: Path, move: bool = True) -> None:
    raw_root = run_dir / "raw_pdfs"
    for item in items:
        if not item.local_path:
            continue
        src = Path(item.local_path)
        if not src.exists() or not src.is_file():
            continue
        source_dir = raw_root / collector.clean_filename(item.source_name, item.source_name).replace(".pdf", "")
        source_dir.mkdir(parents=True, exist_ok=True)
        dst = source_dir / src.name
        if dst.exists():
            dst = source_dir / f"{dst.stem}_{uuid.uuid4().hex[:6]}{dst.suffix}"
        if move:
            shutil.move(str(src), str(dst))
        else:
            shutil.copy2(str(src), str(dst))
        item.local_path = str(dst)


def item_from_row(row: dict) -> collector.Item:
    extra = {}
    try:
        extra = json.loads(row.get("extra_json") or "{}")
    except json.JSONDecodeError:
        extra = {}
    item = collector.Item(
        category=row.get("category", ""),
        source_name=row.get("source_name", ""),
        title=row.get("title", ""),
        url=row.get("url", ""),
        published_date=row.get("published_date", ""),
        file_type=row.get("file_type", "pdf"),
        download_url=row.get("download_url", ""),
        local_path=row.get("local_path", ""),
        notes=row.get("notes", ""),
        original_url=row.get("original_url", ""),
        extra=extra,
    )
    if item.file_type == "article" and collector.is_daum_source_label(item.source_name):
        try:
            collector.normalize_daum_article_item(collector.Http(), item)
        except Exception:
            pass
    if item.source_name == "한국금융연구원":
        item.category = "국가기관"
    return item


def protect_with_dpapi(secret: str) -> bytes:
    data = secret.encode("utf-8")
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptProtectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData)
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def unprotect_with_dpapi(data: bytes) -> str:
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptUnprotectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData).decode("utf-8")
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def save_api_key(api_key: str) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    try:
        KEY_PATH.write_bytes(b"dpapi:" + protect_with_dpapi(api_key))
    except Exception:
        KEY_PATH.write_bytes(b"b64:" + base64.b64encode(api_key.encode("utf-8")))


def load_api_key() -> str:
    if not KEY_PATH.exists():
        return ""
    data = KEY_PATH.read_bytes()
    if data.startswith(b"dpapi:"):
        return unprotect_with_dpapi(data[6:])
    if data.startswith(b"b64:"):
        return base64.b64decode(data[4:]).decode("utf-8")
    return ""


class DATA_BLOB(ctypes.Structure):
    _fields_ = [("cbData", ctypes.wintypes.DWORD), ("pbData", ctypes.POINTER(ctypes.c_char))]


def multipart_request(url: str, fields: dict, files: list[tuple[str, Path]], api_key: str) -> dict:
    boundary = "----codexboundary" + uuid.uuid4().hex
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"),
                str(value).encode("utf-8"),
                b"\r\n",
            ]
        )
    for name, path in files:
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"; filename="{path.name}"\r\n'.encode("utf-8"),
                b"Content-Type: application/pdf\r\n\r\n",
                path.read_bytes(),
                b"\r\n",
            ]
        )
    chunks.append(f"--{boundary}--\r\n".encode("utf-8"))
    body = b"".join(chunks)
    req = urllib.request.Request(
        url,
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
    )
    opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
    with opener.open(req, timeout=180) as resp:
        return json.loads(resp.read().decode("utf-8"))


def openai_json_request(path: str, payload: dict, api_key: str, timeout: int = 300) -> dict:
    req = urllib.request.Request(
        f"https://api.openai.com/v1{path}",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
    )
    try:
        opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
        with opener.open(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", "replace")
        raise RuntimeError(f"OpenAI API error {exc.code}: {detail}") from exc


def call_openai_report(api_key: str, prompt: str, items: list[collector.Item], progress=None) -> dict:
    def emit(event_type: str, payload: dict | None = None) -> None:
        if progress:
            progress(event_type, payload or {})

    emit("generate_status", {"message": "LLM 입력 JSON 골격을 만드는 중입니다."})
    report_base = build_report_base(items)
    report_items = list(iter_report_items(report_base))
    merged_bullets = {"ITEMS": []}
    response_logs = []
    response_ids = []
    batch_size = 6
    total_batches = max((len(report_items) + batch_size - 1) // batch_size, 1)

    for batch_index, start in enumerate(range(0, len(report_items), batch_size), start=1):
        batch = report_items[start:start + batch_size]
        expected_ids = [item.get("ITEM_ID", "") for item in batch]
        pending = batch
        collected: dict[str, dict] = {}
        for attempt in range(1, 3):
            if not pending:
                break
            emit(
                "generate_status",
                {
                    "message": (
                        f"OpenAI에 bullet 생성을 요청하는 중입니다. "
                        f"({batch_index}/{total_batches}, {len(pending)}건)"
                    )
                },
            )
            response, text, parsed = request_openai_bullets(api_key, prompt, pending, emit)
            response_ids.append(response.get("id", ""))
            response_logs.append({
                "batch": batch_index,
                "attempt": attempt,
                "expected_item_ids": [item.get("ITEM_ID", "") for item in pending],
                "response_id": response.get("id", ""),
                "output_text": text,
            })
            for raw in extract_bullet_items(parsed):
                if not isinstance(raw, dict):
                    continue
                item_id = clean_cell_value(raw.get("ITEM_ID"))
                if item_id in expected_ids:
                    collected[item_id] = merge_bullet_object(collected.get(item_id), raw)
            missing = [item_id for item_id in expected_ids if item_id not in collected]
            if not missing:
                break
            pending = [item for item in batch if item.get("ITEM_ID", "") in missing]
            emit("generate_status", {"message": f"누락 bullet {len(missing)}건을 재요청하는 중입니다."})
        merged_bullets["ITEMS"].extend(collected[item_id] for item_id in expected_ids if item_id in collected)

    emit("generate_status", {"message": "LLM 응답을 JSON으로 정리하는 중입니다."})
    merged = merge_bullets_into_report(report_base, merged_bullets)
    merged["_openai_response_id"] = ", ".join(x for x in response_ids if x)
    merged["_openai_output_text"] = json.dumps(
        {
            "batch_size": batch_size,
            "expected_count": len(report_items),
            "received_count": len(merged_bullets["ITEMS"]),
            "responses": response_logs,
        },
        ensure_ascii=False,
        indent=2,
    )
    return merged


def request_openai_bullets(api_key: str, prompt: str, report_items: list[dict], emit) -> tuple[dict, str, dict]:
    materials = build_bullet_request_context(report_items)
    if "{{INPUT_MATERIALS}}" in prompt:
        prompt_text = prompt.replace("{{INPUT_MATERIALS}}", materials)
    else:
        prompt_text = prompt + "\n\n[입력 자료]\n" + materials
    content = [{"type": "input_text", "text": prompt_text}]
    pdf_items = [item for item in report_items if item.get("_LOCAL_PATH")]
    for index, item in enumerate(pdf_items, start=1):
        path = Path(item.get("_LOCAL_PATH", ""))
        if not path or not path.exists() or path.suffix.lower() != ".pdf":
            continue
        emit("generate_upload", {"index": index, "total": len(pdf_items), "title": item.get("TITLE", "")})
        uploaded = multipart_request(
            "https://api.openai.com/v1/files",
            {"purpose": "user_data"},
            [("file", path)],
            api_key,
        )
        content.append({"type": "input_file", "file_id": uploaded["id"]})
    response = openai_json_request(
        "/responses",
        {
            "model": OPENAI_MODEL,
            "input": [{"role": "user", "content": content}],
            "text": {"format": {"type": "json_object"}},
        },
        api_key,
    )
    text = extract_response_text(response)
    parsed = extract_json_object(text)
    return response, text, parsed


def build_bullet_request_context(report_base_or_items) -> str:
    rows = []
    source_items = report_base_or_items if isinstance(report_base_or_items, list) else iter_report_items(report_base_or_items)
    for item in source_items:
        rows.append({
            "ITEM_ID": item.get("ITEM_ID", ""),
            "SECTION": item.get("_SECTION", ""),
            "NO": item.get("NO", ""),
            "SOURCE_NAME": item.get("SOURCE_NAME", ""),
            "TITLE": item.get("TITLE", ""),
            "URL": item.get("URL", ""),
            "PUBLISHED_MM_DD": item.get("PUBLISHED_MM_DD", ""),
            "LOCAL_PATH": item.get("_LOCAL_PATH", ""),
            "CONTENT_TEXT": item.get("_CONTENT_TEXT", ""),
        })
    return json.dumps(rows, ensure_ascii=False, indent=2)


def build_report_base(items: list[collector.Item]) -> dict:
    bank_items = [item for item in items if item.category == "은행/지주사"]
    other_items = [item for item in items if item.category == "그외"]
    agency_items = [item for item in items if item.category == "국가기관"]
    research_items = [item for item in items if item.category == "금융연구소"]
    return {
        "BANK_SECTION": {
            "ITEMS": metadata_items(bank_items, "bank", "1", "BANK_SECTION.ITEMS"),
            "OTHER_ITEMS": metadata_items(other_items, "other", "", "BANK_SECTION.OTHER_ITEMS"),
        },
        "GOVERNMENT_SECTION": {
            "ITEMS": metadata_items(agency_items, "agency", "2", "GOVERNMENT_SECTION.ITEMS"),
        },
        "FINANCIAL_RESEARCH_SECTION": {
            "ITEMS": metadata_items(research_items, "research", "3", "FINANCIAL_RESEARCH_SECTION.ITEMS"),
        },
    }


def metadata_items(items: list[collector.Item], id_prefix: str, section_no: str, section_name: str) -> list[dict]:
    out = []
    for idx, item in enumerate(items, start=1):
        url = item.original_url or item.url
        out.append({
            "ITEM_ID": f"{id_prefix}_{idx}",
            "_SECTION": section_name,
            "_LOCAL_PATH": item.local_path,
            "NO": f"{section_no}.{idx}" if section_no else "",
            "SOURCE_NAME": clean_cell_value(item.source_name),
            "TITLE": clean_title(item.title, item.source_name),
            "URL": clean_cell_value(url),
            "PUBLISHED_MM_DD": clean_mm_dd(item.published_mm_dd),
            "_CONTENT_TEXT": content_text_for_item(item),
            "SUMMARY_BULLET_1": "",
            "SUMMARY_BULLET_2": "",
            "SUMMARY_BULLET_3": "",
        })
    return out


def content_text_for_item(item: collector.Item, limit: int = 5000) -> str:
    text = ""
    if isinstance(item.extra, dict):
        text = item.extra.get("article_text") or item.extra.get("rss_description") or ""
    text = clean_cell_value(text)
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + " ..."


BULLET_KEYS = ["SUMMARY_BULLET_1", "SUMMARY_BULLET_2", "SUMMARY_BULLET_3"]


def merge_bullet_object(dst: dict | None, src: dict) -> dict:
    """같은 ITEM_ID로 여러 조각(부분 bullet 객체)이 와도 빈 슬롯만 채워 병합한다.

    모델이 한 항목을 bullet마다 별도 객체로 쪼개 반환하는 경우, 덮어쓰기 방식이면
    마지막 조각만 남아 bullet이 1개로 붕괴한다. 여기서는 이미 값이 있는 슬롯은
    유지하고 비어 있는 슬롯만 채운다.
    """
    if dst is None:
        dst = {}
    for key, value in src.items():
        if key in BULLET_KEYS:
            if clean_cell_value(value) and not clean_cell_value(dst.get(key)):
                dst[key] = value
        elif not clean_cell_value(dst.get(key)) and clean_cell_value(value):
            dst[key] = value
    return dst


def merge_bullets_into_report(report_base: dict, bullet_json: dict) -> dict:
    by_id = {item.get("ITEM_ID", ""): item for item in iter_report_items(report_base)}
    by_url = {item.get("URL", ""): item for item in iter_report_items(report_base) if item.get("URL")}
    by_title = {item.get("TITLE", ""): item for item in iter_report_items(report_base) if item.get("TITLE")}
    for raw in extract_bullet_items(bullet_json):
        if not isinstance(raw, dict):
            continue
        item_id = clean_cell_value(raw.get("ITEM_ID"))
        target = by_id.get(item_id)
        if not target:
            target = by_url.get(clean_cell_value(raw.get("URL")))
        if not target:
            target = by_title.get(clean_title(raw.get("TITLE"), raw.get("SOURCE_NAME")))
        if not target:
            continue
        for key in BULLET_KEYS:
            value = clean_cell_value(raw.get(key))
            if value and not clean_cell_value(target.get(key)):
                target[key] = value
    return strip_internal_report_fields(report_base)


def extract_bullet_items(data: dict) -> list:
    if isinstance(data.get("ITEMS"), list):
        return data["ITEMS"]
    if isinstance(data.get("BULLETS"), list):
        return data["BULLETS"]
    out = []
    for section in ["BANK_SECTION", "GOVERNMENT_SECTION", "FINANCIAL_RESEARCH_SECTION"]:
        section_data = data.get(section, {})
        if not isinstance(section_data, dict):
            continue
        if isinstance(section_data.get("ITEMS"), list):
            out.extend(section_data["ITEMS"])
        if isinstance(section_data.get("OTHER_ITEMS"), list):
            out.extend(section_data["OTHER_ITEMS"])
    return out


def iter_report_items(report_data: dict):
    yield from report_data.get("BANK_SECTION", {}).get("ITEMS", [])
    yield from report_data.get("BANK_SECTION", {}).get("OTHER_ITEMS", [])
    yield from report_data.get("GOVERNMENT_SECTION", {}).get("ITEMS", [])
    yield from report_data.get("FINANCIAL_RESEARCH_SECTION", {}).get("ITEMS", [])


def strip_internal_report_fields(report_data: dict) -> dict:
    cleaned = deepcopy(report_data)
    for item in iter_report_items(cleaned):
        for key in ["ITEM_ID", "_SECTION", "_LOCAL_PATH", "_CONTENT_TEXT"]:
            item.pop(key, None)
    return cleaned


def extract_response_text(response: dict) -> str:
    if response.get("output_text"):
        return str(response["output_text"])
    parts = []
    for output in response.get("output", []):
        for content in output.get("content", []):
            if "text" in content:
                parts.append(str(content["text"]))
    return "\n".join(parts)


def extract_json_object(text: str) -> dict:
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            return json.loads(match.group(0))
        raise


def generate_docx(report_json: dict, items: list[collector.Item], output_path: Path) -> None:
    template = TEMPLATE_PATH
    doc = Document(template) if template.exists() else Document()
    add_section_bookmarks(doc)
    suppress_word_field_update(doc)
    report_data = normalize_report_json(report_json, items)
    populate_report_tables(doc, report_data)
    replace_docx_placeholders(
        doc,
        {
            "BANK_SECTION_PAGE": "__PAGEREF_BANK_SECTION__",
            "AGENCY_SECTION_PAGE": "__PAGEREF_AGENCY_SECTION__",
            "RESEARCH_SECTION_PAGE": "__PAGEREF_RESEARCH_SECTION__",
        },
    )
    for table in doc.tables[1:]:
        apply_malgun_to_table(table)
    if len(doc.tables) > 1:
        issue_row = doc.tables[1].rows[0]
        if len(issue_row.cells) > 2:
            for para in issue_row.cells[2].paragraphs:
                for run in para.runs:
                    run.bold = True
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(f".{output_path.stem}.{uuid.uuid4().hex[:8]}.tmp{output_path.suffix}")
    try:
        doc.save(tmp_path)
        replace_file_with_retry(tmp_path, output_path)
    finally:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


def replace_file_with_retry(src: Path, dst: Path, attempts: int = 5) -> None:
    last_exc: Exception | None = None
    for attempt in range(attempts):
        try:
            os.replace(src, dst)
            return
        except PermissionError as exc:
            last_exc = exc
            time.sleep(0.5 * (attempt + 1))
    raise PermissionError(
        f"{dst} 파일을 덮어쓸 수 없습니다. 보고서가 열려 있다면 닫은 뒤 다시 재생성하세요."
    ) from last_exc


def build_template_values(report_json: dict, items: list[collector.Item]) -> dict:
    values = {k: stringify_template_value(v) for k, v in report_json.items()}
    values.setdefault("REPORT_WEEK", dt.date.today().strftime("%Y-%m-%d"))
    values.setdefault("DEPARTMENT_NAME", "조사연구")
    groups = {
        "BANK_DETAIL_ITEMS": [x for x in items if "뉴스" in x.category or "은행" in x.category or "지주" in x.category],
        "AGENCY_DETAIL_ITEMS": [x for x in items if "기관" in x.category],
        "RESEARCH_DETAIL_ITEMS": [x for x in items if "연구" in x.category],
    }
    for key, group_items in groups.items():
        values.setdefault(key, format_detail_items(report_json.get(key), group_items))
    values["BANK_SECTION_PAGE"] = "__PAGEREF_BANK_SECTION__"
    values["AGENCY_SECTION_PAGE"] = "__PAGEREF_AGENCY_SECTION__"
    values["RESEARCH_SECTION_PAGE"] = "__PAGEREF_RESEARCH_SECTION__"
    values.setdefault("ITEM_NO", "")
    values.setdefault("TITLE", "")
    values.setdefault("URL", "")
    values.setdefault("SOURCE_NAME", "")
    values.setdefault("PUBLISHED_MM_DD", "")
    values.setdefault("SUMMARY_BULLET_1", "")
    values.setdefault("SUMMARY_BULLET_2", "")
    values.setdefault("SUMMARY_BULLET_3", "")
    values.setdefault("NO", "")
    return values


def normalize_report_json(report_json: dict, items: list[collector.Item]) -> dict:
    data = {
        "BANK_SECTION": {
            "ITEMS": normalize_items(report_json.get("BANK_SECTION", {}).get("ITEMS", []), "1"),
            "OTHER_ITEMS": normalize_items(report_json.get("BANK_SECTION", {}).get("OTHER_ITEMS", []), ""),
        },
        "GOVERNMENT_SECTION": {
            "ITEMS": normalize_items(report_json.get("GOVERNMENT_SECTION", {}).get("ITEMS", []), "2"),
        },
        "FINANCIAL_RESEARCH_SECTION": {
            "ITEMS": normalize_items(report_json.get("FINANCIAL_RESEARCH_SECTION", {}).get("ITEMS", []), "3"),
        },
    }
    has_items = (
        bool(data["BANK_SECTION"]["ITEMS"])
        or bool(data["BANK_SECTION"]["OTHER_ITEMS"])
        or bool(data["GOVERNMENT_SECTION"]["ITEMS"])
        or bool(data["FINANCIAL_RESEARCH_SECTION"]["ITEMS"])
    )
    if not has_items:
        data = fallback_report_data(items)
    return data


def normalize_items(raw_items, section_no: str) -> list[dict]:
    if not isinstance(raw_items, list):
        return []
    out = []
    for idx, raw in enumerate(raw_items, start=1):
        if not isinstance(raw, dict):
            continue
        item = {
            "NO": clean_cell_value(raw.get("NO") or (f"{section_no}.{idx}" if section_no else "")),
            "SOURCE_NAME": clean_cell_value(raw.get("SOURCE_NAME")),
            "TITLE": clean_title(raw.get("TITLE"), raw.get("SOURCE_NAME")),
            "URL": clean_cell_value(raw.get("URL")),
            "PUBLISHED_MM_DD": clean_mm_dd(raw.get("PUBLISHED_MM_DD")),
            "SUMMARY_BULLET_1": clean_cell_value(raw.get("SUMMARY_BULLET_1")),
            "SUMMARY_BULLET_2": clean_cell_value(raw.get("SUMMARY_BULLET_2")),
            "SUMMARY_BULLET_3": clean_cell_value(raw.get("SUMMARY_BULLET_3")),
        }
        out.append(item)
    return out


def fallback_report_data(items: list[collector.Item]) -> dict:
    groups = {
        "BANK_SECTION": [x for x in items if "뉴스" in x.category or "은행" in x.category or "지주" in x.category],
        "OTHER_ITEMS": [x for x in items if x.category == "그외"],
        "GOVERNMENT_SECTION": [x for x in items if "기관" in x.category],
        "FINANCIAL_RESEARCH_SECTION": [x for x in items if "연구" in x.category],
    }
    return {
        "BANK_SECTION": {
            "ITEMS": fallback_items(groups["BANK_SECTION"], "1"),
            "OTHER_ITEMS": fallback_items(groups["OTHER_ITEMS"], ""),
        },
        "GOVERNMENT_SECTION": {"ITEMS": fallback_items(groups["GOVERNMENT_SECTION"], "2")},
        "FINANCIAL_RESEARCH_SECTION": {"ITEMS": fallback_items(groups["FINANCIAL_RESEARCH_SECTION"], "3")},
    }


def fallback_items(items: list[collector.Item], section_no: str) -> list[dict]:
    out = []
    for idx, item in enumerate(items, start=1):
        url = item.original_url or item.url
        out.append(
            {
                "NO": f"{section_no}.{idx}",
                "SOURCE_NAME": item.source_name,
                "TITLE": clean_title(item.title, item.source_name),
                "URL": url,
                "PUBLISHED_MM_DD": item.published_mm_dd,
                "SUMMARY_BULLET_1": "",
                "SUMMARY_BULLET_2": "",
                "SUMMARY_BULLET_3": "",
            }
        )
    return out


def clean_cell_value(value) -> str:
    return "" if value is None else re.sub(r"\s+", " ", str(value)).strip()


def clean_mm_dd(value) -> str:
    text = clean_cell_value(value)
    match = re.fullmatch(r"0?(\d{1,2})[.\-/]0?(\d{1,2})", text)
    if not match:
        return ""
    month = int(match.group(1))
    day = int(match.group(2))
    if not (1 <= month <= 12 and 1 <= day <= 31):
        return ""
    return f"{month}.{day}"


def clean_title(value, source_name: str = "") -> str:
    text = clean_cell_value(value)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"\s+[가-힣]{2,4}(?:\s*,\s*[가-힣]{2,4}){0,5}\s+20\d{2}[.\-/]\d{1,2}[.\-/]\d{1,2}\s+조회수\s*\d+\s*$", "", text)
    text = re.sub(r"\s+20\d{2}[.\-/]\d{1,2}[.\-/]\d{1,2}\s+조회수\s*\d+\s*$", "", text)
    text = re.sub(r"\s+(작성자|등록일|작성일|조회수)\s*[:：]?\s*.*$", "", text)
    source_name = clean_cell_value(source_name)
    if source_name:
        suffix = f" - {source_name}"
        while text.endswith(suffix):
            text = text[: -len(suffix)].strip()
    return text.lstrip("▣").strip()


def row_for_client(item: collector.Item) -> dict:
    row = item.row()
    row["title"] = clean_title(row.get("title", ""), row.get("source_name", ""))
    row["published_date"] = collector.normalize_date(row.get("published_date", "")) or row.get("published_date", "")
    raw_keywords = []
    if isinstance(item.extra, dict):
        raw_keywords = item.extra.get("news_keywords") or []
    if isinstance(raw_keywords, list):
        keywords = [str(keyword).strip() for keyword in raw_keywords if str(keyword or "").strip()]
    else:
        keywords = []
    row["keywords"] = " AND ".join(keywords)
    return row

def populate_report_tables(doc: Document, data: dict) -> None:
    tables = resolve_report_tables(doc)
    if not tables:
        return
    fill_summary_table(tables["bank_summary"], data["BANK_SECTION"]["ITEMS"], keep_min=1)
    fill_other_summary_table(tables["other_summary"], data["BANK_SECTION"]["OTHER_ITEMS"])
    fill_summary_table(tables["agency_summary"], data["GOVERNMENT_SECTION"]["ITEMS"], keep_min=1)
    fill_summary_table(tables["research_summary"], data["FINANCIAL_RESEARCH_SECTION"]["ITEMS"], keep_min=1)
    fill_detail_table(tables["bank_detail"], data["BANK_SECTION"]["ITEMS"], keep_min=1)
    fill_other_detail_table(tables["other_detail"], data["BANK_SECTION"]["OTHER_ITEMS"])
    fill_detail_table(tables["agency_detail"], data["GOVERNMENT_SECTION"]["ITEMS"], keep_min=1)
    fill_detail_table(tables["research_detail"], data["FINANCIAL_RESEARCH_SECTION"]["ITEMS"], keep_min=1)


def resolve_report_tables(doc: Document) -> dict:
    tables = {}
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        first = table.rows[0].cells[0].text if table.rows and table.rows[0].cells else ""
        first_key = re.sub(r"\s+", "", first)
        has_summary_placeholders = "{{NO}}" in text and "{{SOURCE_NAME}}" in text and "{{SUMMARY_BULLET" not in text
        has_detail_placeholders = "{{SUMMARY_BULLET" in text
        if has_summary_placeholders:
            if "1. 은행/지주사" in text:
                tables["bank_summary"] = table
            elif "2. 국가기관" in text:
                tables["agency_summary"] = table
            elif "3. 금융연구소" in text:
                tables["research_summary"] = table
        elif first_key == "그외" and "{{TITLE}}" in text and "{{SUMMARY_BULLET" not in text:
            tables["other_summary"] = table
        elif has_detail_placeholders:
            if "1. 은행/지주사" in first:
                tables["bank_detail"] = table
            elif first_key == "그외":
                tables["other_detail"] = table
            elif "2. 국가기관" in first:
                tables["agency_detail"] = table
            elif "3. 금융연구소" in first:
                tables["research_detail"] = table
    required = {
        "bank_summary", "other_summary", "agency_summary", "research_summary",
        "bank_detail", "other_detail", "agency_detail", "research_detail",
    }
    if required.issubset(tables):
        return tables
    return {}


def fill_summary_table(table, items: list[dict], keep_min: int = 1) -> None:
    desired = max(len(items), keep_min)
    data_start = summary_data_start_row(table)
    adjust_table_rows(table, header_rows=data_start, desired_data_rows=desired)
    set_fixed_table_widths(table, [650, 1600, 7550])
    rows = items + [blank_item()] * (desired - len(items))
    for row, item in zip(table.rows[data_start:], rows):
        cells = row.cells
        title_index = len(cells) - 1 if len(cells) >= 5 else 2
        set_cell_text(cells[0], item.get("NO", ""))
        set_cell_text(cells[1], item.get("SOURCE_NAME", ""))
        set_cell_text(cells[title_index], summary_title(item))


def summary_data_start_row(table) -> int:
    for index, row in enumerate(table.rows):
        if any("{{NO}}" in cell.text for cell in row.cells):
            return index
    return 2


def fill_other_summary_table(table, items: list[dict]) -> None:
    adjust_table_rows(table, header_rows=0, desired_data_rows=1)
    set_fixed_table_widths(table, [650, 1600, 7550])
    row = table.rows[0]
    set_cell_text(row.cells[0], "그\n외")
    set_cell_text(row.cells[1], "\n".join(item.get("SOURCE_NAME", "") for item in items))
    set_cell_text(row.cells[2], "\n".join(summary_title(item) for item in items))
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def fill_detail_table(table, items: list[dict], keep_min: int = 1) -> None:
    desired = max(len(items), keep_min)
    adjust_detail_rows(table, desired)
    set_fixed_table_widths(table, [650, 9150])
    rows = items + [blank_item()] * (desired - len(items))
    for idx, item in enumerate(rows):
        title_row = table.rows[1 + idx * 2]
        bullet_row = table.rows[2 + idx * 2]
        no = item.get("NO", "")
        set_cell_text(title_row.cells[0], no)
        set_cell_text(title_row.cells[1], detail_title(item))
        set_cell_text(bullet_row.cells[0], "")
        set_cell_text(bullet_row.cells[1], detail_summary(item))
        merge_first_column_pair(title_row, bullet_row)


def fill_other_detail_table(table, items: list[dict]) -> None:
    desired = max(len(items), 1)
    adjust_other_detail_rows(table, desired)
    set_fixed_table_widths(table, [650, 9150])
    rows = items + [blank_item()] * (desired - len(items))
    for idx, item in enumerate(rows):
        title_row = table.rows[idx * 2]
        bullet_row = table.rows[idx * 2 + 1]
        set_cell_text(title_row.cells[0], "그\n외" if idx == 0 else "")
        set_cell_text(title_row.cells[1], detail_title(item))
        set_cell_text(bullet_row.cells[0], "")
        set_cell_text(bullet_row.cells[1], detail_summary(item))
    for r_idx, row in enumerate(table.rows):
        if row.cells:
            set_vertical_merge(row.cells[0], restart=(r_idx == 0))


def adjust_table_rows(table, header_rows: int, desired_data_rows: int) -> None:
    while len(table.rows) < header_rows + desired_data_rows:
        clone_row(table, table.rows[-1])
    while len(table.rows) > header_rows + desired_data_rows:
        delete_row(table, table.rows[-1])


def adjust_detail_rows(table, desired_items: int) -> None:
    if len(table.rows) < 3:
        return
    title_row_xml = deepcopy(table.rows[1]._tr)
    bullet_row_xml = deepcopy(table.rows[2]._tr)
    remove_vertical_merges(title_row_xml)
    remove_vertical_merges(bullet_row_xml)
    while len(table.rows) > 1:
        delete_row(table, table.rows[-1])
    for _ in range(desired_items):
        table._tbl.append(deepcopy(title_row_xml))
        table._tbl.append(deepcopy(bullet_row_xml))


def adjust_other_detail_rows(table, desired_items: int) -> None:
    if len(table.rows) < 2:
        return
    title_row_xml = deepcopy(table.rows[0]._tr)
    bullet_row_xml = deepcopy(table.rows[1]._tr)
    remove_vertical_merges(title_row_xml)
    remove_vertical_merges(bullet_row_xml)
    while len(table.rows) > 0:
        delete_row(table, table.rows[-1])
    for _ in range(desired_items):
        table._tbl.append(deepcopy(title_row_xml))
        table._tbl.append(deepcopy(bullet_row_xml))


def clone_row(table, row) -> None:
    table._tbl.append(deepcopy(row._tr))


def delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def remove_vertical_merges(row_xml) -> None:
    for tc_pr in row_xml.iter(qn("w:tcPr")):
        for v_merge in list(tc_pr.findall(qn("w:vMerge"))):
            tc_pr.remove(v_merge)


def merge_first_column_pair(title_row, bullet_row) -> None:
    if not title_row.cells or not bullet_row.cells:
        return
    set_vertical_merge(title_row.cells[0], restart=True)
    set_vertical_merge(bullet_row.cells[0], restart=False)


def set_vertical_merge(cell, restart: bool) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for v_merge in list(tc_pr.findall(qn("w:vMerge"))):
        tc_pr.remove(v_merge)
    v_merge = OxmlElement("w:vMerge")
    if restart:
        v_merge.set(qn("w:val"), "restart")
    tc_pr.append(v_merge)


def set_fixed_table_widths(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        seen = set()
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths) or id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


URL_RE = re.compile(r"https?://[^\s)]+")


def add_hyperlink(paragraph, url: str, text: str | None = None) -> None:
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_hyperlinks(paragraph, line: str) -> None:
    pos = 0
    for match in URL_RE.finditer(line):
        if match.start() > pos:
            paragraph.add_run(line[pos:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        pos = match.end()
    if pos < len(line):
        paragraph.add_run(line[pos:])


MALGUN_FONT = "맑은 고딕"


def set_run_font_malgun(run) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:eastAsia", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), MALGUN_FONT)


def apply_malgun_to_table(table) -> None:
    for row in table.rows:
        seen = set()
        for cell in row.cells:
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font_malgun(run)


def capture_first_run_rpr(paragraph):
    for r in paragraph._p.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            return deepcopy(rPr)
    return None


def apply_rpr_to_paragraph_runs(paragraph, template_rPr) -> None:
    if template_rPr is None:
        return
    for r in paragraph._p.iter(qn("w:r")):
        existing = r.find(qn("w:rPr"))
        if existing is not None:
            r.remove(existing)
        r.insert(0, deepcopy(template_rPr))


def set_cell_text(cell, text: str) -> None:
    text = text or ""
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    remove_paragraph_numbering(first)
    template_rPr = capture_first_run_rpr(first)
    for child in list(first._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            first._p.remove(child)
    lines = str(text).split("\n")
    line0_end_index = None
    for index, line in enumerate(lines):
        if index:
            if line0_end_index is None:
                line0_end_index = len(first._p)
            first.add_run().add_break()
        add_text_with_hyperlinks(first, line)
    if line0_end_index is None:
        line0_end_index = len(first._p)
    if template_rPr is not None:
        for child in list(first._p)[:line0_end_index]:
            if child.tag != qn("w:r"):
                continue
            existing = child.find(qn("w:rPr"))
            if existing is not None:
                child.remove(existing)
            child.insert(0, deepcopy(template_rPr))


def remove_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def blank_item() -> dict:
    return {
        "NO": "",
        "SOURCE_NAME": "",
        "TITLE": "",
        "URL": "",
        "PUBLISHED_MM_DD": "",
        "SUMMARY_BULLET_1": "",
        "SUMMARY_BULLET_2": "",
        "SUMMARY_BULLET_3": "",
    }


def summary_title(item: dict) -> str:
    title = item.get("TITLE", "")
    if not title:
        return ""
    mmdd = item.get("PUBLISHED_MM_DD", "")
    return f"▣ {title}" + (f" ({mmdd})" if mmdd else "")


def detail_title(item: dict) -> str:
    title = item.get("TITLE", "")
    if not title:
        return ""
    url = item.get("URL", "")
    meta = " / ".join(x for x in [item.get("SOURCE_NAME", ""), item.get("PUBLISHED_MM_DD", "")] if x)
    second = f"{url} ({meta})" if url and meta else url or (f"({meta})" if meta else "")
    return f"▣ {title}" + (f"\n{second}" if second else "")


def detail_summary(item: dict) -> str:
    bullets = [
        item.get("SUMMARY_BULLET_1", ""),
        item.get("SUMMARY_BULLET_2", ""),
        item.get("SUMMARY_BULLET_3", ""),
    ]
    return "\n".join(format_detail_bullet(b) for b in bullets if clean_cell_value(b))


def format_detail_bullet(value: str) -> str:
    text = clean_cell_value(value)
    text = re.sub(r"^(?:[\-•·–—]\s*)+", "", text).strip()
    return f"- {text}" if text else ""


def detail_block(item: dict) -> str:
    parts = [detail_title(item), detail_summary(item)]
    return "\n".join(part for part in parts if part)


def compact_url(url: str, limit: int = 76) -> str:
    url = clean_cell_value(url)
    if not url:
        return ""
    parsed = urllib.parse.urlparse(url)
    if parsed.netloc:
        value = parsed.netloc + parsed.path
        if parsed.query:
            value += "?..."
    else:
        value = url
    if len(value) <= limit:
        return value
    return value[: max(0, limit - 3)].rstrip("/") + "..."


def stringify_template_value(value) -> str:
    if isinstance(value, (list, dict)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return "" if value is None else str(value)


def format_detail_items(generated, fallback_items: list[collector.Item]) -> str:
    source = generated if isinstance(generated, list) else []
    lines = []
    if source:
        for idx, item in enumerate(source, start=1):
            if not isinstance(item, dict):
                lines.append(str(item))
                continue
            lines.extend(
                [
                    f"{idx}. {item.get('TITLE') or item.get('title') or ''}",
                    f"{item.get('URL') or item.get('url') or ''} ({item.get('SOURCE_NAME') or item.get('source_name') or ''} / {item.get('PUBLISHED_MM_DD') or ''})",
                    f"- {item.get('SUMMARY_BULLET_1') or ''}",
                    f"- {item.get('SUMMARY_BULLET_2') or ''}",
                    f"- {item.get('SUMMARY_BULLET_3') or ''}",
                    "",
                ]
            )
    else:
        for idx, item in enumerate(fallback_items, start=1):
            lines.extend([
                f"{idx}. {item.title}",
                f"{item.original_url or item.url} ({item.source_name} / {item.published_mm_dd})",
                "",
            ])
    return "\n".join(lines).strip()


def replace_docx_placeholders(doc: Document, values: dict) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, values)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, values)


def replace_in_paragraph(paragraph, values: dict) -> None:
    text = paragraph.text
    if "{{" not in text:
        return
    for key, value in values.items():
        text = text.replace("{{" + key + "}}", str(value))
    template_rPr = capture_first_run_rpr(paragraph)
    paragraph.clear()
    add_text_with_pageref_fields(paragraph, text)
    apply_rpr_to_paragraph_runs(paragraph, template_rPr)


def add_text_with_pageref_fields(paragraph, text: str) -> None:
    markers = {
        "__PAGEREF_BANK_SECTION__": "bank_section",
        "__PAGEREF_AGENCY_SECTION__": "agency_section",
        "__PAGEREF_RESEARCH_SECTION__": "research_section",
    }
    pattern = "(" + "|".join(re.escape(marker) for marker in markers) + ")"
    for part in re.split(pattern, text):
        if not part:
            continue
        bookmark = markers.get(part)
        if bookmark:
            add_pageref_field(paragraph, bookmark)
        else:
            paragraph.add_run(part)


def add_pageref_field(paragraph, bookmark_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    placeholder = OxmlElement("w:t")
    placeholder.text = "0"
    run._r.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_section_bookmarks(doc: Document) -> None:
    targets = {
        "1. 은행/지주사": "bank_section",
        "2. 국가기관": "agency_section",
        "3. 금융연구소": "research_section",
    }
    candidates: dict[str, list] = {bookmark: [] for bookmark in targets.values()}
    for paragraph in iter_doc_paragraphs(doc):
        text = paragraph.text.strip()
        for label, bookmark in targets.items():
            if text == label:
                candidates[bookmark].append(paragraph)
    bookmark_id = 100
    for bookmark, paragraphs in candidates.items():
        if not paragraphs:
            continue
        add_bookmark_to_paragraph(paragraphs[-1], bookmark, bookmark_id)
        bookmark_id += 1


def iter_doc_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def add_bookmark_to_paragraph(paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._p
    if not paragraph.runs:
        paragraph.add_run("")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def suppress_word_field_update(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is not None:
        settings.remove(update_fields)


def is_collect_run_cancelled(run_dir: Path) -> bool:
    try:
        key = str(run_dir.resolve())
    except Exception:
        key = str(run_dir)
    with STATE_LOCK:
        return key in STATE.get("cancelled_run_dirs", set())


def is_collect_job_cancelled(job_id: str) -> bool:
    with STATE_LOCK:
        return job_id in STATE.get("cancelled_collect_jobs", set())


def is_generate_job_cancelled(job_id: str) -> bool:
    with STATE_LOCK:
        return job_id in STATE.get("cancelled_generate_jobs", set())


class GenerateCancelled(RuntimeError):
    pass


def cleanup_cancelled_collect_run(run_dir: Path) -> None:
    try:
        if run_dir.exists() and run_dir.is_dir():
            shutil.rmtree(run_dir)
    except Exception:
        pass
    try:
        key = str(run_dir.resolve())
    except Exception:
        key = str(run_dir)
    with STATE_LOCK:
        STATE.get("cancelled_run_dirs", set()).discard(key)
        if same_path(str(STATE.get("pending_run_dir") or ""), run_dir):
            STATE["pending_run_dir"] = None
        if str(STATE.get("active_collect_job_id") or ""):
            STATE["active_collect_job_id"] = ""


def run_collect_job(job_id: str, data: dict) -> None:
    def emit(event_type: str, payload: dict | None = None) -> None:
        payload = dict(payload or {})
        payload["type"] = event_type
        with STATE_LOCK:
            job = JOBS.get(job_id)
        if job:
            job["queue"].put(payload)

    try:
        news_ranges = parse_news_ranges(data)
        first_news_range = news_ranges[0]
        news_start = first_news_range["start"]
        news_end = first_news_range["end"]
        agency_start = parse_iso_date(data["agency_start"])
        agency_end = parse_iso_date(data["agency_end"])
        research_start = parse_iso_date(data["research_start"])
        research_end = parse_iso_date(data["research_end"])
        agency_max = int(data.get("agency_max") or 10)
        research_max = int(data.get("research_max") or 10)
        news_keywords = base_keywords_from_payload(data)
        primary_news_keywords = first_news_range.get("keywords") or news_keywords
        include_news = bool(data.get("include_news", True))
        include_agency = bool(data.get("include_agency", True))
        include_research = bool(data.get("include_research", True))
        with STATE_LOCK:
            existing_items = list(STATE.get("items") or [])
            current_run = str(STATE.get("run_dir") or "")
        run_dir = resolve_existing_run_dir(current_run) or collector.make_run_dir()
        save_news_keywords(primary_news_keywords)
        save_run_keywords(run_dir, primary_news_keywords)
        with STATE_LOCK:
            STATE["pending_run_dir"] = str(run_dir)
            STATE["report_path"] = ""
            if job_id in STATE.get("cancelled_collect_jobs", set()):
                STATE.setdefault("cancelled_run_dirs", set()).add(str(run_dir.resolve()))
        emit("started", {"run_dir": str(run_dir)})
        if is_collect_job_cancelled(job_id):
            cleanup_cancelled_collect_run(run_dir)
            return

        collected: list[collector.Item] = []
        if include_news:
            for index, range_data in enumerate(news_ranges, start=1):
                emit(
                    "source_start",
                    {
                        "source": f"뉴스 수집 구간 {index}/{len(news_ranges)}",
                        "group": "news",
                    },
                )
                range_items = collector.collect_by_ranges(
                    range_data["start"],
                    range_data["end"],
                    agency_start,
                    agency_end,
                    research_start,
                    research_end,
                    max_per_source=None,
                    dry_run=False,
                    news_bank_max=range_data["bank_max"],
                    news_other_max=range_data["other_max"],
                    agency_max=agency_max,
                    research_max=research_max,
                    output_dir=run_dir,
                    progress=emit,
                    news_keywords=range_data.get("keywords") or news_keywords,
                    include_news=True,
                    include_agency=False,
                    include_research=False,
                )
                collected.extend(range_items)
                emit(
                    "source_done",
                    {
                        "source": f"뉴스 수집 구간 {index}/{len(news_ranges)}",
                        "group": "news",
                        "count": len(range_items),
                    },
                )
        if include_agency or include_research:
            collected.extend(
                collector.collect_by_ranges(
                    news_start,
                    news_end,
                    agency_start,
                    agency_end,
                    research_start,
                    research_end,
                    max_per_source=None,
                    dry_run=False,
                    news_bank_max=first_news_range["bank_max"],
                    news_other_max=first_news_range["other_max"],
                    agency_max=agency_max,
                    research_max=research_max,
                    output_dir=run_dir,
                    progress=emit,
                    news_keywords=news_keywords,
                    include_news=False,
                    include_agency=include_agency,
                    include_research=include_research,
                )
                )
        items = collector.dedupe(existing_items + collected)
        if is_collect_run_cancelled(run_dir):
            cleanup_cancelled_collect_run(run_dir)
            return
        collector.write_outputs(items, output_dir=run_dir)
        if is_collect_run_cancelled(run_dir):
            cleanup_cancelled_collect_run(run_dir)
            return
        payload = []
        try:
            run_key = str(run_dir.resolve())
        except Exception:
            run_key = str(run_dir)
        with STATE_LOCK:
            if run_key in STATE.get("cancelled_run_dirs", set()):
                return
            STATE["items"] = items
            STATE["run_dir"] = str(run_dir)
            STATE["report_path"] = ""
            STATE["pending_run_dir"] = None
            if STATE.get("active_collect_job_id") == job_id:
                STATE["active_collect_job_id"] = ""
            STATE.get("cancelled_collect_jobs", set()).discard(job_id)
            for idx, item in enumerate(items):
                row = row_for_client(item)
                row["id"] = idx
                payload.append(row)
        emit("complete", {"items": payload, "run_dir": str(run_dir), "run_name": run_dir.name, "keywords": primary_news_keywords})
    except Exception as exc:
        traceback.print_exc()
        emit("fatal", {"error": str(exc)})


def write_items(items: list[collector.Item], csv_path: Path, xlsx_path: Path) -> None:
    csv_path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = list(collector.Item("", "", "", "").row().keys())
    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for item in items:
            writer.writerow(row_for_client(item))

    wb = Workbook()
    ws = wb.active
    ws.title = "selected_metadata"
    ws.append(fieldnames)
    for item in items:
        row = row_for_client(item)
        ws.append([row[k] for k in fieldnames])
    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col[:100])
        ws.column_dimensions[col[0].column_letter].width = min(max(max_len + 2, 12), 45)
    wb.save(xlsx_path)


def touch_server_activity() -> None:
    with STATE_LOCK:
        STATE["last_activity"] = time.time()


def _exit_soon(delay: float = 1.2) -> None:
    """응답이 나갈 틈을 준 뒤 프로세스를 끝낸다(업데이트 적용 전용)."""
    time.sleep(delay)
    os._exit(0)


def update_busy_reason() -> str | None:
    """지금 앱을 끊으면 잃는 게 있는가 — 있으면 그 사유.

    업데이트는 프로세스를 죽였다 살리는 일이라, 수집·보고서 생성 도중에 걸면 사용자는 몇 분치
    작업과 OpenAI 호출 비용을 잃는다. 그래서 진행 중이면 거부한다.
    """
    with STATE_LOCK:
        collecting = bool(STATE.get("active_collect_job_id") or STATE.get("pending_run_dir"))
        generating = bool(STATE.get("active_generate_job_id"))
    if collecting:
        return "자료 수집이 진행 중입니다. 끝난 뒤 업데이트해 주세요."
    if generating:
        return "보고서 생성이 진행 중입니다. 끝난 뒤 업데이트해 주세요."
    return None


def server_has_active_jobs() -> bool:
    with STATE_LOCK:
        return bool(
            STATE.get("active_collect_job_id")
            or STATE.get("active_generate_job_id")
            or STATE.get("pending_run_dir")
        )


def start_idle_shutdown_watcher(server: ThreadingHTTPServer) -> None:
    def watch() -> None:
        while True:
            time.sleep(SERVER_IDLE_CHECK_SECONDS)
            with STATE_LOCK:
                idle_for = time.time() - float(STATE.get("last_activity") or time.time())
            if idle_for < SERVER_IDLE_TIMEOUT_SECONDS:
                continue
            if server_has_active_jobs():
                touch_server_activity()
                continue
            print(f"[app] idle for {int(idle_for)}s; shutting down server")
            server.shutdown()
            return

    threading.Thread(target=watch, daemon=True).start()


def main() -> int:
    import socket
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        already_running = s.connect_ex((HOST, PORT)) == 0
    if already_running:
        print(f"이미 실행 중: http://{HOST}:{PORT}")
        webbrowser.open(f"http://{HOST}:{PORT}/")
        return 0
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    start_idle_shutdown_watcher(server)
    print(f"조사연구 도우미: http://{HOST}:{PORT}")
    if os.environ.get("RRA_NO_BROWSER") != "1":
        threading.Timer(0.8, lambda: webbrowser.open(f"http://{HOST}:{PORT}/")).start()
    server.serve_forever()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
